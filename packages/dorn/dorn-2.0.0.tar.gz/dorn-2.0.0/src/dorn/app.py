__program_name__ = "Dorn"
__version__ = "2.0.0"
__release_date__ = "9 October 2025"
__homepage__ = "https://github.com/SAMI-Medical-Physics/dorn"
__author__ = "Jake Forster"
__author_email__ = "jake.forster@sa.gov.au"
__copyright_year__ = "2022-2025"
__copyright_owner__ = "South Australia Medical Imaging"
__license__ = "MIT"

import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
from tkinter import filedialog

from datetime import datetime, date, timedelta
import numpy as np
import pandas as pd

from platformdirs import user_config_dir

import warnings

import os.path
import xmltodict
import collections
import getpass
from pathlib import Path

import importlib.metadata
import platform


try:
    GLOWGREEN_VERSION = importlib.metadata.version("glowgreen")
except importlib.metadata.PackageNotFoundError:
    GLOWGREEN_VERSION = "0"


WINDOWS_OS = platform.system() == "Windows"

# To help restriction periods for once-off patterns converge.
MAX_ADMINISTERED_ACTIVITY_MBQ = 1e9  # 1 PBq

ASSET_DIR = Path(__file__).resolve().parent


def func_exp(t, a, b):
    return a * np.exp(-b * t)


def func_biexp(t, a, b, c, d):
    return a * (b * np.exp(-c * t) + (1 - b) * np.exp(-d * t))


def func_biexp_root(t, a, b, c, d, val):
    return func_biexp(t, a, b, c, d) - val


def str2datetime(astr):
    """Return a datetime.datetime from a string of the form %Y%m%d%H%M,
    according to format codes of 1989 C standard.
    """
    return datetime(
        year=int(astr[0:4]),
        month=int(astr[4:6]),
        day=int(astr[6:8]),
        hour=int(astr[8:10]),
        minute=int(astr[10:12]),
    )


def str2date(astr):
    """Return a datetime.date from a string of the form %Y%m%d,
    according to format codes of 1989 C standard.
    """
    return date(year=int(astr[0:4]), month=int(astr[4:6]), day=int(astr[6:8]))


def datetime2str(dt):
    """Return a string of the form %Y%m%d%H%M from a datetime.datetime,
    according to format codes of 1989 C standard.
    """
    return (
        str(dt.year).zfill(4)
        + str(dt.month).zfill(2)
        + str(dt.day).zfill(2)
        + str(dt.hour).zfill(2)
        + str(dt.minute).zfill(2)
    )


def date2str(d):
    """Return a string of the form %Y%m%d from a datetime.date,
    according to format codes of 1989 C standard.
    """
    return str(d.year).zfill(4) + str(d.month).zfill(2) + str(d.day).zfill(2)


def word_available(windows_os):
    if not windows_os:
        raise NotImplementedError("This feature is only supported on Windows.")
    import win32com.client  # from pywin32

    try:
        win32com.client.Dispatch("Word.Application")
    except:
        # The default exception message is unhelpful.
        raise FileNotFoundError("MS Word was not found.")


def docx2pdf(path_in, path_out):
    # Adapted from https://github.com/AlJohri/docx2pdf
    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    wdFormatPDF = 17
    docx_filepath = Path(path_in).resolve()
    pdf_filepath = Path(path_out).resolve()
    doc = word.Documents.Open(str(docx_filepath))
    try:
        doc.SaveAs(str(pdf_filepath), FileFormat=wdFormatPDF)
    except:
        raise
    finally:
        doc.Close(0)


class RestrictionsWindow:
    def __init__(self, gui):
        from glowgreen import cs_patterns

        self.chks = []
        if gui.odict["data"]["restrictions"]:
            self.df = pd.DataFrame.from_dict(
                gui.odict["data"]["restrictions"]["restriction"]
            )

            # Values are in strings first-thing after reading XML file
            self.df["theta"] = self.df["theta"].map(
                lambda x: (
                    [float(i) for i in x]
                    if isinstance(x, (list, np.ndarray))
                    else float(x)
                )
            )
            self.df["c"] = self.df["c"].map(
                lambda x: (
                    [float(i) for i in x]
                    if isinstance(x, (list, np.ndarray))
                    else float(x)
                )
            )
            self.df["d"] = self.df["d"].map(
                lambda x: (
                    [float(i) for i in x]
                    if isinstance(x, (list, np.ndarray))
                    else float(x)
                )
            )
            self.df["dose_constraint"] = self.df["dose_constraint"].astype(float)
            self.df["per_episode"] = self.df["per_episode"].astype(int)

            if "restriction_period" not in self.df.columns:
                self.compute_restrictions(gui)
        else:
            self.df = cs_patterns()
            self.compute_restrictions(gui)

    def compute_restrictions(self, gui):
        from glowgreen import (
            Clearance_1m,
            restrictions_for,
        )

        admin_datetime = str2datetime(
            gui.odict["data"]["administration_details"]["administration_datetime"]
        )

        model = gui.odict["data"]["clearance_data"]["curve_fit"]["model"]
        meaningful_parameters = [
            float(a)
            for a in gui.odict["data"]["clearance_data"]["curve_fit"][
                "meaningful_parameters"
            ].values()
        ]
        if gui.therapy_options_df.loc[
            gui.odict["data"]["patient_details"]["type_therapy"],
            "generic_clearance",
        ]:
            measurement_distance = 1.0
        else:
            measurement_distance = float(
                gui.odict["data"]["clearance_data"]["measurement_distance"]
            )
        cfit = Clearance_1m(model, meaningful_parameters, measurement_distance)

        num_treatments_in_year = float(
            gui.odict["data"]["patient_details"]["num_treatments_in_year"]
        )

        self.df = restrictions_for(
            self.df,
            cfit,
            num_treatments_in_year,
            admin_datetime=admin_datetime,
        )

        with warnings.catch_warnings():
            # The fsolve used for onceoff patterns doesn't converge
            # for very large activities.  In that case, fail loudly.
            warnings.simplefilter("error", RuntimeWarning)
            try:
                self.df = restrictions_for(
                    self.df,
                    cfit,
                    num_treatments_in_year,
                    admin_datetime=admin_datetime,
                )
            except Exception as e:
                messagebox.showerror(
                    "Error",
                    f"Failed to calculate the restriction periods.  The RuntimeWarning message is shown below.\n\n{e}",
                )
                raise
        self.df["datetime_end"] = self.df["datetime_end"].map(datetime2str)

        gui.odict["data"]["glowgreen_version"] = GLOWGREEN_VERSION

    # Display restrictions in window
    def make_display(self, gui, window):
        for i in range(self.df.shape[0]):
            if self.df["per_episode"].iloc[i]:
                tk.Label(window, text=self.df["name"].iloc[i], fg="#0072B2").grid(
                    row=i + 2
                )
            else:
                tk.Label(window, text=self.df["name"].iloc[i]).grid(row=i + 2)
            datetime_end = str2datetime(self.df["datetime_end"].iloc[i])
            datetime_end_display = (
                (datetime_end + timedelta(hours=1))
                if datetime_end.minute != 0
                else datetime_end
            )
            datetime_end_str = (
                datetime_end_display.strftime("%d %b %Y, %I %p")
                .lstrip("0")
                .replace(", 0", ", ")
            )
            tk.Label(window, text=datetime_end_str).grid(
                row=i + 2, column=1, padx=(5, 5)
            )

            self.chks.append(tk.IntVar())
            tk.Checkbutton(window, variable=self.chks[-1]).grid(row=i + 2, column=2)

            if "applies" in self.df.columns:
                self.chks[-1].set(int(self.df.loc[i, "applies"]))
            else:
                self.chks[-1].set(0)

            tk.Button(
                window,
                text="View",
                command=lambda i=i: self.view_restriction(gui, i),
            ).grid(row=i + 2, column=3, padx=(5, 0))

    # View button alongside each restriction
    def view_restriction(self, gui, i):
        from glowgreen import (
            Clearance_1m,
            ContactPatternRepeating,
            ContactPatternOnceoff,
        )

        admin_datetime = str2datetime(
            gui.odict["data"]["administration_details"]["administration_datetime"]
        )

        model = gui.odict["data"]["clearance_data"]["curve_fit"]["model"]
        meaningful_parameters = [
            float(a)
            for a in gui.odict["data"]["clearance_data"]["curve_fit"][
                "meaningful_parameters"
            ].values()
        ]
        if gui.therapy_options_df.loc[
            gui.odict["data"]["patient_details"]["type_therapy"],
            "generic_clearance",
        ]:
            measurement_distance = 1.0
        else:
            measurement_distance = float(
                gui.odict["data"]["clearance_data"]["measurement_distance"]
            )
        cfit = Clearance_1m(model, meaningful_parameters, measurement_distance)

        num_treatments_in_year = float(
            gui.odict["data"]["patient_details"]["num_treatments_in_year"]
        )

        pd_series = self.df.iloc[i]
        if pd_series["pattern_type"] == "repeating":
            cpat = ContactPatternRepeating(
                pd_series["theta"], pd_series["c"], pd_series["d"]
            )
        elif pd_series["pattern_type"] == "onceoff":
            cpat = ContactPatternOnceoff(
                pd_series["theta"], pd_series["c"], pd_series["d"]
            )

        if pd_series["per_episode"] == 0:
            dose_constraint = pd_series["dose_constraint"] / num_treatments_in_year
        elif pd_series["per_episode"] == 1:
            dose_constraint = pd_series["dose_constraint"]

        cpat.plot(
            name=pd_series["name"],
            cfit=cfit,
            dose_constraint=dose_constraint,
            admin_datetime=admin_datetime,
        )

    # Submit button in window
    def submit_restrictions(self, gui, window):
        self.df["applies"] = [e.get() for e in self.chks]

        gui.odict["data"]["restrictions"]["restriction"] = self.df.to_dict(
            "records", into=collections.OrderedDict
        )
        window.withdraw()

        gui.odict["data"]["reports_generated"] = "0"

        gui.unsaved_data = True

        if int(gui.odict["data"]["patient_finished"]):
            gui.odict["data"]["patient_finished"] = "0"
            gui.odict["data"]["patient_finished_by"] = "0"
            gui.viewing_completed_patient_label.place_forget()

        filepath = gui.filepath
        if filepath is not None:
            gui.root.title(f"*{os.path.basename(filepath)} - {__program_name__}")

        gui.update_buttons()


class Gui:
    def __init__(
        self, add_therapy_measured=None, add_therapy_generic=None, write_info=False
    ):
        """Start GUI program."""
        self.SOFTWARE_ICON = str(ASSET_DIR / "shovel.ico")
        self.GUI_BKGD = str(ASSET_DIR / "gui_background.png")
        self.FILENAME_SETTINGS = (
            Path(user_config_dir(__program_name__)) / "settings.xml"
        )
        if Path("settings.xml").is_file():
            # Use the settings.xml in the current directory instead.
            self.FILENAME_SETTINGS = "settings.xml"
        self.SEX_OPTIONS = ["Male", "Female", "Other"]
        self.N_MEASUREMENTS_MAX = 15

        radionuclide_options = [
            self.radionuclide_dict(
                "I-131",
                8.0252 * 24.0,
                activity_outpatient=600,
                specific_dose_rate_1m=7.647e-2,
            ),
            self.radionuclide_dict(
                "Lu-177", 6.6443 * 24.0, specific_dose_rate_1m=7.636e-3
            ),
            self.radionuclide_dict(
                "In-111",
                2.8047 * 24,
                activity_outpatient=400,
                specific_dose_rate_1m=1.356e-1,
            ),
            self.radionuclide_dict("P-32", 14.268 * 24, activity_outpatient=1200),
            self.radionuclide_dict(
                "Re-188",
                17.005,
                activity_outpatient=4000,
                specific_dose_rate_1m=1.094e-2,
            ),
            self.radionuclide_dict(
                "Sm-153",
                46.284,
                activity_outpatient=4000,
                specific_dose_rate_1m=2.440e-2,
            ),
            self.radionuclide_dict("Sr-89", 50.563 * 24, activity_outpatient=300),
            self.radionuclide_dict("Y-90", 64.05, activity_outpatient=4000),
            self.radionuclide_dict("Cu-67", 61.83, specific_dose_rate_1m=2.363e-2),
            self.radionuclide_dict("Zr-89", 78.41, specific_dose_rate_1m=2.655e-1),
            self.radionuclide_dict(
                "I-124", 4.1760 * 24.0, specific_dose_rate_1m=2.050e-1
            ),
            self.radionuclide_dict(
                "I-125", 59.400 * 24.0, specific_dose_rate_1m=7.432e-2
            ),
            self.radionuclide_dict("Tc-99m", 6.0072, specific_dose_rate_1m=3.317e-2),
        ]
        self.radionuclide_options_df = pd.DataFrame(radionuclide_options)
        self.radionuclide_options_df.set_index("name", inplace=True)

        therapy_options = [
            self.therapy_dict("I-131 thyroid cancer, rhTSH", "I-131", True),
            self.therapy_dict(
                "I-131 thyroid cancer, thyroxine withdrawal", "I-131", True
            ),
            self.therapy_dict("I-131 MIBG", "I-131", True),
            self.therapy_dict(
                "I-131 thyrotoxicosis (generic clearance)",
                "I-131",
                False,
                generic_clearance=True,
                generic_model="biexponential",
                generic_parameters=[0.05, 0.73, 6.7, 139],
            ),
            self.therapy_dict("Lu-177-dotatate", "Lu-177", False),
            self.therapy_dict("Lu-177-PSMA-617", "Lu-177", False),
        ]
        self.therapy_options_df = pd.DataFrame(therapy_options)

        if add_therapy_measured is not None:
            therapy_name, radionuclide, inpatient = add_therapy_measured
            if therapy_name in self.therapy_options_df["name"].to_list():
                raise ValueError(f'Therapy named "{therapy_name}" already exists')
            radionuclide_list = self.radionuclide_options_df.index.to_list()
            if radionuclide not in radionuclide_list:
                raise ValueError(
                    f"Radionuclide {radionuclide} not one of {radionuclide_list}"
                )
            measured_therapy_dict = self.therapy_dict(
                therapy_name, radionuclide, inpatient
            )
            self.therapy_options_df = pd.concat(
                [self.therapy_options_df, pd.DataFrame([measured_therapy_dict])],
                ignore_index=True,
            )

        if add_therapy_generic is not None:
            (
                therapy_name,
                radionuclide,
                inpatient,
                generic_model,
                generic_parameters,
            ) = add_therapy_generic
            if therapy_name in self.therapy_options_df["name"].to_list():
                raise ValueError(f'Therapy named "{therapy_name}" already exists')
            radionuclide_list = self.radionuclide_options_df.index.to_list()
            if radionuclide not in radionuclide_list:
                raise ValueError(
                    f"Radionuclide {radionuclide} not one of {radionuclide_list}"
                )
            model_list = ["exponential", "biexponential"]
            if generic_model not in model_list:
                raise ValueError(f"Model {generic_model} not one of {model_list}")
            if generic_model == "exponential" and len(generic_parameters) != 2:
                raise IndexError(
                    f"Exponential model needs 2 parameters, {len(generic_parameters)} supplied"
                )
            if generic_model == "biexponential" and len(generic_parameters) != 4:
                raise IndexError(
                    f"Biexponential model needs 4 parameters, {len(generic_parameters)} supplied"
                )
            generic_therapy_dict = self.therapy_dict(
                therapy_name,
                radionuclide,
                inpatient,
                generic_clearance=True,
                generic_model=generic_model,
                generic_parameters=generic_parameters,
            )
            self.therapy_options_df = pd.concat(
                [self.therapy_options_df, pd.DataFrame([generic_therapy_dict])],
                ignore_index=True,
            )

        self.therapy_options_df.set_index("name", inplace=True)

        if write_info:
            self.radionuclide_options_df.to_csv("radionuclide_options.csv")
            self.therapy_options_df.to_csv("therapy_options.csv")
            return

        self.previous_data_directory = None
        self.previous_report_directory = None
        self.unsaved_data = False
        self.filepath = None
        self.odict = self.get_new_odict()

        # Create the GUI
        self.root = tk.Tk()
        self.root.geometry("450x450")
        self.root.title(f"Main Menu - {__program_name__}")
        if WINDOWS_OS:
            self.root.resizable(width=False, height=False)
            self.root.iconbitmap(self.SOFTWARE_ICON)

        bkgd_img = tk.PhotoImage(file=self.GUI_BKGD)
        tk.Label(self.root, image=bkgd_img).place(x=-3, y=-20)

        self.info_label = tk.Label(self.root, justify=tk.LEFT)
        self.viewing_completed_patient_label = tk.Label(self.root, bg="yellow")

        tk.Label(
            self.root,
            text=f"{__program_name__} {__version__}",
            bg="white",
            font="Arial 16 bold",
        ).grid(row=0, pady=(55, 0))
        tk.Label(
            self.root,
            text="Close contact restrictions\nfor radionuclide therapy patients",
            bg="white",
            font="Arial 10",
        ).grid(row=1, padx=130, pady=(0, 0))

        self.button_patient = tk.Button(
            self.root, text="Patient Details", command=self.patient_details
        )
        self.button_patient.grid(row=2, pady=(15, 0))

        self.button_administration = tk.Button(
            self.root,
            text="Administration Details",
            command=self.administration_details,
        )
        self.button_administration.grid(row=3, pady=(5, 0))

        self.button_clearance = tk.Button(
            self.root, text="Measured Clearance Data", command=self.clearance_data
        )
        self.button_clearance.grid(row=4, pady=(5, 0))

        self.button_discharge = tk.Button(
            self.root, text="Patient Discharge", command=self.patient_discharge
        )
        self.button_discharge.grid(row=5, pady=(5, 0))

        self.button_restrictions = tk.Button(
            self.root, text="Restrictions", command=self.restrictions_window
        )
        self.button_restrictions.grid(row=6, pady=(5, 0))

        self.button_comments = tk.Button(
            self.root, text="Comments", command=self.comments
        )
        self.button_comments.grid(row=7, pady=(5, 0))

        self.button_reports = tk.Button(
            self.root, text="Generate Reports", command=self.generate_reports
        )
        self.button_reports.grid(row=8, pady=(5, 0))

        self.button_finish = tk.Button(
            self.root, text="Finish", command=self.finish_patient
        )
        self.button_finish.grid(row=9, pady=(5, 0))

        mainmenu = tk.Menu(self.root)
        mainmenu.add_command(label="New", command=self.new_patient)
        mainmenu.add_command(label="Open", command=self.load_patient)
        mainmenu.add_command(label="Save", command=self.save_patient)
        mainmenu.add_command(label="Save As", command=self.save_patient_as)
        mainmenu.add_command(label="Info", command=self.info_box)

        settings = tk.Menu(mainmenu, tearoff=0)
        mainmenu.add_cascade(label="Settings", menu=settings)
        settings.add_command(label="Organisation", command=self.settings_organisation)
        settings.add_command(label="Initial Values", command=self.settings_init_vals)
        settings.add_command(label="Other", command=self.settings_other)
        self.root.config(menu=mainmenu)

        self.update_buttons()

        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)

        # Use hard-coded settings, then try to overwrite them by reading in settings file
        self.set_default_settings()
        if os.path.isfile(self.FILENAME_SETTINGS):
            read_success = False
            try:
                with open(self.FILENAME_SETTINGS) as fd:
                    settings_odict = xmltodict.parse(
                        fd.read(),
                        dict_constructor=collections.OrderedDict,
                        postprocessor=Gui.my_postprocessor_settings,
                    )
                read_success = True
            except Exception as e:
                messagebox.showinfo(
                    "Information",
                    f"Ignoring {self.FILENAME_SETTINGS} due to:\n{e}",
                )

            if read_success:
                try:
                    (
                        self.email,
                        self.url,
                        self.report_logo,
                    ) = Gui.read_organisation_info(settings_odict)
                except Exception as e:
                    messagebox.showinfo(
                        "Information",
                        f"Ignoring organisation email, URL and logo file in {self.FILENAME_SETTINGS} due to:\n{e}",
                    )
                try:
                    self.site_options_df = Gui.read_sites(settings_odict)
                except Exception as e:
                    messagebox.showinfo(
                        "Information",
                        f"Ignoring sites in {self.FILENAME_SETTINGS} due to:\n{e}",
                    )
                try:
                    self.init_vals = self.read_init_vals(settings_odict)
                except Exception as e:
                    messagebox.showinfo(
                        "Information",
                        f"Ignoring 'initial_values' in {self.FILENAME_SETTINGS}:\n    {e}",
                    )
                # 'other_settings' was added in Dorn 1.12.
                # It will remain missing from the settings file until the user submits a setting.
                if "other_settings" in settings_odict["root"]:
                    try:
                        self.other_settings = self.read_other_settings(settings_odict)
                    except Exception as e:
                        messagebox.showinfo(
                            "Information",
                            f"Ignoring 'other_settings' in {self.FILENAME_SETTINGS}:\n    {e}",
                        )

        self.root.mainloop()

    @staticmethod
    def site_dict(name, address_line1, address_line2, phone):
        """A hospital/site where therapies are administered."""
        return {
            "name": name,
            "address_line1": address_line1,
            "address_line2": address_line2,
            "phone": phone,
        }

    @staticmethod
    def radionuclide_dict(
        name, half_life, activity_outpatient=None, specific_dose_rate_1m=None
    ):
        """A therapeutic radionuclide.

        Args:
            name (str): Name of radionuclide.
            half_life (float): Physical half-life of radionuclide (h).
            activity_outpatient (float, optional): Recommended maximum administered activity (MBq) for outpatient therapy
                with this radionuclide per RPS 4. Defaults to None.
            specific_dose_rate_1m (float, optional): Curve fit upper bound on initial dose rate at 1 m from patient per unit
                administered activity (uSv/h/MBq). Typically use specific gamma dose rate constant for unattenuated point source.
                See e.g. Unger and Trubey ORNL 1981. Defaults to None.

        Returns:
            dict: Keys same as args.
        """
        return {
            "name": name,
            "half_life": half_life,
            "specific_dose_rate_1m": specific_dose_rate_1m,
            "activity_outpatient": activity_outpatient,
        }

    def therapy_dict(
        self,
        name,
        radionuclide,
        inpatient,
        generic_clearance=False,
        generic_model=None,
        generic_parameters=None,
    ):
        if radionuclide not in self.radionuclide_options_df.index.to_list():
            raise ValueError(f'Radionuclide "{radionuclide}" not recognised')
        return {
            "name": name,
            "radionuclide": radionuclide,
            "inpatient": inpatient,
            "generic_clearance": generic_clearance,
            "generic_model": generic_model,
            "generic_parameters": generic_parameters,
        }

    def generic_updates(self):
        if "type_therapy" in self.odict["data"]["patient_details"]:
            if (
                self.odict["data"]["patient_details"]["type_therapy"]
                in self.therapy_options_df.index.to_list()
            ):
                if self.therapy_options_df.loc[
                    self.odict["data"]["patient_details"]["type_therapy"],
                    "generic_clearance",
                ]:
                    if (
                        "administered_activity"
                        in self.odict["data"]["administration_details"]
                    ):
                        a0 = float(
                            self.odict["data"]["administration_details"][
                                "administered_activity"
                            ]
                        )
                        generic_parameters = self.therapy_options_df.loc[
                            self.odict["data"]["patient_details"]["type_therapy"],
                            "generic_parameters",
                        ].copy()
                        generic_parameters[0] *= a0
                        generic_model = self.therapy_options_df.loc[
                            self.odict["data"]["patient_details"]["type_therapy"],
                            "generic_model",
                        ]
                        if generic_model == "exponential":
                            self.odict["data"]["clearance_data"]["curve_fit"] = (
                                collections.OrderedDict(
                                    [
                                        ("generic", "1"),
                                        ("model", "exponential"),
                                        (
                                            "fit_parameters",
                                            collections.OrderedDict(
                                                [
                                                    ("a", generic_parameters[0]),
                                                    (
                                                        "b",
                                                        np.log(2)
                                                        / generic_parameters[1],
                                                    ),
                                                ]
                                            ),
                                        ),
                                        (
                                            "meaningful_parameters",
                                            collections.OrderedDict(
                                                [
                                                    (
                                                        "dose_rate_1m_init",
                                                        generic_parameters[0],
                                                    ),
                                                    (
                                                        "effective_half_life",
                                                        generic_parameters[1],
                                                    ),
                                                ]
                                            ),
                                        ),
                                    ]
                                )
                            )
                        elif generic_model == "biexponential":
                            self.odict["data"]["clearance_data"]["curve_fit"] = (
                                collections.OrderedDict(
                                    [
                                        ("generic", "1"),
                                        ("model", "biexponential"),
                                        (
                                            "fit_parameters",
                                            collections.OrderedDict(
                                                [
                                                    ("a", generic_parameters[0]),
                                                    ("b", generic_parameters[1]),
                                                    (
                                                        "c",
                                                        np.log(2)
                                                        / generic_parameters[2],
                                                    ),
                                                    (
                                                        "d",
                                                        np.log(2)
                                                        / generic_parameters[3],
                                                    ),
                                                ]
                                            ),
                                        ),
                                        (
                                            "meaningful_parameters",
                                            collections.OrderedDict(
                                                [
                                                    (
                                                        "dose_rate_1m_init",
                                                        generic_parameters[0],
                                                    ),
                                                    (
                                                        "fraction_1",
                                                        generic_parameters[1],
                                                    ),
                                                    (
                                                        "half_life_1",
                                                        generic_parameters[2],
                                                    ),
                                                    (
                                                        "half_life_2",
                                                        generic_parameters[3],
                                                    ),
                                                ]
                                            ),
                                        ),
                                    ]
                                )
                            )

                    self.button_clearance.grid_forget()
                else:
                    self.button_clearance.grid(row=4, pady=(5, 0))
                if not self.therapy_options_df.loc[
                    self.odict["data"]["patient_details"]["type_therapy"], "inpatient"
                ]:
                    self.button_discharge.grid_forget()
                else:
                    self.button_discharge.grid(row=5, pady=(5, 0))

    # Text in top left corner of Main Menu (self.info_label)
    def get_info_str(self):
        info_str = ""
        if "name" in self.odict["data"]["patient_details"]:
            last_name, first_name = self.odict["data"]["patient_details"]["name"].split(
                "^"
            )
            info_str += f"{last_name.upper()} {first_name}"
        if "type_therapy" in self.odict["data"]["patient_details"]:
            info_str += "\n{}".format(
                self.odict["data"]["patient_details"]["type_therapy"]
                .split(",")[0]
                .split("(")[0]
            )
        if "administration_datetime" in self.odict["data"]["administration_details"]:
            admin_datetime = str2datetime(
                self.odict["data"]["administration_details"]["administration_datetime"]
            )
            info_str += "\n"
            info_str += admin_datetime.strftime("%d %b %Y").lstrip("0")
        return info_str

    def set_default_settings(self):
        self.email = "info@example.org"
        self.url = "https://example.org"
        self.report_logo = str(ASSET_DIR / "organisation_logo.png")
        site_options = [
            self.site_dict(
                "Your Hospital Here",
                "123 Main Street",
                "Anytown SA 5002",
                "(08) 1234 5678",
            ),
            self.site_dict(
                "Another Hospital",
                "124 Main Street",
                "Anytown SA 5002",
                "(08) 1234 5679",
            ),
        ]
        self.site_options_df = pd.DataFrame(site_options)
        self.site_options_df.set_index("name", inplace=True)

        patient_handout_comments = (
            "For the next two days, continue to take precautions with your urine and saliva.  "
            "Sit down while urinating, shut the toilet lid before flushing, and flush twice.  "
            "Avoid contacting other people with your saliva "
            "and wash your toothbrush thoroughly after each use."
        )

        self.init_vals = {
            "therapy_type": self.therapy_options_df.index.to_list()[0],
            "site": self.site_options_df.index.to_list()[0],
            "num_treatments_in_year": 1,
            "measurement_distance": 2,
            "detector_calibration_factor": 1,
            "curve_fit_model": "biexponential",
            "data_directory": os.getcwd(),
            "patient_handout_directory": os.getcwd(),
            "summary_report_directory": os.getcwd(),
            "patient_handout_comments": patient_handout_comments,
        }

        self.other_settings = {"export_pdf": 0, "allow_admin_before_calib": 0}

    @staticmethod
    def read_organisation_info(settings_odict):
        if "organisation" not in settings_odict["root"]:
            raise KeyError('Missing "organisation" key')

        organisation_keys = ["email", "url", "logo"]
        if any(
            x not in settings_odict["root"]["organisation"] for x in organisation_keys
        ):
            raise KeyError(
                f"Missing one or more of the following keys: {organisation_keys}"
            )

        email = settings_odict["root"]["organisation"]["email"]
        url = settings_odict["root"]["organisation"]["url"]
        report_logo = settings_odict["root"]["organisation"]["logo"]

        return email, url, report_logo

    @staticmethod
    def read_sites(settings_odict):
        if "organisation" not in settings_odict["root"]:
            raise KeyError('Missing "organisation" key')
        if "site" not in settings_odict["root"]["organisation"]:
            raise KeyError('Missing "site" key')
        site_keys = ["name", "address_line1", "address_line2", "phone"]
        for site in settings_odict["root"]["organisation"]["site"]:
            if any(x not in site for x in site_keys):
                raise KeyError(
                    "There is a site missing one or more of the keys: {}".format(
                        ", ".join(site_keys)
                    )
                )

        df_site_options = pd.DataFrame(settings_odict["root"]["organisation"]["site"])
        df_site_options = df_site_options[df_site_options["name"] != ""]
        if df_site_options.empty:
            raise ValueError("Provide a name for at least 1 site")
        df_site_options.set_index("name", inplace=True)

        return df_site_options

    def read_init_vals(self, settings_odict):
        if "initial_values" not in settings_odict["root"]:
            raise KeyError('Missing "initial_values" key')
        init_vals = settings_odict["root"]["initial_values"]

        if any(x not in init_vals for x in self.init_vals):
            raise KeyError(
                f"Missing one of the following keys: {list(self.init_vals.keys())}"
            )

        therapy = init_vals["therapy_type"]
        site = init_vals["site"]
        curve_fit_model = init_vals["curve_fit_model"]

        therapy_options = self.therapy_options_df.index.to_list()
        if therapy not in therapy_options:
            raise ValueError(f'"{therapy}" is not a valid therapy option')

        if site not in self.site_options_df.index.to_list():
            raise ValueError(f'"{site}" is not a valid site option')

        if curve_fit_model not in ["exponential", "biexponential"]:
            raise ValueError(f'"{curve_fit_model}" is not a valid curve fit model')

        return init_vals

    def read_other_settings(self, settings_odict):
        other_settings = settings_odict["root"]["other_settings"]
        if any(x not in other_settings for x in self.other_settings):
            raise KeyError(
                f"Missing one of the following keys: {list(self.other_settings.keys())}"
            )
        return other_settings

    def write_settings(self):
        settings_odict = collections.OrderedDict()
        settings_odict["root"] = collections.OrderedDict()

        settings_odict["root"]["organisation"] = collections.OrderedDict()
        settings_odict["root"]["organisation"]["email"] = self.email
        settings_odict["root"]["organisation"]["url"] = self.url
        settings_odict["root"]["organisation"]["logo"] = self.report_logo

        self.site_options_df.reset_index(inplace=True)
        settings_odict["root"]["organisation"]["site"] = self.site_options_df.to_dict(
            "records", into=collections.OrderedDict
        )
        self.site_options_df.set_index("name", inplace=True)

        settings_odict["root"]["initial_values"] = collections.OrderedDict(
            self.init_vals
        )

        settings_odict["root"]["other_settings"] = collections.OrderedDict(
            self.other_settings
        )

        self.FILENAME_SETTINGS.parent.mkdir(parents=True, exist_ok=True)
        with open(self.FILENAME_SETTINGS, "w") as fd:
            fd.write(xmltodict.unparse(settings_odict, pretty=True))

    # Settings > Organisation
    def settings_organisation(self):
        # button next to logo path that can fill the path
        def fill_stringvar(stringvar):
            f_old = stringvar.get()
            if os.path.isfile(f_old):
                startdir = os.path.dirname(f_old)
            else:
                startdir = ASSET_DIR

            filep = filedialog.askopenfilename(
                initialdir=startdir,
                parent=window,
                title="Select organisation logo",
                filetypes=[("Image Files", (".png", ".jpg", ".jpeg"))],
            )
            if filep != "" and filep != ():
                stringvar.set(filep)

        def submit_settings_organisation(self):
            site_options = [e.get() for e in site]
            addressA_list = [e.get() for e in addressA]
            addressB_list = [e.get() for e in addressB]
            phone_list = [e.get() for e in phone]

            site_options_df = pd.DataFrame(
                data={
                    "name": site_options,
                    "address_line1": addressA_list,
                    "address_line2": addressB_list,
                    "phone": phone_list,
                }
            )
            site_options_df = site_options_df[site_options_df["name"] != ""]

            if site_options_df.empty:
                messagebox.showerror(
                    "Error",
                    "You must enter a name for at least one site",
                    parent=window,
                )
                return

            self.site_options_df = site_options_df
            self.site_options_df.set_index("name", inplace=True)

            self.url = url_entry.get()
            self.email = email_entry.get()
            self.report_logo = logo_entry.get()

            window.withdraw()

            self.init_vals["site"] = self.site_options_df.index.to_list()[0]

            self.write_settings()

        window = tk.Toplevel()
        window.geometry("580x730")
        window.title("Settings - Organisation")
        if WINDOWS_OS:
            window.resizable(width=False, height=False)
            window.iconbitmap(self.SOFTWARE_ICON)

        tk.Label(window, text="Site options", font=("Arial", 10, "bold")).grid(
            row=0, column=0, pady=(10, 0), columnspan=5
        )

        tk.Label(window, text="Site").grid(row=1, column=1, pady=(10, 0))
        tk.Label(window, text="Address").grid(row=1, column=2, pady=(10, 0))
        tk.Label(window, text="Phone").grid(row=1, column=3, pady=(10, 0))

        site = []
        addressA = []
        addressB = []
        phone = []
        site_options = self.site_options_df.index.to_list()
        N_SITES_MAX = 8
        j = 0
        for x in range(0, N_SITES_MAX * 2, 2):
            tk.Label(window, text="{}.".format(str(int(x / 2 + 1)))).grid(
                row=x + 2, column=0, sticky="e", padx=(10, 0)
            )

            site.append(tk.Entry(window, width=32))
            site[int(x / 2)].grid(row=x + 2, column=1, padx=(5, 0))

            addressA.append(tk.Entry(window, width=40))
            addressA[int(x / 2)].grid(row=x + 2, column=2, padx=(5, 0))

            addressB.append(tk.Entry(window, width=40))
            addressB[int(x / 2)].grid(row=x + 3, column=2, pady=(0, 15), padx=(5, 0))

            phone.append(tk.Entry(window, width=14))
            phone[int(x / 2)].grid(row=x + 2, column=3, padx=(5, 0))

            if j < len(site_options):
                site_str = site_options[j]
                site[int(x / 2)].insert(0, site_str)
                addressA[int(x / 2)].insert(
                    0, self.site_options_df.loc[site_str, "address_line1"]
                )
                addressB[int(x / 2)].insert(
                    0, self.site_options_df.loc[site_str, "address_line2"]
                )
                phone[int(x / 2)].insert(0, self.site_options_df.loc[site_str, "phone"])
            j += 1

        tk.Label(
            window,
            text="To appear in summary reports and patient handouts",
            font=("Arial", 10, "bold"),
        ).grid(row=4 + x, column=0, pady=(10, 10), columnspan=4)

        tk.Label(window, text="URL").grid(row=5 + x, column=1, sticky="e", pady=(10, 0))
        url_entry = tk.Entry(window, width=40)
        url_entry.insert(0, self.url)
        url_entry.grid(row=5 + x, column=2, sticky="we", pady=(10, 0), padx=(5, 0))

        tk.Label(window, text="Email").grid(
            row=6 + x, column=1, sticky="e", pady=(10, 0)
        )
        email_entry = tk.Entry(window, width=40)
        email_entry.insert(0, self.email)
        email_entry.grid(row=6 + x, column=2, sticky="we", pady=(10, 0), padx=(5, 0))

        tk.Label(window, text="Logo image file").grid(
            row=7 + x, column=1, sticky="e", pady=(10, 0), padx=(5, 0)
        )
        logo_str = tk.StringVar()
        logo_str.set(self.report_logo)
        logo_entry = tk.Entry(window, textvariable=logo_str)
        logo_entry.grid(row=7 + x, column=2, sticky="we", pady=(10, 0), padx=(5, 0))
        logo_xs = ttk.Scrollbar(window, orient="horizontal", command=logo_entry.xview)
        logo_entry["xscrollcommand"] = logo_xs.set
        logo_xs.grid(row=8 + x, column=2, sticky="we", pady=(0, 30))
        tk.Button(
            window,
            text="\N{HORIZONTAL ELLIPSIS}",
            command=lambda: fill_stringvar(logo_str),
        ).grid(row=7 + x, column=3, sticky="w", pady=(5, 0))

        tk.Button(
            window, text="Submit", command=lambda: submit_settings_organisation(self)
        ).grid(row=9 + x, column=0, columnspan=4)

    # Settings > Initial Values
    def settings_init_vals(self):
        def fill_stringvar(stringvar):
            mydir = filedialog.askdirectory(initialdir=stringvar.get(), parent=window)
            if mydir != "" and mydir != ():
                stringvar.set(mydir)

        def submit_settings_init_vals(self):
            window.withdraw()

            init_vals = {}
            init_vals["therapy_type"] = therapy_combo.get()
            init_vals["site"] = site_combo.get()
            init_vals["num_treatments_in_year"] = num_treatments_in_year_entry.get()
            init_vals["measurement_distance"] = measurement_distance_entry.get()
            init_vals["detector_calibration_factor"] = (
                detector_calibration_factor_entry.get()
            )
            init_vals["curve_fit_model"] = cfit_combo.get()
            init_vals["data_directory"] = data_directory_entry.get()
            init_vals["patient_handout_directory"] = (
                patient_handout_directory_entry.get()
            )
            init_vals["summary_report_directory"] = summary_report_directory_entry.get()
            init_vals["patient_handout_comments"] = patient_handout_comments.get(
                "1.0", "end"
            )

            self.init_vals = init_vals
            self.write_settings()

        window = tk.Toplevel()

        therapy_options = self.therapy_options_df.index.to_list()
        width_therapy = max(
            [len(therapy_options[i]) for i in range(len(therapy_options))]
        )
        window_width = 6 * width_therapy + 218
        if window_width < 470:
            window_width = 470
        window.geometry(f"{int(window_width)}x600")
        window.title("Settings - Initial Values")
        if WINDOWS_OS:
            window.resizable(width=False, height=False)
            window.iconbitmap(self.SOFTWARE_ICON)

        tk.Label(window, text="Set initial values", font=("Arial", 10, "bold")).grid(
            row=0, column=0, columnspan=3, pady=(10, 0)
        )

        tk.Label(window, text="Type of therapy").grid(
            row=1, column=0, pady=(20, 0), sticky="e"
        )
        therapy_combo = ttk.Combobox(
            window,
            state="readonly",
            values=therapy_options,
            width=width_therapy,
        )
        therapy_combo.set(self.init_vals["therapy_type"])
        therapy_combo.grid(row=1, column=1, padx=(10, 0), pady=(20, 0), sticky="w")

        tk.Label(window, text="Site").grid(row=2, column=0, sticky="e", pady=(5, 0))
        site_options = self.site_options_df.index.to_list()
        site_combo = ttk.Combobox(
            window,
            state="readonly",
            values=site_options,
            width=max([len(site_options[i]) for i in range(len(site_options))]),
        )
        site_combo.set(self.init_vals["site"])
        site_combo.grid(row=2, column=1, padx=(10, 0), sticky="w", pady=(5, 0))

        tk.Label(
            window, text="Number of treatments\nexpected in a year", justify=tk.RIGHT
        ).grid(row=3, column=0, sticky="e", pady=(5, 0))
        num_treatments_in_year_entry = tk.Entry(window, width=3)
        num_treatments_in_year_entry.insert(0, self.init_vals["num_treatments_in_year"])
        num_treatments_in_year_entry.grid(
            row=3, column=1, padx=(10, 0), sticky="w", pady=(5, 0)
        )

        tk.Label(window, text="Measurement distance (m)").grid(
            row=4, column=0, sticky="e", pady=(5, 0)
        )
        measurement_distance_entry = tk.Entry(window, width=5)
        measurement_distance_entry.insert(0, self.init_vals["measurement_distance"])
        measurement_distance_entry.grid(
            row=4, column=1, padx=(10, 0), sticky="w", pady=(5, 0)
        )

        tk.Label(
            window,
            text="Detector calibration factor\n(true/displayed)",
            justify=tk.RIGHT,
        ).grid(row=5, column=0, sticky="e", pady=(5, 0))
        detector_calibration_factor_entry = tk.Entry(window, width=5)
        detector_calibration_factor_entry.insert(
            0, self.init_vals["detector_calibration_factor"]
        )
        detector_calibration_factor_entry.grid(
            row=5, column=1, padx=(10, 0), sticky="w", pady=(5, 0)
        )

        cfit_options = ["exponential", "biexponential"]
        tk.Label(window, text="Clearance curve fit model").grid(
            row=6, column=0, sticky="e", pady=(5, 0)
        )
        cfit_combo = ttk.Combobox(
            window,
            state="readonly",
            values=cfit_options,
            width=max([len(cfit_options[i]) for i in range(len(cfit_options))]),
        )
        cfit_combo.set(self.init_vals["curve_fit_model"])
        cfit_combo.grid(row=6, column=1, padx=(10, 0), sticky="w", pady=(5, 0))

        tk.Label(window, text="XML output directory").grid(
            row=7, column=0, sticky="e", pady=(5, 0)
        )
        data_directory_str = tk.StringVar()
        data_directory_str.set(self.init_vals["data_directory"])
        data_directory_entry = tk.Entry(window, textvariable=data_directory_str)
        data_directory_entry.grid(
            row=7, column=1, padx=(10, 0), sticky="we", pady=(5, 0)
        )
        data_directory_xs = ttk.Scrollbar(
            window, orient="horizontal", command=data_directory_entry.xview
        )
        data_directory_entry["xscrollcommand"] = data_directory_xs.set
        data_directory_xs.grid(row=8, column=1, sticky="we")
        tk.Button(
            window,
            text="\N{HORIZONTAL ELLIPSIS}",
            command=lambda: fill_stringvar(data_directory_str),
        ).grid(row=7, column=2)

        tk.Label(window, text="Patient handout directory").grid(
            row=9, column=0, sticky="e", pady=(5, 0)
        )
        patient_handout_directory_str = tk.StringVar()
        patient_handout_directory_str.set(self.init_vals["patient_handout_directory"])
        patient_handout_directory_entry = tk.Entry(
            window, textvariable=patient_handout_directory_str
        )
        patient_handout_directory_entry.grid(
            row=9, column=1, padx=(10, 0), sticky="we", pady=(5, 0)
        )
        patient_handout_directory_xs = ttk.Scrollbar(
            window, orient="horizontal", command=patient_handout_directory_entry.xview
        )
        patient_handout_directory_entry["xscrollcommand"] = (
            patient_handout_directory_xs.set
        )
        patient_handout_directory_xs.grid(row=10, column=1, sticky="we")
        tk.Button(
            window,
            text="\N{HORIZONTAL ELLIPSIS}",
            command=lambda: fill_stringvar(patient_handout_directory_str),
        ).grid(row=9, column=2)

        tk.Label(window, text="Summary report directory").grid(
            row=11, column=0, sticky="e", pady=(5, 0)
        )
        summary_report_directory_str = tk.StringVar()
        summary_report_directory_str.set(self.init_vals["summary_report_directory"])
        summary_report_directory_entry = tk.Entry(
            window, textvariable=summary_report_directory_str
        )
        summary_report_directory_entry.grid(
            row=11, column=1, padx=(10, 0), sticky="we", pady=(5, 0)
        )
        summary_report_directory_xs = ttk.Scrollbar(
            window, orient="horizontal", command=summary_report_directory_entry.xview
        )
        summary_report_directory_entry["xscrollcommand"] = (
            summary_report_directory_xs.set
        )
        summary_report_directory_xs.grid(row=12, column=1, sticky="we")
        tk.Button(
            window,
            text="\N{HORIZONTAL ELLIPSIS}",
            command=lambda: fill_stringvar(summary_report_directory_str),
        ).grid(row=11, column=2)

        tk.Label(
            window, text="Additional comments\non patient handout", justify=tk.RIGHT
        ).grid(row=13, column=0, sticky="ne", pady=(5, 0))
        patient_handout_comments = tk.Text(window, width=35, height=10, wrap=tk.WORD)
        patient_handout_comments.insert(
            "1.0", self.init_vals["patient_handout_comments"]
        )
        patient_handout_comments.grid(
            row=13, column=1, padx=(10, 0), sticky="we", pady=(5, 0), columnspan=2
        )

        tk.Button(
            window, text="Submit", command=lambda: submit_settings_init_vals(self)
        ).grid(row=14, column=0, columnspan=3, pady=(20, 0))

    # Settings > Other
    def settings_other(self):
        def submit_settings_other(self):
            window.withdraw()

            other_settings = {}
            other_settings["export_pdf"] = chk_export_pdf.get()
            other_settings["allow_admin_before_calib"] = (
                chk_allow_admin_before_calib.get()
            )
            self.other_settings = other_settings
            self.write_settings()

        window = tk.Toplevel()
        window.geometry("355x155")
        window.title("Settings - Other")
        if WINDOWS_OS:
            window.resizable(width=False, height=False)
            window.iconbitmap(self.SOFTWARE_ICON)

        tk.Label(window, text="Other settings", font=("Arial", 10, "bold")).grid(
            row=0, column=0, columnspan=2, pady=(10, 10)
        )

        tk.Label(
            window, text="Export reports to PDF if MS Word is installed (Windows only)"
        ).grid(row=1, column=0, sticky="W", padx=(5, 0))
        chk_export_pdf = tk.IntVar()
        tk.Checkbutton(window, variable=chk_export_pdf).grid(row=1, column=1)
        chk_export_pdf.set(self.other_settings["export_pdf"])

        tk.Label(window, text="Allow administration time before calibration time").grid(
            row=2, column=0, sticky="W", padx=(5, 0)
        )
        chk_allow_admin_before_calib = tk.IntVar()
        tk.Checkbutton(window, variable=chk_allow_admin_before_calib).grid(
            row=2, column=1
        )
        chk_allow_admin_before_calib.set(
            self.other_settings["allow_admin_before_calib"]
        )

        tk.Button(
            window, text="Submit", command=lambda: submit_settings_other(self)
        ).grid(row=3, column=0, columnspan=2, pady=(20, 0))

    # New
    def new_patient(self):
        if int(self.odict["data"]["patient_finished"]):
            self.viewing_completed_patient_label.place_forget()
        else:
            if self.unsaved_data:
                if not messagebox.askokcancel(
                    "Warning - New",
                    "You have unsaved data. Are you sure you want to start over?",
                    default="cancel",
                    icon="warning",
                ):
                    return
        self.odict = self.get_new_odict()
        self.info_label.place_forget()
        self.button_clearance.grid(row=4, pady=(5, 0))
        self.button_discharge.grid(row=5, pady=(5, 0))
        self.update_buttons()
        self.root.title(f"Main Menu - {__program_name__}")
        self.filepath = None
        self.unsaved_data = False

    # Initialise self.odict or see if it is empty
    def get_new_odict(self):
        odict = collections.OrderedDict()
        odict["data"] = collections.OrderedDict()
        odict["data"][f"{__program_name__.lower()}_version"] = __version__
        odict["data"]["glowgreen_version"] = GLOWGREEN_VERSION
        odict["data"]["patient_details"] = collections.OrderedDict()
        odict["data"]["administration_details"] = collections.OrderedDict()
        odict["data"]["clearance_data"] = collections.OrderedDict()
        odict["data"]["patient_discharge"] = collections.OrderedDict()
        odict["data"]["patient_discharge"][
            "recommended_datetime"
        ] = collections.OrderedDict()
        odict["data"]["restrictions"] = collections.OrderedDict()
        odict["data"]["additional_comments_to_patient"] = "0"
        odict["data"]["reports_generated"] = "0"
        odict["data"]["patient_finished"] = "0"
        odict["data"]["patient_finished_by"] = "0"
        return odict

    # Load
    def load_patient(self):
        if self.unsaved_data:
            messagebox.showwarning(
                "Warning - Open",
                "You have unsaved data, which will be lost if you open a file.",
            )

        filepath = filedialog.askopenfilename(
            initialdir=(
                self.init_vals["data_directory"]
                if self.previous_data_directory is None
                else self.previous_data_directory
            ),
            title="Select patient XML file",
            filetypes=[("XML files (*.xml)", "*.xml")],
        )

        if filepath != "" and filepath != ():
            self.previous_data_directory = os.path.dirname(filepath)

            try:
                with open(filepath) as fd:
                    odict_from_file = xmltodict.parse(
                        fd.read(),
                        dict_constructor=collections.OrderedDict,
                        postprocessor=Gui.my_postprocessor_patient,
                    )
            except Exception as e:
                messagebox.showerror(
                    title="File Open Error",
                    message=f"Unable to open {os.path.basename(filepath)} due to:\n{e}",
                )
                return

            if "data" not in odict_from_file:
                messagebox.showerror(
                    title="File Open Error",
                    message=f'Unable to open {os.path.basename(filepath)}. Missing key "data"',
                )
                return

            init_odict = self.get_new_odict()
            for key in init_odict["data"]:
                if (
                    key
                    not in [
                        "patient_finished_by",
                        "additional_comments_to_patient",
                        f"{__program_name__.lower()}_version",
                        "glowgreen_version",
                    ]
                ) and (key not in odict_from_file["data"]):
                    messagebox.showerror(
                        title="File Open Error",
                        message=f'Unable to open {os.path.basename(filepath)}. Missing key "{key}".',
                    )
                    return
            if (
                "recommended_datetime"
                not in odict_from_file["data"]["patient_discharge"]
            ):
                messagebox.showerror(
                    title="File Open Error",
                    message=f'Unable to open {os.path.basename(filepath)}. Missing key "recommended_datetime".',
                )
                return

            self.odict = odict_from_file
            self.unsaved_data = False

            self.filepath = filepath

            # Actions to maintain backward compatibility
            if "number_of_treatments" in self.odict["data"]["patient_details"]:
                self.odict["data"]["patient_details"]["num_treatments_in_year"] = (
                    self.odict["data"]["patient_details"].pop("number_of_treatments")
                )

            update_therapy_names = {
                "I-131 remnant thyroid ablation, thyrogen": "I-131 thyroid cancer, rhTSH",
                "I-131 remnant thyroid ablation, withdrawal": "I-131 thyroid cancer, thyroxine withdrawal",
            }
            if "type_therapy" in self.odict["data"]["patient_details"]:
                type_therapy = self.odict["data"]["patient_details"]["type_therapy"]
                if type_therapy in update_therapy_names:
                    self.odict["data"]["patient_details"]["type_therapy"] = (
                        update_therapy_names[type_therapy]
                    )

            if "restriction" in self.odict["data"]["restrictions"]:
                if (
                    "dose_constraint_corrected_for_number_of_treatments"
                    in self.odict["data"]["restrictions"]["restriction"][0]
                ):  # min 2 restrictions
                    for child1 in self.odict["data"]["restrictions"]["restriction"]:
                        child1["dose_constraint_corrected"] = child1.pop(
                            "dose_constraint_corrected_for_number_of_treatments"
                        )

                if (
                    "type" in self.odict["data"]["restrictions"]["restriction"][0]
                ):  # min 2 restrictions
                    for child1 in self.odict["data"]["restrictions"]["restriction"]:
                        child1["name"] = child1.pop("type")
                        child1.move_to_end("name", last=False)

                if (
                    "per_episode"
                    not in self.odict["data"]["restrictions"]["restriction"][0]
                ):  # min 2 restrictions
                    for child1 in self.odict["data"]["restrictions"]["restriction"]:
                        if child1["name"] in [
                            "Close contact with informed persons caring for patient",
                            "Sleeping with informed person supporting patient",
                        ]:
                            child1["per_episode"] = "1"
                        else:
                            child1["per_episode"] = "0"

                if "tau" in self.odict["data"]["restrictions"]["restriction"][0]:
                    for child1 in self.odict["data"]["restrictions"]["restriction"]:
                        child1["restriction_period"] = child1.pop("tau")

            if "type_therapy" in self.odict["data"]["patient_details"]:
                if (
                    self.odict["data"]["patient_details"]["type_therapy"]
                    not in self.therapy_options_df.index.to_list()
                ):
                    messagebox.showwarning(
                        "Warning - Therapy Type",
                        "Therapy '{}' not recognised. Select a new one to proceed.".format(
                            self.odict["data"]["patient_details"]["type_therapy"]
                        ),
                    )
                    self.odict["data"]["patient_details"].pop("type_therapy")

                    if (
                        "administered_activity"
                        in self.odict["data"]["administration_details"]
                    ):
                        self.odict["data"]["administration_details"].pop(
                            "administered_activity"
                        )

                    # type_therapy is used in retrieve_fit button to get clearance curve_fit
                    if "curve_fit" in self.odict["data"]["clearance_data"]:
                        self.odict["data"]["clearance_data"].pop("curve_fit")

                    # Hence quantities that use the clearance curve_fit are also invaldiated...
                    # curve_fit is used in discharge button to calculate discharge_activity
                    if "discharge_activity" in self.odict["data"]["patient_discharge"]:
                        self.odict["data"]["patient_discharge"].pop(
                            "discharge_activity"
                        )

                    # curve_fit is used in discharge button to calculate calculated_discharge_dose_rate_xm/1m
                    if (
                        "calculated_discharge_dose_rate_xm"
                        in self.odict["data"]["patient_discharge"]
                    ):
                        self.odict["data"]["patient_discharge"].pop(
                            "calculated_discharge_dose_rate_xm"
                        )
                    if (
                        "calculated_discharge_dose_rate_1m"
                        in self.odict["data"]["patient_discharge"]
                    ):
                        self.odict["data"]["patient_discharge"].pop(
                            "calculated_discharge_dose_rate_1m"
                        )

                    # curve_fit is used in discharge window to calculate recommended_datetime for discharge
                    self.odict["data"]["patient_discharge"][
                        "recommended_datetime"
                    ] = collections.OrderedDict()

                    # curve_fit is used to calculate restrictions
                    if "restriction" in self.odict["data"]["restrictions"]:
                        if (
                            "restriction_period"
                            in self.odict["data"]["restrictions"]["restriction"][0]
                        ):  # min 2 restrictions
                            for child1 in self.odict["data"]["restrictions"][
                                "restriction"
                            ]:
                                child1.pop("dose_constraint_corrected")
                                child1.pop("restriction_period")
                                child1.pop("dose")
                                child1.pop("datetime_end")

                    # type_therapy is used in discharge window to calculate recommended_datetime for discharge (radionuclide -> RPS 4 activity limit)

                    self.odict["data"]["reports_generated"] = "0"

            if (
                "calculated_discharge_dose_rate_2m"
                in self.odict["data"]["patient_discharge"]
            ):
                self.odict["data"]["patient_discharge"][
                    "calculated_discharge_dose_rate_xm"
                ] = self.odict["data"]["patient_discharge"].pop(
                    "calculated_discharge_dose_rate_2m"
                )

            if "curve_fit" in self.odict["data"]["clearance_data"]:
                if "generic" not in self.odict["data"]["clearance_data"]["curve_fit"]:
                    if (
                        self.odict["data"]["patient_details"]["type_therapy"]
                        in self.therapy_options_df.index.to_list()
                    ):
                        if self.therapy_options_df.loc[
                            self.odict["data"]["patient_details"]["type_therapy"],
                            "generic_clearance",
                        ]:
                            self.odict["data"]["clearance_data"]["curve_fit"][
                                "generic"
                            ] = "1"
                        else:
                            self.odict["data"]["clearance_data"]["curve_fit"][
                                "generic"
                            ] = "0"
                        self.odict["data"]["clearance_data"]["curve_fit"].move_to_end(
                            "generic", last=False
                        )

                if "parameters" in self.odict["data"]["clearance_data"]["curve_fit"]:
                    model = self.odict["data"]["clearance_data"]["curve_fit"]["model"]
                    if model == "biexponential":
                        x = [
                            float(a)
                            for a in self.odict["data"]["clearance_data"]["curve_fit"][
                                "parameters"
                            ].values()
                        ]
                        dose_rate_xm_init = x[0] + x[2]
                        fraction_1 = x[0] / dose_rate_xm_init
                        lambda_1 = x[1]
                        lambda_2 = x[3]
                        fit_parameters = collections.OrderedDict(
                            [
                                ("a", dose_rate_xm_init),
                                ("b", fraction_1),
                                ("c", lambda_1),
                                ("d", lambda_2),
                            ]
                        )
                        self.odict["data"]["clearance_data"]["curve_fit"].pop(
                            "parameters"
                        )
                        self.odict["data"]["clearance_data"]["curve_fit"][
                            "fit_parameters"
                        ] = fit_parameters
                    elif model == "exponential":
                        self.odict["data"]["clearance_data"]["curve_fit"][
                            "fit_parameters"
                        ] = self.odict["data"]["clearance_data"]["curve_fit"].pop(
                            "parameters"
                        )

                if "output_data" in self.odict["data"]["clearance_data"]["curve_fit"]:
                    model = self.odict["data"]["clearance_data"]["curve_fit"]["model"]
                    if model == "biexponential":
                        dose_rate_xm_init, fraction_1, half_life_1, _, half_life_2 = [
                            float(a)
                            for a in self.odict["data"]["clearance_data"]["curve_fit"][
                                "output_data"
                            ].values()
                        ]
                        meaningful_parameters = collections.OrderedDict(
                            [
                                ("dose_rate_xm_init", dose_rate_xm_init),
                                ("fraction_1", fraction_1),
                                ("half_life_1", half_life_1),
                                ("half_life_2", half_life_2),
                            ]
                        )
                        self.odict["data"]["clearance_data"]["curve_fit"].pop(
                            "output_data"
                        )
                        self.odict["data"]["clearance_data"]["curve_fit"][
                            "meaningful_parameters"
                        ] = meaningful_parameters
                    elif model == "exponential":
                        dose_rate_xm_init, effective_half_life = [
                            float(a)
                            for a in self.odict["data"]["clearance_data"]["curve_fit"][
                                "output_data"
                            ].values()
                        ]
                        meaningful_parameters = collections.OrderedDict(
                            [
                                ("dose_rate_xm_init", dose_rate_xm_init),
                                ("effective_half_life", effective_half_life),
                            ]
                        )
                        self.odict["data"]["clearance_data"]["curve_fit"].pop(
                            "output_data"
                        )
                        self.odict["data"]["clearance_data"]["curve_fit"][
                            "meaningful_parameters"
                        ] = meaningful_parameters

            if "unsaved_data" in self.odict["data"]:
                self.odict["data"].pop("unsaved_data")

            if "additional_comments_to_patient" not in self.odict["data"]:
                self.odict["data"]["additional_comments_to_patient"] = "0"

            if "filepath" in self.odict["data"]:
                self.odict["data"].pop("filepath")

            if "morningstar_version" in self.odict["data"]:
                self.odict["data"].pop("morningstar_version")
                self.odict["data"][f"{__program_name__.lower()}_version"] = "0"
                self.odict["data"]["glowgreen_version"] = "0"
                self.odict["data"].move_to_end("glowgreen_version", last=False)
                self.odict["data"].move_to_end(
                    f"{__program_name__.lower()}_version", last=False
                )

            self.root.title(f"{os.path.basename(filepath)} - {__program_name__}")

            info_str = self.get_info_str()
            if info_str:
                self.info_label["text"] = info_str
                self.info_label.place(relx=0.01, rely=0.01)

            self.generic_updates()
            self.update_buttons()

            if int(self.odict["data"]["patient_finished"]):
                self.viewing_completed_patient_label.place(
                    relx=0.5, rely=0.05, anchor="center"
                )

                if self.odict["data"]["patient_finished_by"] != "0":
                    self.viewing_completed_patient_label["text"] = (
                        "Viewing completed patient\n(author: {})".format(
                            self.odict["data"]["patient_finished_by"]
                        )
                    )
                else:
                    self.viewing_completed_patient_label["text"] = (
                        "Viewing completed patient."
                    )
            else:
                self.viewing_completed_patient_label.place_forget()

    # Save
    def save_patient(self):
        filepath = self.filepath
        if self.unsaved_data or filepath is None:
            if filepath is None:
                _ = self.save_patient_as()
            else:
                self.save_odict_to_xml(filepath)

    def save_odict_to_xml(self, filepath):
        if not filepath.endswith(".xml"):
            filepath += ".xml"

        self.filepath = filepath
        self.unsaved_data = False
        self.odict["data"][f"{__program_name__.lower()}_version"] = __version__

        with open(filepath, "w") as fd:
            fd.write(xmltodict.unparse(self.odict, pretty=True))

        self.root.title(f"{os.path.basename(filepath)} - {__program_name__}")

    # Save As
    def save_patient_as(self):
        filepath = self.filepath
        if filepath is None:
            filename = ""
            if "name" in self.odict["data"]["patient_details"]:
                last_name, first_name = self.odict["data"]["patient_details"][
                    "name"
                ].split("^")
                filename += last_name.upper() + "_" + first_name
            if (
                "administration_datetime"
                in self.odict["data"]["administration_details"]
            ):
                admin_datetime = str2datetime(
                    self.odict["data"]["administration_details"][
                        "administration_datetime"
                    ]
                )
                filename += "_" + admin_datetime.strftime("%b%Y")
        else:
            filename = os.path.basename(filepath)

        filepath_new = filedialog.asksaveasfilename(
            initialdir=(
                self.init_vals["data_directory"]
                if self.previous_data_directory is None
                else self.previous_data_directory
            ),
            title="Save XML file",
            initialfile=filename,
            filetypes=[("XML file (*.xml)", "*.xml")],
        )

        if filepath_new:
            self.previous_data_directory = os.path.dirname(filepath_new)
            self.save_odict_to_xml(filepath_new)
            return True
        else:
            return False

    # X
    def on_closing(self):
        if self.unsaved_data:
            filepath = self.filepath
            if filepath is not None:
                response = messagebox.askyesnocancel(
                    title="Save Changes?",
                    message=f"{os.path.basename(filepath)} has been modified, save changes?",
                    default=messagebox.CANCEL,
                )
                if response is None:
                    pass
                elif response:
                    self.save_patient()
                    self.root.quit()
                    self.root.destroy()
                else:
                    self.root.quit()
                    self.root.destroy()
            else:
                response = messagebox.askyesnocancel(
                    title="Save Progress?",
                    message="Do you want to save the progress?",
                    default=messagebox.CANCEL,
                )
                if response is None:
                    pass
                elif response:
                    t = self.save_patient_as()
                    if t:
                        self.root.quit()
                        self.root.destroy()
                else:
                    self.root.quit()
                    self.root.destroy()
        else:
            self.root.quit()
            self.root.destroy()

    # Info
    def info_box(self):
        window = tk.Toplevel()
        window.geometry("390x310")
        window.title("Info")
        if WINDOWS_OS:
            window.resizable(width=False, height=False)
            window.iconbitmap(self.SOFTWARE_ICON)

        tk.Label(
            window,
            text=__program_name__,
            font="Arial 10 bold",
        ).pack(pady=(10, 0))
        tk.Label(window, text=f"version {__version__}").pack()
        tk.Label(
            window,
            text=f"released {__release_date__}",
        ).pack()
        tk.Label(window, text=f"using glowgreen {GLOWGREEN_VERSION}").pack()

        tk.Label(window, text=f"Author: {__author__}").pack(pady=(20, 0))
        tk.Label(window, text=f"Email: {__author_email__}").pack()
        tk.Label(window, text=__homepage__).pack()

        tk.Label(
            window,
            text=f"Copyright \N{COPYRIGHT SIGN} {__copyright_year__} {__copyright_owner__}",
        ).pack(pady=(20, 0))

        # tk.Label(window, text="License: Apache-2.0", font="Arial 9 bold").pack()
        tk.Label(window, text=f"License: {__license__}").pack()

        tk.Button(window, text="OK", command=window.withdraw).pack(pady=(20, 0))

    # When reading XML file into self.odict with xmltodict.parse, empty XML tags are initialised to ordered dicts (instead of None)
    @staticmethod
    def my_postprocessor_patient(path, key, value):
        if value is None:
            return key, collections.OrderedDict()
        return key, value

    # When reading XML file into settings_odict with xmltodict.parse, empty XML tags initialised to empty strings (instead of None)
    @staticmethod
    def my_postprocessor_settings(path, key, value):
        if value is None:
            return key, ""
        return key, value

    # Patient Details
    def patient_details(self):
        # 'Submit' button
        def retrieve_patient_details(self):
            # QC
            if len(dob_year_entry.get()) != 4:
                messagebox.showerror(
                    "Error", "Please enter a 4-digit year for DOB", parent=window
                )
                return
            try:
                e_dob = date(
                    year=int(dob_year_entry.get()),
                    month=int(dob_month_entry.get()),
                    day=int(dob_day_entry.get()),
                )
            except ValueError:
                messagebox.showerror("Error", "Bad DOB", parent=window)
                return

            # Don't let the user submit an invalid number of treatments in a year
            try:
                e_num_treatments_in_year = float(num_treatments_in_year_entry.get())
            except ValueError:
                messagebox.showerror(
                    "Error", "Bad number of treatments in a year", parent=window
                )
                return
            if (
                e_num_treatments_in_year <= 0.0
                or (1 / e_num_treatments_in_year == 0)
                or np.isnan(e_num_treatments_in_year)
            ):
                messagebox.showerror(
                    "Error", "Bad number of treatments in a year", parent=window
                )
                return

            # Force user to tick these boxes
            e_pregnancy_excluded = chk_pregnancy.get()
            e_breastfeeding_excluded = chk_breastfeeding.get()
            e_hygiene = chk_hygiene.get()
            missing_data = []
            if not int(e_pregnancy_excluded):
                missing_data.append("is not pregnant")
            if not int(e_breastfeeding_excluded):
                missing_data.append("is not breastfeeding")
            if not int(e_hygiene):
                missing_data.append("can maintain good hygiene upon discharge")

            if missing_data:
                if len(missing_data) == 1:
                    missing_data = missing_data[0]
                else:
                    missing_data = (
                        ", ".join(missing_data[:-1]) + " and " + missing_data[-1]
                    )
                messagebox.showerror(
                    "Error",
                    "Please confirm that the patient " + missing_data,
                    parent=window,
                )
                return

            # User is allowed to proceed
            window.withdraw()

            # Write the entered values into self.odict
            first_name = first_name_entry.get()
            last_name = last_name_entry.get()
            if last_name != "" or first_name != "":
                self.odict["data"]["patient_details"]["name"] = (
                    last_name + "^" + first_name
                )
            if last_name == "" and first_name == "":
                if "name" in self.odict["data"]["patient_details"]:
                    self.odict["data"]["patient_details"].pop("name")

            e_pid = pid_entry.get()
            if e_pid != "":
                self.odict["data"]["patient_details"]["id"] = e_pid
            elif "id" in self.odict["data"]["patient_details"]:
                self.odict["data"]["patient_details"].pop("id")

            self.odict["data"]["patient_details"]["dob"] = date2str(e_dob)

            e_sex = sex_combo.get()
            if e_sex != "":
                self.odict["data"]["patient_details"]["sex"] = e_sex

            self.odict["data"]["patient_details"]["type_therapy"] = therapy_combo.get()
            self.odict["data"]["patient_details"]["site"] = site_combo.get()
            self.odict["data"]["patient_details"][
                "num_treatments_in_year"
            ] = e_num_treatments_in_year
            self.odict["data"]["patient_details"][
                "pregnancy_excluded"
            ] = e_pregnancy_excluded
            self.odict["data"]["patient_details"][
                "breastfeeding_excluded"
            ] = e_breastfeeding_excluded
            self.odict["data"]["patient_details"]["hygiene"] = e_hygiene

            # Wipe quantities from self.odict that depend on quantities that the user may have just changed...
            # If key is in self.get_new_odict(), set to empty ordered dict, else pop.

            # type_therapy impacts administered_activity
            if "administered_activity" in self.odict["data"]["administration_details"]:
                self.odict["data"]["administration_details"].pop(
                    "administered_activity"
                )

            # type_therapy is used in retrieve_fit button to get clearance curve_fit
            if "curve_fit" in self.odict["data"]["clearance_data"]:
                self.odict["data"]["clearance_data"].pop("curve_fit")

            # Hence quantities that use the clearance curve_fit are also invaldiated...
            # curve_fit is used in discharge button to calculate discharge_activity
            if "discharge_activity" in self.odict["data"]["patient_discharge"]:
                self.odict["data"]["patient_discharge"].pop("discharge_activity")

            # curve_fit is used in discharge button to calculate calculated_discharge_dose_rate_xm/1m
            if (
                "calculated_discharge_dose_rate_xm"
                in self.odict["data"]["patient_discharge"]
            ):
                self.odict["data"]["patient_discharge"].pop(
                    "calculated_discharge_dose_rate_xm"
                )
            if (
                "calculated_discharge_dose_rate_1m"
                in self.odict["data"]["patient_discharge"]
            ):
                self.odict["data"]["patient_discharge"].pop(
                    "calculated_discharge_dose_rate_1m"
                )

            # curve_fit is used in discharge window to calculate recommended_datetime for discharge
            self.odict["data"]["patient_discharge"][
                "recommended_datetime"
            ] = collections.OrderedDict()

            # curve_fit is used to calculate restrictions
            if "restriction" in self.odict["data"]["restrictions"]:
                if (
                    "restriction_period"
                    in self.odict["data"]["restrictions"]["restriction"][0]
                ):  # min 2 restrictions
                    for child1 in self.odict["data"]["restrictions"]["restriction"]:
                        child1.pop("dose_constraint_corrected")
                        child1.pop("restriction_period")
                        child1.pop("dose")
                        child1.pop("datetime_end")

            # type_therapy is used in discharge window to calculate recommended_datetime for discharge (radionuclide -> RPS 4 activity limit)

            self.odict["data"]["reports_generated"] = "0"

            self.unsaved_data = True

            if int(self.odict["data"]["patient_finished"]):
                self.odict["data"]["patient_finished"] = "0"
                self.odict["data"]["patient_finished_by"] = "0"
                self.viewing_completed_patient_label.place_forget()

            filepath = self.filepath
            if filepath is not None:
                self.root.title(f"*{os.path.basename(filepath)} - {__program_name__}")

            self.generic_updates()

            self.update_buttons()
            self.info_label["text"] = self.get_info_str()
            self.info_label.place(relx=0.01, rely=0.01)

        # Read in previously entered values
        if "name" in self.odict["data"]["patient_details"]:
            name = self.odict["data"]["patient_details"]["name"].split("^")
        else:
            name = None
        if "id" in self.odict["data"]["patient_details"]:
            pid = self.odict["data"]["patient_details"]["id"]
        else:
            pid = None
        if "dob" in self.odict["data"]["patient_details"]:
            dob = str2date(self.odict["data"]["patient_details"]["dob"])
        else:
            dob = None
        if "sex" in self.odict["data"]["patient_details"]:
            sex = self.odict["data"]["patient_details"]["sex"]
        else:
            sex = None
        if "type_therapy" in self.odict["data"]["patient_details"]:
            type_therapy = self.odict["data"]["patient_details"]["type_therapy"]
        else:
            type_therapy = None
        if "site" in self.odict["data"]["patient_details"]:
            site = self.odict["data"]["patient_details"]["site"]
        else:
            site = None
        if "num_treatments_in_year" in self.odict["data"]["patient_details"]:
            num_treatments_in_year = float(
                self.odict["data"]["patient_details"]["num_treatments_in_year"]
            )
            if num_treatments_in_year.is_integer():
                num_treatments_in_year = int(num_treatments_in_year)
        else:
            num_treatments_in_year = None
        if "hygiene" in self.odict["data"]["patient_details"]:
            hygiene = int(self.odict["data"]["patient_details"]["hygiene"])
        else:
            hygiene = None
        if "breastfeeding_excluded" in self.odict["data"]["patient_details"]:
            breastfeeding = int(
                self.odict["data"]["patient_details"]["breastfeeding_excluded"]
            )
        else:
            breastfeeding = None
        if "pregnancy_excluded" in self.odict["data"]["patient_details"]:
            pregnancy = int(self.odict["data"]["patient_details"]["pregnancy_excluded"])
        else:
            pregnancy = None

        # Create the 'Patient Details' window
        window = tk.Toplevel()
        therapy_options = self.therapy_options_df.index.to_list()
        width_therapy = max(
            [len(therapy_options[i]) for i in range(len(therapy_options))]
        )
        window_width = (130 / 20) * width_therapy + 97
        if window_width < 380:
            window_width = 380
        window.geometry(f"{int(window_width)}x350")
        window.title("Patient Details")
        if WINDOWS_OS:
            window.resizable(width=False, height=False)
            window.iconbitmap(self.SOFTWARE_ICON)

        tk.Label(window, text="Last name").grid(
            row=0, column=0, pady=(5, 0), sticky="E"
        )
        frame3 = tk.Frame(window)
        last_name_entry = tk.Entry(frame3, width=20)
        last_name_entry.grid(row=0, column=0)
        tk.Label(
            frame3,
            text="(letter case to appear in reports)",
            fg="green",
            font="Arial 8",
        ).grid(row=0, column=1, padx=(5, 0))
        frame3.grid(row=0, column=1, sticky="W", pady=(5, 0))

        tk.Label(window, text="First name").grid(row=1, column=0, sticky="E")
        first_name_entry = tk.Entry(window, width=20)
        first_name_entry.grid(row=1, column=1, sticky="W")

        if name is not None:
            last_name_entry.insert(0, name[0])
            first_name_entry.insert(0, name[1])

        tk.Label(window, text="ID number").grid(row=2, column=0, sticky="E")
        pid_entry = tk.Entry(window, width=20)
        if pid is not None:
            pid_entry.insert(0, pid)
        pid_entry.grid(row=2, column=1, sticky="W")

        tk.Label(window, text="DOB").grid(row=3, column=0, sticky="E")

        frame1 = tk.Frame(window)
        dob_day_entry = tk.Entry(frame1, width=3)
        dob_day_entry.grid(row=0, column=0, sticky="W")
        tk.Label(frame1, text="(D or DD)", anchor="w").grid(
            row=0, column=1, sticky="W", padx=(0, 5)
        )

        dob_month_entry = tk.Entry(frame1, width=3)
        dob_month_entry.grid(row=0, column=2, sticky="W")
        tk.Label(frame1, text="(M or MM)", anchor="w").grid(
            row=0, column=3, sticky="W", padx=(0, 5)
        )

        dob_year_entry = tk.Entry(frame1, width=5)
        dob_year_entry.grid(row=0, column=4, sticky="W")
        tk.Label(frame1, text="(YYYY)", anchor="w").grid(row=0, column=5, sticky="W")

        frame1.grid(row=3, column=1, sticky="W")

        if dob is not None:
            dob_day_entry.insert(0, dob.day)
            dob_month_entry.insert(0, dob.month)
            dob_year_entry.insert(0, str(dob.year).zfill(4))

        tk.Label(window, text="Sex").grid(row=4, column=0, sticky="E")
        sex_combo = ttk.Combobox(
            window,
            state="readonly",
            values=self.SEX_OPTIONS,
            width=max([len(self.SEX_OPTIONS[i]) for i in range(len(self.SEX_OPTIONS))])
            + 1,
        )
        if sex is not None:
            sex_combo.set(sex)
        sex_combo.grid(row=4, column=1, sticky="W")

        tk.Label(window, text="Type of therapy").grid(
            row=5, column=0, pady=(20, 0), sticky="E"
        )
        therapy_options = self.therapy_options_df.index.to_list()
        if width_therapy < 42:
            width_therapy = 42
        therapy_combo = ttk.Combobox(
            window, state="readonly", values=therapy_options, width=width_therapy
        )
        if type_therapy is not None:
            therapy_combo.set(type_therapy)
        else:
            therapy_combo.set(self.init_vals["therapy_type"])
        therapy_combo.grid(row=5, column=1, pady=(20, 0), sticky="W")

        tk.Label(window, text="Site").grid(row=6, column=0, pady=(5, 0), sticky="E")
        site_options = self.site_options_df.index.to_list()
        if site is not None:
            if site not in site_options:
                site_options.append(site)
        site_combo = ttk.Combobox(
            window,
            state="readonly",
            values=site_options,
            width=max([len(site_options[i]) for i in range(len(site_options))]),
        )
        if site is not None:
            site_combo.set(site)
        else:
            site_combo.set(self.init_vals["site"])
        site_combo.grid(row=6, column=1, sticky="w", pady=(5, 0))

        frame2 = tk.Frame(window)

        tk.Label(frame2, text="Number of treatments expected in a year").grid(
            row=0, column=0, sticky="E", pady=(5, 0)
        )
        num_treatments_in_year_entry = tk.Entry(frame2, width=3)
        if num_treatments_in_year is not None:
            num_treatments_in_year_entry.insert(0, num_treatments_in_year)
        else:
            num_treatments_in_year_entry.insert(
                0, self.init_vals["num_treatments_in_year"]
            )
        num_treatments_in_year_entry.grid(row=0, column=1, pady=(5, 0))

        chk_pregnancy = tk.IntVar()
        tk.Label(frame2, text="Pregnancy excluded").grid(row=1, column=0, sticky="E")
        tk.Checkbutton(frame2, variable=chk_pregnancy).grid(row=1, column=1)

        if pregnancy is not None:
            chk_pregnancy.set(pregnancy)

        chk_breastfeeding = tk.IntVar()
        tk.Label(frame2, text="Breastfeeding excluded").grid(
            row=2, column=0, sticky="E"
        )
        tk.Checkbutton(frame2, variable=chk_breastfeeding).grid(row=2, column=1)

        if breastfeeding is not None:
            chk_breastfeeding.set(breastfeeding)

        chk_hygiene = tk.IntVar()
        tk.Label(
            frame2,
            text="Patient can be properly cared for\n(and good hygiene maintained) at home",
            justify="right",
        ).grid(
            row=3, sticky="E", column=0
        )  # , padx=(40,0))
        tk.Checkbutton(frame2, variable=chk_hygiene).grid(row=3, column=1)

        if hygiene is not None:
            chk_hygiene.set(hygiene)

        frame2.grid(row=7, column=0, pady=(5, 0), columnspan=2)

        tk.Button(
            window, text="Submit", command=lambda: retrieve_patient_details(self)
        ).grid(row=8, column=0, columnspan=2, pady=(15, 0))

    # Administration Details
    def administration_details(self):
        def compute_administered_activity(
            self,
            c_calib_activity=None,
            c_calib_datetime=None,
            c_admin_datetime=None,
            c_residual=None,
        ):
            if c_calib_activity is None:
                if (
                    "calibrated_activity"
                    in self.odict["data"]["administration_details"]
                ):
                    c_calib_activity = float(
                        self.odict["data"]["administration_details"][
                            "calibrated_activity"
                        ]
                    )

            if c_calib_datetime is None:
                if (
                    "calibration_datetime"
                    in self.odict["data"]["administration_details"]
                ):
                    c_calib_datetime = str2datetime(
                        self.odict["data"]["administration_details"][
                            "calibration_datetime"
                        ]
                    )

            if c_admin_datetime is None:
                if (
                    "administration_datetime"
                    in self.odict["data"]["administration_details"]
                ):
                    c_admin_datetime = str2datetime(
                        self.odict["data"]["administration_details"][
                            "administration_datetime"
                        ]
                    )

            if None in [c_calib_activity, c_calib_datetime, c_admin_datetime]:
                return None

            radionuclide = self.therapy_options_df.loc[
                self.odict["data"]["patient_details"]["type_therapy"], "radionuclide"
            ]
            half_life = self.radionuclide_options_df.loc[radionuclide, "half_life"]

            if c_residual is None:
                if "residual_activity" in self.odict["data"]["administration_details"]:
                    c_resid_activity = float(
                        self.odict["data"]["administration_details"][
                            "residual_activity"
                        ]
                    )
                    c_resid_datetime = str2datetime(
                        self.odict["data"]["administration_details"][
                            "residual_datetime"
                        ]
                    )
                    c_residual = (c_resid_activity, c_resid_datetime)

            with warnings.catch_warnings():
                # Don't print the RuntimeWarning
                warnings.filterwarnings("ignore", "overflow encountered in exp")
                c_a0 = Gui.administered_activity(
                    c_calib_activity,
                    c_calib_datetime,
                    c_admin_datetime,
                    half_life,
                    c_residual,
                )

            return c_a0

        # First submit button
        def retrieve_calibration(self):
            r_calib_a = calib_a_entry.get()
            try:
                r_calib_a = float(r_calib_a)
            except ValueError:
                update_window_admin(self)
                messagebox.showerror("Error", "Bad calibrated activity", parent=window)
                return
            if r_calib_a < 0.0 or np.isinf(r_calib_a) or np.isnan(r_calib_a):
                update_window_admin(self)
                messagebox.showerror("Error", "Bad calibrated activity", parent=window)
                return

            if len(calib_year_entry.get()) != 4:
                update_window_admin(self)
                messagebox.showerror(
                    "Error",
                    "Please enter a 4-digit year for calibration",
                    parent=window,
                )
                return
            try:
                r_calib_datetime = datetime(
                    year=int(calib_year_entry.get()),
                    month=int(calib_month_entry.get()),
                    day=int(calib_day_entry.get()),
                    hour=int(calib_hour_entry.get()),
                    minute=int(calib_minute_entry.get()),
                )
            except ValueError:
                update_window_admin(self)
                messagebox.showerror(
                    "Error", "Bad calibration date/time", parent=window
                )
                return

            allow_admin_before_calib = int(
                self.other_settings["allow_admin_before_calib"]
            )
            if not allow_admin_before_calib:
                if (
                    "administration_datetime"
                    in self.odict["data"]["administration_details"]
                ):
                    if r_calib_datetime > str2datetime(
                        self.odict["data"]["administration_details"][
                            "administration_datetime"
                        ]
                    ):
                        update_window_admin(self)
                        messagebox.showerror(
                            "Error",
                            "Calibrated after administration.  This is usually a mistake.\nIf it was intentional, you will need to check the box in 'Settings \u2192 Other' to proceed.",
                            parent=window,
                        )
                        return

            # No failing from here.
            #
            # Returns None if admin time has not previously been entered.
            r_admin_activity = compute_administered_activity(
                self, c_calib_activity=r_calib_a, c_calib_datetime=r_calib_datetime
            )
            info_msg = None
            if r_admin_activity is not None:
                # If calculated A0 is unacceptable, accept the
                # calibration details entered just now and remove the
                # previously entered admin time and residual details.
                # This is an escape hatch to be able to make large
                # changes to calibration and admin times.
                if r_admin_activity <= 0.0:
                    self.odict["data"][
                        "administration_details"
                    ] = collections.OrderedDict()
                    info_msg = f"The previously entered administration time and residual have been CLEARED because they gave an administered activity of {r_admin_activity:.3g} MBq, which is less than or equal to zero."
                    r_admin_activity = None
                elif np.isinf(r_admin_activity):
                    self.odict["data"][
                        "administration_details"
                    ] = collections.OrderedDict()
                    # Since Dorn does not accept admin and residual
                    # times that give an infinitely large residual
                    # (the second submit button), A0 here is
                    # infinitely large unless the XML file was edited
                    # by the user.
                    info_msg = "The previously entered administration time and residual have been CLEARED because they gave an administered activity that is infinitely large or otherwise cannot be computed."
                    r_admin_activity = None
                elif r_admin_activity > MAX_ADMINISTERED_ACTIVITY_MBQ:
                    self.odict["data"][
                        "administration_details"
                    ] = collections.OrderedDict()
                    info_msg = f"The previously entered administration time and residual have been CLEARED because they gave an administered activity of {r_admin_activity:.3g} MBq, which is too large (>{(MAX_ADMINISTERED_ACTIVITY_MBQ*1e-6):g} TBq)."
                    r_admin_activity = None

            self.odict["data"]["administration_details"][
                "calibrated_activity"
            ] = r_calib_a
            self.odict["data"]["administration_details"]["calibration_datetime"] = (
                datetime2str(r_calib_datetime)
            )

            # Calculate and write to odict the derived quantity administered_activity if administration datetime was already submitted
            if r_admin_activity is not None:
                self.odict["data"]["administration_details"][
                    "administered_activity"
                ] = r_admin_activity

            update_window_admin(self)

            # admin_activity is used in discharge button to calculate discharge_activity
            if "discharge_activity" in self.odict["data"]["patient_discharge"]:
                self.odict["data"]["patient_discharge"].pop("discharge_activity")

            # admin_activity is used in discharge window to calculate recommended_datetime for discharge
            self.odict["data"]["patient_discharge"][
                "recommended_datetime"
            ] = collections.OrderedDict()

            # admin_activity is used to get upper bound for curve fit parameter 'a'
            # This invalidates the curve_fit
            if "curve_fit" in self.odict["data"]["clearance_data"]:
                self.odict["data"]["clearance_data"].pop("curve_fit")

            # curve_fit is used in discharge button to calculate discharge_activity
            # curve_fit is used in discharge button to calculate calculated_discharge_dose_rate_1m/xm
            if (
                "calculated_discharge_dose_rate_xm"
                in self.odict["data"]["patient_discharge"]
            ):
                self.odict["data"]["patient_discharge"].pop(
                    "calculated_discharge_dose_rate_xm"
                )
            if (
                "calculated_discharge_dose_rate_1m"
                in self.odict["data"]["patient_discharge"]
            ):
                self.odict["data"]["patient_discharge"].pop(
                    "calculated_discharge_dose_rate_1m"
                )

            # curve_fit is used in discharge window to calculate recommended_datetime for discharge
            # curve_fit is used to calculate restrictions
            if "restriction" in self.odict["data"]["restrictions"]:
                if (
                    "restriction_period"
                    in self.odict["data"]["restrictions"]["restriction"][0]
                ):  # min 2 restrictions
                    for child1 in self.odict["data"]["restrictions"]["restriction"]:
                        child1.pop("dose_constraint_corrected")
                        child1.pop("restriction_period")
                        child1.pop("dose")
                        child1.pop("datetime_end")

            self.odict["data"]["reports_generated"] = "0"

            self.unsaved_data = True

            if int(self.odict["data"]["patient_finished"]):
                self.odict["data"]["patient_finished"] = "0"
                self.odict["data"]["patient_finished_by"] = "0"
                self.viewing_completed_patient_label.place_forget()

            filepath = self.filepath
            if filepath is not None:
                self.root.title(f"*{os.path.basename(filepath)} - {__program_name__}")

            self.generic_updates()
            self.update_buttons()
            # we may have removed the admin time
            self.info_label["text"] = self.get_info_str()
            self.info_label.place(relx=0.01, rely=0.01)

            if info_msg is not None:
                messagebox.showinfo(
                    "Information",
                    info_msg,
                    parent=window,
                )
            else:
                # We didn't clear the administration time and residual.
                if (
                    "administration_datetime"
                    in self.odict["data"]["administration_details"]
                ):
                    if r_calib_datetime > str2datetime(
                        self.odict["data"]["administration_details"][
                            "administration_datetime"
                        ]
                    ):
                        assert allow_admin_before_calib
                        messagebox.showwarning(
                            "Warning",
                            "Calibrated after administration.  This is just a warning.",
                            parent=window,
                        )

        # Second submit button
        def retrieve_administration(self):
            if len(admin_year_entry.get()) != 4:
                update_window_admin(self)
                messagebox.showerror(
                    "Error",
                    "Please enter a 4-digit year for administration",
                    parent=window,
                )
                return
            try:
                r_admin_datetime = datetime(
                    year=int(admin_year_entry.get()),
                    month=int(admin_month_entry.get()),
                    day=int(admin_day_entry.get()),
                    hour=int(admin_hour_entry.get()),
                    minute=int(admin_minute_entry.get()),
                )
            except ValueError:
                update_window_admin(self)
                messagebox.showerror(
                    "Error", "Bad administration date/time", parent=window
                )
                return

            # ignore residual date/time if residual activity is not entered
            r_residual = None
            r_resid_a = resid_a_entry.get()
            if r_resid_a != "":
                try:
                    r_resid_a = float(r_resid_a)
                except ValueError:
                    update_window_admin(self)
                    messagebox.showerror(
                        "Error", "Bad residual activity", parent=window
                    )
                    return
                if r_resid_a < 0.0 or np.isinf(r_resid_a) or np.isnan(r_resid_a):
                    update_window_admin(self)
                    messagebox.showerror(
                        "Error", "Bad residual activity", parent=window
                    )
                    return

                if len(resid_year_entry.get()) != 4:
                    update_window_admin(self)
                    messagebox.showerror(
                        "Error",
                        "Please enter a 4-digit year for residual",
                        parent=window,
                    )
                    return
                try:
                    r_resid_datetime = datetime(
                        year=int(resid_year_entry.get()),
                        month=int(resid_month_entry.get()),
                        day=int(resid_day_entry.get()),
                        hour=int(resid_hour_entry.get()),
                        minute=int(resid_minute_entry.get()),
                    )
                except ValueError:
                    update_window_admin(self)
                    messagebox.showerror(
                        "Error", "Bad residual date/time", parent=window
                    )
                    return
                if r_resid_datetime < r_admin_datetime:
                    update_window_admin(self)
                    messagebox.showerror(
                        "Error",
                        "The residual cannot be measured before administration.  Please try again.",
                        parent=window,
                    )
                    return
                r_residual = (r_resid_a, r_resid_datetime)

            if "calibration_datetime" in self.odict["data"]["administration_details"]:
                if r_admin_datetime < str2datetime(
                    self.odict["data"]["administration_details"]["calibration_datetime"]
                ):
                    allow_admin_before_calib = int(
                        self.other_settings["allow_admin_before_calib"]
                    )
                    if not allow_admin_before_calib:
                        update_window_admin(self)
                        messagebox.showerror(
                            "Error",
                            "Administered before calibration.  This is usually a mistake.\nIf it was intentional, you will need to check the box in 'Settings \u2192 Other' to proceed.",
                            parent=window,
                        )
                        return
                    else:
                        # Show the warning before the error, because it
                        # may help identify the cause of the error.
                        messagebox.showwarning(
                            "Warning",
                            "Administered before calibration.  This is just a warning.",
                            parent=window,
                        )

            # Returns None if calibration info was not been previously
            # entered.
            r_admin_activity = compute_administered_activity(
                self,
                c_admin_datetime=r_admin_datetime,
                # c_residual = None makes it look up the stored value,
                # but in this case the user entered no residual.
                c_residual=(0, r_admin_datetime) if r_residual is None else r_residual,
            )
            if r_admin_activity is not None:
                if r_admin_activity <= 0.0:
                    update_window_admin(self)
                    messagebox.showerror(
                        "Error",
                        f"The calculated administered activity of {r_admin_activity:.3g} MBq is less than or equal to zero.  Please try again.",
                        parent=window,
                    )
                    return
                if np.isinf(r_admin_activity):
                    update_window_admin(self)
                    messagebox.showerror(
                        "Error",
                        "The calculated administered activity is infinitely large or otherwise cannot be computed.  Please try again.",
                        parent=window,
                    )
                    return
                if r_admin_activity > MAX_ADMINISTERED_ACTIVITY_MBQ:
                    update_window_admin(self)
                    messagebox.showerror(
                        "Error",
                        f"The calculated administered activity of {r_admin_activity:.3g} MBq is too large (>{(MAX_ADMINISTERED_ACTIVITY_MBQ*1e-6):g} TBq).  Please try again.",
                        parent=window,
                    )
                    return

            # No failing from here.
            #
            self.odict["data"]["administration_details"]["administration_datetime"] = (
                datetime2str(r_admin_datetime)
            )

            if r_residual is not None:
                self.odict["data"]["administration_details"][
                    "residual_activity"
                ] = r_resid_a
                self.odict["data"]["administration_details"]["residual_datetime"] = (
                    datetime2str(r_resid_datetime)
                )
            else:
                if "residual_activity" in self.odict["data"]["administration_details"]:
                    self.odict["data"]["administration_details"].pop(
                        "residual_activity"
                    )
                if "residual_datetime" in self.odict["data"]["administration_details"]:
                    self.odict["data"]["administration_details"].pop(
                        "residual_datetime"
                    )

            if r_admin_activity is not None:
                self.odict["data"]["administration_details"][
                    "administered_activity"
                ] = r_admin_activity

            update_window_admin(self)

            # In Measured Clearance Data window, when data is submitted, admin_datetime is used to calculate hours_elapsed
            if "measurement" in self.odict["data"]["clearance_data"]:
                if isinstance(
                    self.odict["data"]["clearance_data"]["measurement"], list
                ):
                    for child1 in self.odict["data"]["clearance_data"]["measurement"]:
                        if "hours_elapsed" in child1:
                            child1.pop("hours_elapsed")
                elif (
                    "hours_elapsed"
                    in self.odict["data"]["clearance_data"]["measurement"]
                ):
                    self.odict["data"]["clearance_data"]["measurement"].pop(
                        "hours_elapsed"
                    )

            # This invalidates the curve_fit
            if "curve_fit" in self.odict["data"]["clearance_data"]:
                self.odict["data"]["clearance_data"].pop("curve_fit")

            # curve_fit is used in discharge button to calculate discharge_activity
            if "discharge_activity" in self.odict["data"]["patient_discharge"]:
                self.odict["data"]["patient_discharge"].pop("discharge_activity")

            # curve_fit is used in discharge button to calculate calculated_discharge_dose_rate_1m/xm
            if (
                "calculated_discharge_dose_rate_xm"
                in self.odict["data"]["patient_discharge"]
            ):
                self.odict["data"]["patient_discharge"].pop(
                    "calculated_discharge_dose_rate_xm"
                )
            if (
                "calculated_discharge_dose_rate_1m"
                in self.odict["data"]["patient_discharge"]
            ):
                self.odict["data"]["patient_discharge"].pop(
                    "calculated_discharge_dose_rate_1m"
                )

            # curve_fit is used in discharge window to calculate recommended_datetime for discharge
            self.odict["data"]["patient_discharge"][
                "recommended_datetime"
            ] = collections.OrderedDict()

            # curve_fit is used to calculate restrictions
            if "restriction" in self.odict["data"]["restrictions"]:
                if (
                    "restriction_period"
                    in self.odict["data"]["restrictions"]["restriction"][0]
                ):  # min 2 restrictions
                    for child1 in self.odict["data"]["restrictions"]["restriction"]:
                        child1.pop("dose_constraint_corrected")
                        child1.pop("restriction_period")
                        child1.pop("dose")
                        child1.pop("datetime_end")

            # admin_datetime and admin_activity are used in discharge button to calculate discharge_activity
            # admin_datetime is used in discharge button to calculate calculated_discharge_dose_rate_xm/1m from curve_fit
            # admin_datetime and administered_activity are used in discharge window to calculate recommended_datetime for discharge
            # admin datetime is used to calculate restrictions (already wiped)
            self.odict["data"]["reports_generated"] = "0"

            self.unsaved_data = True

            if int(self.odict["data"]["patient_finished"]):
                self.odict["data"]["patient_finished"] = "0"
                self.odict["data"]["patient_finished_by"] = "0"
                self.viewing_completed_patient_label.place_forget()

            filepath = self.filepath
            if filepath is not None:
                self.root.title(f"*{os.path.basename(filepath)} - {__program_name__}")

            self.generic_updates()
            self.update_buttons()
            self.info_label["text"] = self.get_info_str()
            self.info_label.place(relx=0.01, rely=0.01)

        def update_window_admin(self):
            # Retrieve the stored values
            if "calibrated_activity" in self.odict["data"]["administration_details"]:
                calib_a = float(
                    self.odict["data"]["administration_details"]["calibrated_activity"]
                )
                if calib_a.is_integer():
                    calib_a = int(calib_a)
            else:
                calib_a = None
            if "calibration_datetime" in self.odict["data"]["administration_details"]:
                calib_datetime = str2datetime(
                    self.odict["data"]["administration_details"]["calibration_datetime"]
                )
            else:
                calib_datetime = None
            if (
                "administration_datetime"
                in self.odict["data"]["administration_details"]
            ):
                admin_datetime = str2datetime(
                    self.odict["data"]["administration_details"][
                        "administration_datetime"
                    ]
                )
            else:
                admin_datetime = None
            if "administered_activity" in self.odict["data"]["administration_details"]:
                admin_activity = float(
                    self.odict["data"]["administration_details"][
                        "administered_activity"
                    ]
                )
            else:
                admin_activity = None
            if "residual_datetime" in self.odict["data"]["administration_details"]:
                resid_datetime = str2datetime(
                    self.odict["data"]["administration_details"]["residual_datetime"]
                )
            else:
                resid_datetime = None
            if "residual_activity" in self.odict["data"]["administration_details"]:
                resid_a = float(
                    self.odict["data"]["administration_details"]["residual_activity"]
                )
                if resid_a.is_integer():
                    resid_a = int(resid_a)
            else:
                resid_a = None

            # update the values showing in the window
            calib_a_entry.delete(0, tk.END)
            calib_hour_entry.delete(0, tk.END)
            calib_minute_entry.delete(0, tk.END)
            calib_day_entry.delete(0, tk.END)
            calib_month_entry.delete(0, tk.END)
            calib_year_entry.delete(0, tk.END)
            admin_hour_entry.delete(0, tk.END)
            admin_minute_entry.delete(0, tk.END)
            admin_day_entry.delete(0, tk.END)
            admin_month_entry.delete(0, tk.END)
            admin_year_entry.delete(0, tk.END)
            resid_a_entry.delete(0, tk.END)
            resid_hour_entry.delete(0, tk.END)
            resid_minute_entry.delete(0, tk.END)
            resid_day_entry.delete(0, tk.END)
            resid_month_entry.delete(0, tk.END)
            resid_year_entry.delete(0, tk.END)
            if calib_a is not None:
                calib_a_entry.insert(0, calib_a)
            if calib_datetime is not None:
                calib_hour_entry.insert(0, calib_datetime.hour)
                calib_minute_entry.insert(0, calib_datetime.minute)
                calib_day_entry.insert(0, calib_datetime.day)
                calib_month_entry.insert(0, calib_datetime.month)
                calib_year_entry.insert(0, str(calib_datetime.year).zfill(4))
            if admin_datetime is not None:
                admin_hour_entry.insert(0, admin_datetime.hour)
                admin_minute_entry.insert(0, admin_datetime.minute)
                admin_day_entry.insert(0, admin_datetime.day)
                admin_month_entry.insert(0, admin_datetime.month)
                admin_year_entry.insert(0, str(admin_datetime.year).zfill(4))
            elif calib_datetime is not None:
                # Prefill entries
                admin_day_entry.insert(0, calib_datetime.day)
                admin_month_entry.insert(0, calib_datetime.month)
                admin_year_entry.insert(0, str(calib_datetime.year).zfill(4))
            if resid_a is not None:
                resid_a_entry.insert(0, resid_a)
            if resid_datetime is not None:
                resid_hour_entry.insert(0, resid_datetime.hour)
                resid_minute_entry.insert(0, resid_datetime.minute)
                resid_day_entry.insert(0, resid_datetime.day)
                resid_month_entry.insert(0, resid_datetime.month)
                resid_year_entry.insert(0, str(resid_datetime.year).zfill(4))
            elif admin_datetime is not None:
                # Prefill entries
                resid_day_entry.insert(0, admin_datetime.day)
                resid_month_entry.insert(0, admin_datetime.month)
                resid_year_entry.insert(0, str(admin_datetime.year).zfill(4))
            elif calib_datetime is not None:
                # Prefill entries
                resid_day_entry.insert(0, calib_datetime.day)
                resid_month_entry.insert(0, calib_datetime.month)
                resid_year_entry.insert(0, str(calib_datetime.year).zfill(4))
            if admin_activity is not None:
                admin_activity_str.set(
                    f"The administered activity was {admin_activity:g} MBq"
                )
                l1.grid(row=2, pady=(10, 0))
                button_ok.grid(row=3, pady=(10, 0))
            else:
                l1.grid_forget()
                button_ok.grid_forget()

        window = tk.Toplevel()
        window.geometry("342x400")
        window.title("Administration Details")
        if WINDOWS_OS:
            window.resizable(width=False, height=False)
            window.iconbitmap(self.SOFTWARE_ICON)

        frame1 = tk.LabelFrame(window)

        tk.Label(frame1, text="Calibrated activity").grid(row=0, sticky="W")
        calib_a_entry = tk.Entry(frame1, width=6)
        calib_a_entry.grid(row=0, column=1, sticky="E")
        tk.Label(frame1, text="MBq").grid(row=0, column=2, sticky="W")

        tk.Label(frame1, text="Calibration time").grid(row=1, sticky="W")
        calib_hour_entry = tk.Entry(frame1, width=3)
        calib_hour_entry.grid(row=1, column=1, sticky="E")
        tk.Label(frame1, text="(hh)").grid(row=1, column=2, padx=(0, 5), sticky="W")
        calib_minute_entry = tk.Entry(frame1, width=3)
        calib_minute_entry.grid(row=1, column=3)
        tk.Label(frame1, text="(mm)").grid(row=1, column=4, sticky="W")
        tk.Label(frame1, text="(24-hour time)", fg="green", font="Arial 8").grid(
            row=1, column=5, columnspan=2
        )

        tk.Label(frame1, text="Calibration date").grid(row=2, sticky="W")
        calib_day_entry = tk.Entry(frame1, width=3)
        calib_day_entry.grid(row=2, column=1, sticky="E")
        tk.Label(frame1, text="(DD)").grid(row=2, column=2, padx=(0, 5), sticky="W")
        calib_month_entry = tk.Entry(frame1, width=3)
        calib_month_entry.grid(row=2, column=3)
        tk.Label(frame1, text="(MM)").grid(row=2, column=4, padx=(0, 5), sticky="W")
        calib_year_entry = tk.Entry(frame1, width=5)
        calib_year_entry.grid(row=2, column=5)
        tk.Label(frame1, text="(YYYY)").grid(row=2, column=6, sticky="W")

        tk.Button(
            frame1, text="Submit", command=lambda: retrieve_calibration(self)
        ).grid(row=3, columnspan=7, pady=(10, 10))

        frame1.grid(pady=(10, 10), padx=(10, 10), sticky="WE")

        frame2 = tk.LabelFrame(window)

        tk.Label(frame2, text="Administration time").grid(row=0, sticky="W")
        admin_hour_entry = tk.Entry(frame2, width=3)
        admin_hour_entry.grid(row=0, column=1, sticky="E")
        tk.Label(frame2, text="(hh)").grid(row=0, column=2, padx=(0, 5), sticky="W")
        admin_minute_entry = tk.Entry(frame2, width=3)
        admin_minute_entry.grid(row=0, column=3)
        tk.Label(frame2, text="(mm)").grid(row=0, column=4, sticky="W")

        tk.Label(frame2, text="Administration date").grid(row=1, sticky="W")
        admin_day_entry = tk.Entry(frame2, width=3)
        admin_day_entry.grid(row=1, column=1, sticky="E")
        tk.Label(frame2, text="(DD)").grid(row=1, column=2, padx=(0, 5), sticky="W")
        admin_month_entry = tk.Entry(frame2, width=3)
        admin_month_entry.grid(row=1, column=3)
        tk.Label(frame2, text="(MM)").grid(row=1, column=4, padx=(0, 5), sticky="W")
        admin_year_entry = tk.Entry(frame2, width=5)
        admin_year_entry.grid(row=1, column=5)
        tk.Label(frame2, text="(YYYY)").grid(row=1, column=6, sticky="W")

        tk.Label(frame2, text="Residual (optional)", font="Arial 9 bold").grid(
            row=2, pady=(10, 0), columnspan=7
        )
        tk.Label(frame2, text="Residual activity").grid(row=3, sticky="W")
        resid_a_entry = tk.Entry(frame2, width=6)
        resid_a_entry.grid(row=3, column=1, sticky="E")
        tk.Label(frame2, text="MBq").grid(row=3, column=2, sticky="W")

        tk.Label(frame2, text="Residual time").grid(row=4, sticky="W")
        resid_hour_entry = tk.Entry(frame2, width=3)
        resid_hour_entry.grid(row=4, column=1, sticky="E")
        tk.Label(frame2, text="(hh)").grid(row=4, column=2, padx=(0, 5), sticky="W")
        resid_minute_entry = tk.Entry(frame2, width=3)
        resid_minute_entry.grid(row=4, column=3)
        tk.Label(frame2, text="(mm)").grid(row=4, column=4, sticky="W")

        tk.Label(frame2, text="Residual date").grid(row=5, sticky="W")
        resid_day_entry = tk.Entry(frame2, width=3)
        resid_day_entry.grid(row=5, column=1, sticky="E")
        tk.Label(frame2, text="(DD)").grid(row=5, column=2, padx=(0, 5), sticky="W")
        resid_month_entry = tk.Entry(frame2, width=3)
        resid_month_entry.grid(row=5, column=3)
        tk.Label(frame2, text="(MM)").grid(row=5, column=4, padx=(0, 5), sticky="W")
        resid_year_entry = tk.Entry(frame2, width=5)
        resid_year_entry.grid(row=5, column=5)
        tk.Label(frame2, text="(YYYY)").grid(row=5, column=6, sticky="W")

        tk.Button(
            frame2, text="Submit", command=lambda: retrieve_administration(self)
        ).grid(row=6, columnspan=7, pady=(10, 10))

        frame2.grid(row=1)

        admin_activity_str = tk.StringVar()
        l1 = tk.Label(window, textvariable=admin_activity_str)
        button_ok = tk.Button(window, text="OK", command=window.withdraw)
        update_window_admin(self)

    # Measured Clearance Data
    def clearance_data(self):
        from scipy.optimize import curve_fit
        import matplotlib.pyplot as plt
        from matplotlib.figure import Figure
        from matplotlib.backends.backend_tkagg import (
            FigureCanvasTkAgg,
            NavigationToolbar2Tk,
        )

        def fit_exponential(self, time_lapse, dose_rate_xm):
            # Fit to y = a * np.exp(-b * t)
            # i.e. fit_parameters = [a, b]
            radionuclide = self.therapy_options_df.loc[
                self.odict["data"]["patient_details"]["type_therapy"], "radionuclide"
            ]
            half_life_phys = self.radionuclide_options_df.loc[radionuclide, "half_life"]
            spec_dose_rate = self.radionuclide_options_df.loc[
                radionuclide, "specific_dose_rate_1m"
            ]
            if np.isnan(spec_dose_rate):
                spec_dose_rate = (
                    0.6820  # largest value I can find for any radionuclide (Bi-206)
                )
            a0 = float(
                self.odict["data"]["administration_details"]["administered_activity"]
            )
            measurement_distance = float(
                self.odict["data"]["clearance_data"]["measurement_distance"]
            )
            a_bound = a0 * spec_dose_rate / (measurement_distance**1.5)

            mbounds = ([0.0, np.log(2) / half_life_phys], [a_bound, 1000.0])

            a_guess = a_bound * 0.5
            b_guess = np.log(2) * 3 / half_life_phys
            guess = np.array([a_guess, b_guess])

            # Levenberg-Marquardt algorithm
            fit_parameters, _ = curve_fit(
                func_exp,
                time_lapse,
                dose_rate_xm,
                p0=guess,
                bounds=mbounds,
                method="trf",
            )

            dose_rate_xm_init = fit_parameters[0]
            effective_half_life = np.log(2) / fit_parameters[1]
            meaningful_parameters = [dose_rate_xm_init, effective_half_life]

            return fit_parameters, meaningful_parameters

        def fit_biexponential(self, time_lapse, dose_rate_xm):
            # Fit to y = a * ( b * np.exp(-c * t) + (1-b) * np.exp(-d * t) )
            # i.e. fit_parameters = [a, b, c, d]
            radionuclide = self.therapy_options_df.loc[
                self.odict["data"]["patient_details"]["type_therapy"], "radionuclide"
            ]
            half_life_phys = self.radionuclide_options_df.loc[radionuclide, "half_life"]
            spec_dose_rate = self.radionuclide_options_df.loc[
                radionuclide, "specific_dose_rate_1m"
            ]
            if np.isnan(spec_dose_rate):
                spec_dose_rate = (
                    0.682  # largest value I can find for any radionuclide (Bi-206)
                )
            a0 = float(
                self.odict["data"]["administration_details"]["administered_activity"]
            )
            measurement_distance = float(
                self.odict["data"]["clearance_data"]["measurement_distance"]
            )
            a_bound = a0 * spec_dose_rate / (measurement_distance**1.5)

            mbounds = (
                [0.0, 0.0, np.log(2) / half_life_phys, np.log(2) / half_life_phys],
                [a_bound, 1.0, 200.0, 200.0],
            )

            a_guess = a_bound * 0.5
            b_guess = 0.5
            c_guess = np.log(2) * 3 / half_life_phys
            d_guess = np.log(2) * 3 / half_life_phys
            guess = np.array([a_guess, b_guess, c_guess, d_guess])

            # Levenberg-Marquardt algorithm
            fit_parameters, _ = curve_fit(
                func_biexp,
                time_lapse,
                dose_rate_xm,
                p0=guess,
                bounds=mbounds,
                method="trf",
            )

            dose_rate_xm_init = fit_parameters[0]
            fraction_1 = fit_parameters[1]
            half_life_1 = np.log(2) / fit_parameters[2]
            half_life_2 = np.log(2) / fit_parameters[3]
            meaningful_parameters = [
                dose_rate_xm_init,
                fraction_1,
                half_life_1,
                half_life_2,
            ]

            return fit_parameters, meaningful_parameters

        # Plot on plot_dr as much as can be plotted
        def plot_clearance(self):
            plot_dr.clear()

            hrs_elapsed = []
            dr_xm = []
            hrs_elapsed_excluded = []
            dr_xm_excluded = []
            if "measurement" in self.odict["data"]["clearance_data"]:
                if isinstance(
                    self.odict["data"]["clearance_data"]["measurement"], list
                ):
                    if (
                        "hours_elapsed"
                        in self.odict["data"]["clearance_data"]["measurement"][0]
                    ):
                        hrs_elapsed = [
                            float(child1["hours_elapsed"])
                            for child1 in self.odict["data"]["clearance_data"][
                                "measurement"
                            ]
                            if not int(child1["exclude"])
                        ]
                        dr_xm = [
                            float(child1["doserate_corrected"])
                            for child1 in self.odict["data"]["clearance_data"][
                                "measurement"
                            ]
                            if not int(child1["exclude"])
                        ]
                        hrs_elapsed_excluded = [
                            float(child1["hours_elapsed"])
                            for child1 in self.odict["data"]["clearance_data"][
                                "measurement"
                            ]
                            if int(child1["exclude"])
                        ]
                        dr_xm_excluded = [
                            float(child1["doserate_corrected"])
                            for child1 in self.odict["data"]["clearance_data"][
                                "measurement"
                            ]
                            if int(child1["exclude"])
                        ]
                else:
                    if (
                        "hours_elapsed"
                        in self.odict["data"]["clearance_data"]["measurement"]
                    ):
                        if int(
                            self.odict["data"]["clearance_data"]["measurement"][
                                "exclude"
                            ]
                        ):
                            hrs_elapsed_excluded = [
                                float(
                                    self.odict["data"]["clearance_data"]["measurement"][
                                        "hours_elapsed"
                                    ]
                                )
                            ]
                            dr_xm_excluded = [
                                float(
                                    self.odict["data"]["clearance_data"]["measurement"][
                                        "doserate_corrected"
                                    ]
                                )
                            ]
                        else:
                            hrs_elapsed = [
                                float(
                                    self.odict["data"]["clearance_data"]["measurement"][
                                        "hours_elapsed"
                                    ]
                                )
                            ]
                            dr_xm = [
                                float(
                                    self.odict["data"]["clearance_data"]["measurement"][
                                        "doserate_corrected"
                                    ]
                                )
                            ]

            if hrs_elapsed or dr_xm or hrs_elapsed_excluded or dr_xm_excluded:
                if hrs_elapsed:
                    plot_dr.plot(
                        hrs_elapsed,
                        dr_xm,
                        "o",
                        label="Data",
                        markeredgecolor="k",
                        markerfacecolor=(86 / 255, 180 / 255, 233 / 255),
                    )
                if hrs_elapsed_excluded:
                    plot_dr.plot(
                        hrs_elapsed_excluded,
                        dr_xm_excluded,
                        "x",
                        label="Excluded",
                        markeredgecolor="red",
                        markerfacecolor="red",
                    )
                plot_dr.legend()

            if "measurement_distance" in self.odict["data"]["clearance_data"]:
                p_measurement_distance = float(
                    self.odict["data"]["clearance_data"]["measurement_distance"]
                )
                if p_measurement_distance.is_integer():
                    p_measurement_distance = int(p_measurement_distance)
                plot_dr.set_ylabel(
                    "Dose rate at {} m (\N{GREEK SMALL LETTER MU}Sv/h)".format(
                        p_measurement_distance
                    )
                )
            else:
                plot_dr.set_ylabel("Dose rate (\N{GREEK SMALL LETTER MU}Sv/h)")

            if "curve_fit" in self.odict["data"]["clearance_data"]:
                p_model = self.odict["data"]["clearance_data"]["curve_fit"]["model"]
                fit_parameters = [
                    float(a)
                    for a in self.odict["data"]["clearance_data"]["curve_fit"][
                        "fit_parameters"
                    ].values()
                ]
                meaningful_parameters = [
                    float(a)
                    for a in self.odict["data"]["clearance_data"]["curve_fit"][
                        "meaningful_parameters"
                    ].values()
                ]

                time_more = np.linspace(
                    0.0, max(hrs_elapsed + hrs_elapsed_excluded), num=1000
                )
                if p_model == "biexponential":
                    plot_dr.plot(
                        time_more,
                        func_biexp(time_more, *fit_parameters),
                        label="Biexponential",
                        color=(0, 158 / 255, 115 / 255),
                    )
                    (
                        dose_rate_xm_init,
                        fraction_1,
                        half_life_1,
                        half_life_2,
                    ) = meaningful_parameters
                    plot_dr.annotate(
                        "D'(0) = {:.1f} \N{GREEK SMALL LETTER MU}Sv/h\n{:.1f}% {:.1f} hours\n{:.1f}% {:.1f} hours".format(
                            dose_rate_xm_init,
                            fraction_1 * 100,
                            half_life_1,
                            (1.0 - fraction_1) * 100,
                            half_life_2,
                        ),
                        xy=(0.6, 0.62),
                        xycoords="figure fraction",
                    )
                elif p_model == "exponential":
                    plot_dr.plot(
                        time_more,
                        func_exp(time_more, *fit_parameters),
                        label="Exponential",
                        color=(204 / 255, 121 / 255, 167 / 255),
                    )
                    plot_dr.annotate(
                        "D'(0) = {:.1f} \N{GREEK SMALL LETTER MU}Sv/h\n{:.1f} hours".format(
                            meaningful_parameters[0], meaningful_parameters[1]
                        ),
                        xy=(0.6, 0.62),
                        xycoords="figure fraction",
                    )

                handles, labels = plot_dr.get_legend_handles_labels()
                new_handles = [handles[1], handles[0]]
                new_labels = [labels[1], labels[0]]
                if len(handles) == 3:
                    new_handles.append(handles[2])
                    new_labels.append(labels[2])
                plot_dr.legend(new_handles, new_labels)

                rss = float(
                    self.odict["data"]["clearance_data"]["curve_fit"][
                        "sum_of_squared_residuals"
                    ]
                )
                plot_dr.annotate(
                    "Sum of squared residuals\n= {:.1f} \N{GREEK SMALL LETTER MU}Sv$^2$/h$^2$".format(
                        rss
                    ),
                    xy=(0.6, 0.53),
                    xycoords="figure fraction",
                    fontsize=8,
                    ha="left",
                )

                # Warn the user if the last included measurement differs too much from the curve fit
                hrs_elapsed_last = max(hrs_elapsed)
                dr_xm_last = dr_xm[hrs_elapsed.index(hrs_elapsed_last)]
                dr_xm_last_curve = 0.0
                if p_model == "biexponential":
                    dr_xm_last_curve = func_biexp(hrs_elapsed_last, *fit_parameters)
                elif p_model == "exponential":
                    dr_xm_last_curve = func_exp(hrs_elapsed_last, *fit_parameters)
                rel_diff = (dr_xm_last - dr_xm_last_curve) / dr_xm_last_curve
                if abs(rel_diff) > 0.2:
                    plot_dr.set_title(
                        "WARNING:\nThe last dose rate measurement differs from the curve fit by more than 20%\n({:.1f} vs {:.1f} \N{GREEK SMALL LETTER MU}Sv/h)".format(
                            dr_xm_last, dr_xm_last_curve
                        ),
                        loc="left",
                        fontsize=8,
                        color="red",
                        pad=10,
                    )

            plot_dr.set_xlabel("Time from administration (h)")
            plot_dr.set_ylim(bottom=0)
            plot_dr.set_xlim(left=0)
            plot_dr.grid(alpha=0.2)

            canvas.draw()

        # Submit and Plot button calls the functions retrieve_clearance_data, retrieve_fit_choice, then plot_clearance
        def retrieve_clearance_data(self):
            try:
                r_detector_calibration_factor = float(
                    detector_calibration_factor_entry.get()
                )
            except ValueError:
                update_window_measure(self)
                messagebox.showerror(
                    "Error", "Bad detector calibration factor", parent=window
                )
                return
            if (
                r_detector_calibration_factor <= 0.0
                or np.isinf(r_detector_calibration_factor)
                or np.isnan(r_detector_calibration_factor)
            ):
                update_window_measure(self)
                messagebox.showerror(
                    "Error", "Bad detector calibration factor", parent=window
                )
                return

            try:
                r_measurement_distance = float(measurement_distance_entry.get())
            except ValueError:
                update_window_measure(self)
                messagebox.showerror("Error", "Bad measurement distance", parent=window)
                return

            if not (1.0 <= r_measurement_distance <= 3.0):
                update_window_measure(self)
                messagebox.showerror(
                    "Error",
                    "Measurement distance must be between 1 and 3 m",
                    parent=window,
                )
                return

            doserate_list = [e.get() for e in doserate]
            for qdoserate in doserate_list:
                if qdoserate:
                    try:
                        qqdoserate = float(qdoserate)
                    except ValueError:
                        update_window_measure(self)
                        messagebox.showerror(
                            "Error",
                            f"Bad dose rate measurement: '{qdoserate}'",
                            parent=window,
                        )
                        return

                    if qqdoserate < 0.0 or np.isinf(qqdoserate) or np.isnan(qqdoserate):
                        update_window_measure(self)
                        messagebox.showerror(
                            "Error",
                            f"Bad dose rate measurement: '{qqdoserate}'",
                            parent=window,
                        )
                        return

            year_list = [e.get() for e in year]
            month_list = [e.get() for e in month]
            day_list = [e.get() for e in day]
            hour_list = [e.get() for e in hour]
            minute_list = [e.get() for e in minute]
            chks_list = [e.get() for e in chks]

            year_list = [
                year_list[i] for i in range(len(doserate_list)) if doserate_list[i]
            ]
            month_list = [
                month_list[i] for i in range(len(doserate_list)) if doserate_list[i]
            ]
            day_list = [
                day_list[i] for i in range(len(doserate_list)) if doserate_list[i]
            ]
            hour_list = [
                hour_list[i] for i in range(len(doserate_list)) if doserate_list[i]
            ]
            minute_list = [
                minute_list[i] for i in range(len(doserate_list)) if doserate_list[i]
            ]
            chks_list = [
                chks_list[i] for i in range(len(doserate_list)) if doserate_list[i]
            ]

            doserate_list = [x for x in doserate_list if x]

            for r_year in year_list:
                if len(r_year) != 4:
                    update_window_measure(self)
                    messagebox.showerror(
                        "Error", "Please enter 4-digit years", parent=window
                    )
                    return

            datetime_str_list = []
            for r_year, r_month, r_day, r_hour, r_minute in zip(
                year_list, month_list, day_list, hour_list, minute_list
            ):
                try:
                    r_datetime = datetime(
                        year=int(r_year),
                        month=int(r_month),
                        day=int(r_day),
                        hour=int(r_hour),
                        minute=int(r_minute),
                    )
                    datetime_str_list.append(datetime2str(r_datetime))
                except ValueError:
                    update_window_measure(self)
                    messagebox.showerror(
                        "Error",
                        "Bad date/time for at least one dose rate measurement",
                        parent=window,
                    )
                    return

            r_admin_datetime = str2datetime(
                self.odict["data"]["administration_details"]["administration_datetime"]
            )
            for datetime_str in datetime_str_list:
                if str2datetime(datetime_str) < r_admin_datetime:
                    update_window_measure(self)
                    messagebox.showerror(
                        "Error",
                        "Dose rate measurement before administration",
                        parent=window,
                    )
                    return

            self.odict["data"]["clearance_data"][
                "detector_calibration_factor"
            ] = r_detector_calibration_factor
            self.odict["data"]["clearance_data"][
                "measurement_distance"
            ] = r_measurement_distance

            # Overwrite measurements in self.odict
            if "measurement" in self.odict["data"]["clearance_data"]:
                self.odict["data"]["clearance_data"].pop("measurement")

            if doserate_list:
                self.odict["data"]["clearance_data"]["measurement"] = []
                for i in range(len(doserate_list)):
                    record = collections.OrderedDict(
                        [
                            ("datetime", datetime_str_list[i]),
                            ("doserate", doserate_list[i]),
                        ]
                    )
                    if chks_list[i]:
                        record["exclude"] = "1"
                    else:
                        record["exclude"] = "0"
                    record["hours_elapsed"] = (
                        str2datetime(record["datetime"]) - r_admin_datetime
                    ).total_seconds() / 3600
                    record["doserate_corrected"] = float(record["doserate"]) * float(
                        self.odict["data"]["clearance_data"][
                            "detector_calibration_factor"
                        ]
                    )
                    self.odict["data"]["clearance_data"]["measurement"].append(record)

                if len(self.odict["data"]["clearance_data"]["measurement"]) == 1:
                    self.odict["data"]["clearance_data"]["measurement"] = self.odict[
                        "data"
                    ]["clearance_data"]["measurement"][0]

            # hours_elapsed and doserate_corrected are used in retrieve_fit_choice to get curve_fit
            if "curve_fit" in self.odict["data"]["clearance_data"]:
                self.odict["data"]["clearance_data"].pop("curve_fit")

            # curve_fit is used in discharge button to calculate discharge_activity
            if "discharge_activity" in self.odict["data"]["patient_discharge"]:
                self.odict["data"]["patient_discharge"].pop("discharge_activity")

            # curve_fit is used in discharge button to calculate calculated_discharge_dose_rate_xm/1m
            if (
                "calculated_discharge_dose_rate_xm"
                in self.odict["data"]["patient_discharge"]
            ):
                self.odict["data"]["patient_discharge"].pop(
                    "calculated_discharge_dose_rate_xm"
                )
            if (
                "calculated_discharge_dose_rate_1m"
                in self.odict["data"]["patient_discharge"]
            ):
                self.odict["data"]["patient_discharge"].pop(
                    "calculated_discharge_dose_rate_1m"
                )

            # curve_fit is used in discharge window to calculate recommended_datetime for discharge
            self.odict["data"]["patient_discharge"][
                "recommended_datetime"
            ] = collections.OrderedDict()

            # curve_fit is used to calculate restrictions
            if "restriction" in self.odict["data"]["restrictions"]:
                if (
                    "restriction_period"
                    in self.odict["data"]["restrictions"]["restriction"][0]
                ):  # min 2 restrictions
                    for child1 in self.odict["data"]["restrictions"]["restriction"]:
                        child1.pop("dose_constraint_corrected")
                        child1.pop("restriction_period")
                        child1.pop("dose")
                        child1.pop("datetime_end")

            # measurement distance is used in the calculation of recommended datetime (datetime/based on) - discharge -> compute recommended datetime
            # measurement distance is used in the calculation of discharge dose rate 1m and xm - discharge -> compute discharge dose rate in curve fit
            # measurement distance is used in calculating restrictions (compute restrictions and view restrictions)
            # measurement distance is used in generating reports

            self.odict["data"]["reports_generated"] = "0"

            self.unsaved_data = True

            if int(self.odict["data"]["patient_finished"]):
                self.odict["data"]["patient_finished"] = "0"
                self.odict["data"]["patient_finished_by"] = "0"
                self.viewing_completed_patient_label.place_forget()

            filepath = self.filepath
            if filepath is not None:
                self.root.title(f"*{os.path.basename(filepath)} - {__program_name__}")

            self.update_buttons()

            button_view_residuals["state"] = "disabled"

        def retrieve_fit_choice(self):
            if "measurement" not in self.odict["data"]["clearance_data"]:
                return
            if not isinstance(
                self.odict["data"]["clearance_data"]["measurement"], list
            ):
                return
            if (
                not "hours_elapsed"
                in self.odict["data"]["clearance_data"]["measurement"][0]
            ):
                return

            hrs_elapsed = [
                float(child1["hours_elapsed"])
                for child1 in self.odict["data"]["clearance_data"]["measurement"]
                if not int(child1["exclude"])
            ]
            dr_xm = [
                float(child1["doserate_corrected"])
                for child1 in self.odict["data"]["clearance_data"]["measurement"]
                if not int(child1["exclude"])
            ]

            if len(hrs_elapsed) < 2:
                return

            # Goodness of fit (sum of squared residuals)
            residuals = 0

            with warnings.catch_warnings():
                warnings.filterwarnings(
                    "ignore", "Covariance of the parameters could not be estimated"
                )
                if var1.get() == 2:
                    try:
                        params_biexp, meaningful_params_biexp = fit_biexponential(
                            self, hrs_elapsed, dr_xm
                        )
                    except Exception as e:
                        update_window_measure(self)
                        messagebox.showerror(
                            "Error",
                            f"Unable to fit a biexponential curve.\n{e}",
                            parent=window,
                        )
                        return
                    self.odict["data"]["clearance_data"]["curve_fit"] = (
                        collections.OrderedDict(
                            [
                                ("generic", "0"),
                                ("model", "biexponential"),
                                (
                                    "fit_parameters",
                                    collections.OrderedDict(
                                        [
                                            ("a", params_biexp[0]),
                                            ("b", params_biexp[1]),
                                            ("c", params_biexp[2]),
                                            ("d", params_biexp[3]),
                                        ]
                                    ),
                                ),
                                (
                                    "meaningful_parameters",
                                    collections.OrderedDict(
                                        [
                                            (
                                                "dose_rate_xm_init",
                                                meaningful_params_biexp[0],
                                            ),
                                            ("fraction_1", meaningful_params_biexp[1]),
                                            ("half_life_1", meaningful_params_biexp[2]),
                                            ("half_life_2", meaningful_params_biexp[3]),
                                        ]
                                    ),
                                ),
                            ]
                        )
                    )
                    residuals = np.array(dr_xm) - func_biexp(
                        np.array(hrs_elapsed), *params_biexp
                    )
                elif var1.get() == 1:
                    try:
                        params_exp, meaningful_params_exp = fit_exponential(
                            self, hrs_elapsed, dr_xm
                        )
                    except Exception as e:
                        update_window_measure(self)
                        messagebox.showerror(
                            "Error",
                            f"Unable to fit an exponential curve.\n{e}",
                            parent=window,
                        )
                        return
                    self.odict["data"]["clearance_data"]["curve_fit"] = (
                        collections.OrderedDict(
                            [
                                ("generic", "0"),
                                ("model", "exponential"),
                                (
                                    "fit_parameters",
                                    collections.OrderedDict(
                                        [("a", params_exp[0]), ("b", params_exp[1])]
                                    ),
                                ),
                                (
                                    "meaningful_parameters",
                                    collections.OrderedDict(
                                        [
                                            (
                                                "dose_rate_xm_init",
                                                meaningful_params_exp[0],
                                            ),
                                            (
                                                "effective_half_life",
                                                meaningful_params_exp[1],
                                            ),
                                        ]
                                    ),
                                ),
                            ]
                        )
                    )
                    residuals = np.array(dr_xm) - func_exp(
                        np.array(hrs_elapsed), *params_exp
                    )

            rss = np.sum(residuals**2)
            self.odict["data"]["clearance_data"]["curve_fit"][
                "sum_of_squared_residuals"
            ] = rss

            # curve_fit is used in discharge button to calculate discharge_activity
            if "discharge_activity" in self.odict["data"]["patient_discharge"]:
                self.odict["data"]["patient_discharge"].pop("discharge_activity")

            # curve_fit is used in discharge button to calculate calculated_discharge_dose_rate_xm/1m
            if (
                "calculated_discharge_dose_rate_xm"
                in self.odict["data"]["patient_discharge"]
            ):
                self.odict["data"]["patient_discharge"].pop(
                    "calculated_discharge_dose_rate_xm"
                )
            if (
                "calculated_discharge_dose_rate_1m"
                in self.odict["data"]["patient_discharge"]
            ):
                self.odict["data"]["patient_discharge"].pop(
                    "calculated_discharge_dose_rate_1m"
                )

            # curve_fit is used in discharge window to calculate recommended_datetime for discharge
            self.odict["data"]["patient_discharge"][
                "recommended_datetime"
            ] = collections.OrderedDict()

            # curve_fit is used to calculate restrictions
            if "restriction" in self.odict["data"]["restrictions"]:
                if (
                    "restriction_period"
                    in self.odict["data"]["restrictions"]["restriction"][0]
                ):  # min 2 restrictions
                    for child1 in self.odict["data"]["restrictions"]["restriction"]:
                        child1.pop("dose_constraint_corrected")
                        child1.pop("restriction_period")
                        child1.pop("dose")
                        child1.pop("datetime_end")

            self.odict["data"]["reports_generated"] = "0"

            self.unsaved_data = True

            if int(self.odict["data"]["patient_finished"]):
                self.odict["data"]["patient_finished"] = "0"
                self.odict["data"]["patient_finished_by"] = "0"
                self.viewing_completed_patient_label.place_forget()

            filepath = self.filepath
            if filepath is not None:
                self.root.title(f"*{os.path.basename(filepath)} - {__program_name__}")

            button_view_residuals["state"] = "normal"

            self.update_buttons()

        def view_residuals(self):
            if "curve_fit" in self.odict["data"]["clearance_data"]:
                v_model = self.odict["data"]["clearance_data"]["curve_fit"]["model"]
                fit_parameters = [
                    float(a)
                    for a in self.odict["data"]["clearance_data"]["curve_fit"][
                        "fit_parameters"
                    ].values()
                ]

                # if there is a curve fit, there are at least 2 measurements
                hrs_elapsed = [
                    float(child1["hours_elapsed"])
                    for child1 in self.odict["data"]["clearance_data"]["measurement"]
                    if not int(child1["exclude"])
                ]
                dr_xm = [
                    float(child1["doserate_corrected"])
                    for child1 in self.odict["data"]["clearance_data"]["measurement"]
                    if not int(child1["exclude"])
                ]

                residuals = 0
                if v_model == "biexponential":
                    residuals = np.array(dr_xm) - func_biexp(
                        np.array(hrs_elapsed), *fit_parameters
                    )
                elif v_model == "exponential":
                    residuals = np.array(dr_xm) - func_exp(
                        np.array(hrs_elapsed), *fit_parameters
                    )

                plt.plot(
                    np.array(hrs_elapsed),
                    residuals,
                    "o",
                    markeredgecolor="k",
                    markerfacecolor=(86 / 255, 180 / 255, 233 / 255),
                )
                plt.xlim(left=0)
                plt.plot(
                    np.linspace(0.0, plt.gca().get_xlim()[1], num=1000),
                    np.zeros(1000),
                    "--",
                    color=(0, 114 / 255, 178 / 255),
                )
                plt.xlabel("Time from administration (h)")
                plt.ylabel("Residual (\N{GREEK SMALL LETTER MU}Sv/h)")
                plt.title(
                    "If the fit is good, the residuals will be normally distributed with a mean of zero.",
                    loc="right",
                    fontsize=8,
                    pad=20,
                )
                plt.grid(alpha=0.2)
                plt.tight_layout()
                plt.show()

        def update_window_measure(self):
            # Retrieve values that the user entered previously
            if "detector_calibration_factor" in self.odict["data"]["clearance_data"]:
                detector_calibration_factor = self.odict["data"]["clearance_data"][
                    "detector_calibration_factor"
                ]
            else:
                detector_calibration_factor = None

            if "measurement_distance" in self.odict["data"]["clearance_data"]:
                measurement_distance = float(
                    self.odict["data"]["clearance_data"]["measurement_distance"]
                )
                if measurement_distance.is_integer():
                    measurement_distance = int(measurement_distance)
            else:
                measurement_distance = None

            if (
                "administration_datetime"
                in self.odict["data"]["administration_details"]
            ):
                admin_datetime = str2datetime(
                    self.odict["data"]["administration_details"][
                        "administration_datetime"
                    ]
                )
            else:
                admin_datetime = None

            if "curve_fit" in self.odict["data"]["clearance_data"]:
                model = self.odict["data"]["clearance_data"]["curve_fit"]["model"]
            else:
                model = None

            prev_datetime_list = []
            prev_doserate_list = []
            prev_exclude_list = []
            if "measurement" in self.odict["data"]["clearance_data"]:
                if isinstance(
                    self.odict["data"]["clearance_data"]["measurement"], list
                ):
                    for child in self.odict["data"]["clearance_data"]["measurement"]:
                        prev_datetime_list.append(child["datetime"])
                        prev_doserate_list.append(child["doserate"])
                        prev_exclude_list.append(child["exclude"])
                else:
                    prev_datetime_list.append(
                        self.odict["data"]["clearance_data"]["measurement"]["datetime"]
                    )
                    prev_doserate_list.append(
                        self.odict["data"]["clearance_data"]["measurement"]["doserate"]
                    )
                    prev_exclude_list.append(
                        self.odict["data"]["clearance_data"]["measurement"]["exclude"]
                    )

            for x in range(self.N_MEASUREMENTS_MAX):
                if x < len(prev_doserate_list):
                    prev_datetime = str2datetime(prev_datetime_list[x])
                    prev_doserate = prev_doserate_list[x]
                    prev_exclude = prev_exclude_list[x]
                else:
                    prev_datetime = None
                    prev_doserate = None
                    prev_exclude = None

                day[x].delete(0, tk.END)
                month[x].delete(0, tk.END)
                year[x].delete(0, tk.END)
                hour[x].delete(0, tk.END)
                minute[x].delete(0, tk.END)
                if prev_datetime is not None:
                    day[x].insert(0, prev_datetime.day)
                    month[x].insert(0, prev_datetime.month)
                    year[x].insert(0, str(prev_datetime.year).zfill(4))
                    hour[x].insert(0, prev_datetime.hour)
                    minute[x].insert(0, prev_datetime.minute)
                elif admin_datetime is not None and x == 0:
                    day[x].insert(0, admin_datetime.day)
                    month[x].insert(0, admin_datetime.month)
                    year[x].insert(0, str(admin_datetime.year).zfill(4))
                elif admin_datetime is not None:
                    month[x].insert(0, admin_datetime.month)
                    year[x].insert(0, str(admin_datetime.year).zfill(4))

                doserate[x].delete(0, tk.END)
                if prev_doserate is not None:
                    doserate[x].insert(0, prev_doserate)

                if prev_exclude is not None:
                    chks[x].set(prev_exclude)

            measurement_distance_entry.delete(0, tk.END)
            if measurement_distance is not None:
                measurement_distance_entry.insert(0, measurement_distance)
            else:
                measurement_distance_entry.insert(
                    0, self.init_vals["measurement_distance"]
                )

            detector_calibration_factor_entry.delete(0, tk.END)
            if detector_calibration_factor is not None:
                detector_calibration_factor_entry.insert(0, detector_calibration_factor)
            else:
                detector_calibration_factor_entry.insert(
                    0, self.init_vals["detector_calibration_factor"]
                )

            if "curve_fit" not in self.odict["data"]["clearance_data"]:
                button_view_residuals["state"] = "disabled"

            if model is not None:
                if model == "biexponential":
                    var1.set(2)
                elif model == "exponential":
                    var1.set(1)
            else:
                if self.init_vals["curve_fit_model"] == "biexponential":
                    var1.set(2)
                elif self.init_vals["curve_fit_model"] == "exponential":
                    var1.set(1)

            plot_clearance(self)

        window = tk.Toplevel()
        window_height = 24 * self.N_MEASUREMENTS_MAX + 300
        if window_height < 570:
            window_height = 570
        window.geometry(f"940x{int(window_height)}")
        window.title("Measured Clearance Data")
        if WINDOWS_OS:
            window.resizable(width=False, height=False)
            window.iconbitmap(self.SOFTWARE_ICON)

        tk.Label(window, text="Dose Rate Measurements", font="Arial 10 bold").grid(
            row=0, column=0, sticky="s"
        )

        frame4 = tk.Frame(window)

        tk.Label(frame4, text="Date").grid(row=0, columnspan=3, sticky="s", padx=(5, 0))
        tk.Label(frame4, text="DD").grid(row=1, sticky="s", padx=(5, 0))
        tk.Label(frame4, text="MM").grid(row=1, column=1, sticky="s")
        tk.Label(frame4, text="YYYY").grid(row=1, column=2, sticky="s")

        tk.Label(frame4, text="Time").grid(
            row=0, column=3, columnspan=2, padx=(20, 0), sticky="s"
        )
        tk.Label(frame4, text="hh").grid(row=1, column=3, padx=(20, 0), sticky="s")
        tk.Label(frame4, text="mm").grid(row=1, column=4, sticky="s")

        tk.Label(
            frame4, text="Displayed dose\nrate (\N{GREEK SMALL LETTER MU}Sv/h)"
        ).grid(row=0, column=5, rowspan=2, sticky="s")

        tk.Label(frame4, text="Exclude?").grid(row=0, column=6, rowspan=2, sticky="s")

        day = []
        month = []
        year = []
        hour = []
        minute = []
        doserate = []
        chks = []
        for x in range(self.N_MEASUREMENTS_MAX):
            day.append(tk.Entry(frame4, width=3))
            day[x].grid(row=x + 2, padx=(5, 0))
            month.append(tk.Entry(frame4, width=3))
            month[x].grid(row=x + 2, column=1)
            year.append(tk.Entry(frame4, width=5))
            year[x].grid(row=x + 2, column=2)
            hour.append(tk.Entry(frame4, width=3))
            hour[x].grid(row=x + 2, column=3, padx=(20, 0))
            minute.append(tk.Entry(frame4, width=3))
            minute[x].grid(row=x + 2, column=4)

            doserate.append(tk.Entry(frame4, width=5))
            doserate[x].grid(row=x + 2, column=5)

            chks.append(tk.IntVar())
            tk.Checkbutton(frame4, variable=chks[-1]).grid(row=x + 2, column=6)

        # This is to let the user press "Enter" to go down to next row in the same column
        def go_to_next_entry(event, entry_list, this_index):
            next_index = (this_index + 6) % len(
                entry_list
            )  # needs to move 6 entries across to go down 1
            entry_list[next_index].focus_set()

        entries = [
            child for child in frame4.winfo_children() if isinstance(child, tk.Entry)
        ]
        for idx, entry in enumerate(entries):
            entry.bind("<Return>", lambda e, idx=idx: go_to_next_entry(e, entries, idx))

        frame4.grid(row=1, column=0)

        frame1 = tk.Frame(window)
        tk.Label(frame1, text="Measurement distance (m)").grid(
            row=0, column=0, sticky="e"
        )
        tk.Label(
            frame1,
            text="Detector calibration factor\n(true/displayed)",
            justify=tk.RIGHT,
        ).grid(row=1, column=0, sticky="e")

        measurement_distance_entry = tk.Entry(frame1, width=5)
        measurement_distance_entry.grid(row=0, column=1, padx=(5, 0), pady=(5, 0))

        detector_calibration_factor_entry = tk.Entry(frame1, width=5)

        detector_calibration_factor_entry.grid(row=1, column=1, padx=(5, 0))
        frame1.grid(row=2, column=0, pady=(10, 0))

        frame5 = tk.Frame(window)
        button_view_residuals = tk.Button(
            frame5, text="View residuals", command=lambda: view_residuals(self)
        )
        button_view_residuals.grid(row=0, sticky="e")

        fig = Figure(figsize=(6.4, 4.8), dpi=100)
        plot_dr = fig.add_subplot(111)
        canvas = FigureCanvasTkAgg(fig, master=frame5)
        canvas.get_tk_widget().grid(row=1)
        toolbarFrame = tk.Frame(frame5)
        toolbarFrame.grid(row=2)
        NavigationToolbar2Tk(canvas, toolbarFrame)

        frame5.grid(row=0, column=1, rowspan=5)

        frame2 = tk.LabelFrame(window, text="Curve Fit", font="Arial 10 bold")
        var1 = tk.IntVar()

        tk.Radiobutton(frame2, text="Exponential", variable=var1, value=1).grid(
            sticky="w"
        )
        tk.Radiobutton(frame2, text="Biexponential", variable=var1, value=2).grid(
            row=1, sticky="w"
        )
        frame2.grid(row=3, pady=(10, 0))

        frame3 = tk.Frame(window)
        tk.Button(
            frame3,
            text="Submit and Plot",
            command=lambda: [
                retrieve_clearance_data(self),
                retrieve_fit_choice(self),
                plot_clearance(self),
            ],
        ).pack(side="left")

        tk.Button(frame3, text="OK", command=window.withdraw).pack(
            side="left", padx=(20, 0)
        )

        frame3.grid(row=4, pady=(20, 0))

        update_window_measure(self)

    # Patient Discharge
    def patient_discharge(self):
        from glowgreen import Clearance_1m

        def compute_recommended_discharge(self):
            model = self.odict["data"]["clearance_data"]["curve_fit"]["model"]
            meaningful_parameters = [
                float(a)
                for a in self.odict["data"]["clearance_data"]["curve_fit"][
                    "meaningful_parameters"
                ].values()
            ]
            if self.therapy_options_df.loc[
                self.odict["data"]["patient_details"]["type_therapy"],
                "generic_clearance",
            ]:
                measurement_distance = 1.0
            else:
                measurement_distance = float(
                    self.odict["data"]["clearance_data"]["measurement_distance"]
                )
            cfit = Clearance_1m(model, meaningful_parameters, measurement_distance)
            dr_1m_init = cfit.model_params[0]
            hrs = None
            method = None
            if dr_1m_init > 25.0:
                hrs = cfit.get_timedelta(25.0)
                method = "25 uSv/h at 1 m"

            radionuclide = self.therapy_options_df.loc[
                self.odict["data"]["patient_details"]["type_therapy"], "radionuclide"
            ]
            activity_limit_for_discharge = self.radionuclide_options_df.loc[
                radionuclide, "activity_outpatient"
            ]
            if not np.isnan(activity_limit_for_discharge):
                a0 = float(
                    self.odict["data"]["administration_details"][
                        "administered_activity"
                    ]
                )
                hrs_new = None
                if a0 > activity_limit_for_discharge:
                    hrs_new = cfit.get_timedelta(activity_limit_for_discharge, init=a0)
                if activity_limit_for_discharge.is_integer():
                    activity_limit_for_discharge = int(activity_limit_for_discharge)
                    activity_limit_for_discharge_str = f"{activity_limit_for_discharge}"
                else:
                    activity_limit_for_discharge_str = (
                        f"{activity_limit_for_discharge:.1f}"
                    )

                if hrs_new is not None:
                    if hrs is None:
                        hrs = hrs_new
                        method = f"{activity_limit_for_discharge_str} MBq retained"
                    elif hrs_new > hrs:
                        hrs = hrs_new
                        method = f"{activity_limit_for_discharge_str} MBq retained"
            if hrs is None:
                hrs = 0.0
                method = "immediately following administration"

            c_admin_datetime = str2datetime(
                self.odict["data"]["administration_details"]["administration_datetime"]
            )
            add_mins = 1 if hrs != 0 else 0
            recommended_discharge_datetime = c_admin_datetime + timedelta(
                hours=hrs, minutes=add_mins
            )

            self.odict["data"]["patient_discharge"]["recommended_datetime"][
                "datetime"
            ] = datetime2str(recommended_discharge_datetime)
            self.odict["data"]["patient_discharge"]["recommended_datetime"][
                "based_on"
            ] = method

        def set_patient_discharge(self):
            if len(discharge_year_entry.get()) != 4:
                display_discharge_datetime(self)
                messagebox.showerror(
                    "Error", "Please enter a 4-digit year", parent=window
                )
                return
            try:
                s_discharge_datetime = datetime(
                    year=int(discharge_year_entry.get()),
                    month=int(discharge_month_entry.get()),
                    day=int(discharge_day_entry.get()),
                    hour=int(discharge_hour_entry.get()),
                    minute=int(discharge_minute_entry.get()),
                )
            except ValueError:
                display_discharge_datetime(self)
                messagebox.showerror("Error", "Bad discharge date/time", parent=window)
                return
            if s_discharge_datetime < str2datetime(
                self.odict["data"]["administration_details"]["administration_datetime"]
            ):
                display_discharge_datetime(self)
                messagebox.showerror(
                    "Error", "Discharge before administration", parent=window
                )
                return

            self.odict["data"]["patient_discharge"]["actual_datetime"] = datetime2str(
                s_discharge_datetime
            )

            compute_discharge_activity(self)
            compute_discharge_dose_rate_in_curve_fit(self)

            set_discharge_activity_strings(self)
            set_discharge_dose_rate_strings(self)

            button_ok.grid(row=7, columnspan=7, pady=(10, 0))

            self.odict["data"]["reports_generated"] = "0"

            self.unsaved_data = True

            if int(self.odict["data"]["patient_finished"]):
                self.odict["data"]["patient_finished"] = "0"
                self.odict["data"]["patient_finished_by"] = "0"
                self.viewing_completed_patient_label.place_forget()

            filepath = self.filepath
            if filepath is not None:
                self.root.title(f"*{os.path.basename(filepath)} - {__program_name__}")

            self.update_buttons()

        def compute_discharge_activity(self):
            c_admin_datetime = str2datetime(
                self.odict["data"]["administration_details"]["administration_datetime"]
            )
            c_discharge_datetime = str2datetime(
                self.odict["data"]["patient_discharge"]["actual_datetime"]
            )

            model = self.odict["data"]["clearance_data"]["curve_fit"]["model"]
            meaningful_parameters = [
                float(a)
                for a in self.odict["data"]["clearance_data"]["curve_fit"][
                    "meaningful_parameters"
                ].values()
            ]
            a0 = float(
                self.odict["data"]["administration_details"]["administered_activity"]
            )

            discharge_activity = Gui.discharge_activity(
                c_admin_datetime, c_discharge_datetime, model, meaningful_parameters, a0
            )

            self.odict["data"]["patient_discharge"][
                "discharge_activity"
            ] = discharge_activity

        def compute_discharge_dose_rate_in_curve_fit(self):
            c_admin_datetime = str2datetime(
                self.odict["data"]["administration_details"]["administration_datetime"]
            )
            c_discharge_datetime = str2datetime(
                self.odict["data"]["patient_discharge"]["actual_datetime"]
            )
            model = self.odict["data"]["clearance_data"]["curve_fit"]["model"]
            meaningful_parameters = [
                float(a)
                for a in self.odict["data"]["clearance_data"]["curve_fit"][
                    "meaningful_parameters"
                ].values()
            ]
            if self.therapy_options_df.loc[
                self.odict["data"]["patient_details"]["type_therapy"],
                "generic_clearance",
            ]:
                measurement_distance = 1.0
            else:
                measurement_distance = float(
                    self.odict["data"]["clearance_data"]["measurement_distance"]
                )

            (
                discharge_dose_rate_1m,
                discharge_dose_rate_xm,
            ) = Gui.discharge_dose_rate(
                c_admin_datetime,
                c_discharge_datetime,
                model,
                meaningful_parameters,
                measurement_distance,
            )

            self.odict["data"]["patient_discharge"][
                "calculated_discharge_dose_rate_1m"
            ] = discharge_dose_rate_1m
            self.odict["data"]["patient_discharge"][
                "calculated_discharge_dose_rate_xm"
            ] = discharge_dose_rate_xm

        def set_discharge_activity_strings(self):
            discharge_activity = float(
                self.odict["data"]["patient_discharge"]["discharge_activity"]
            )
            discharge_activity_str.set(
                f"Activity retained at\ndischarge: {discharge_activity:.0f} MBq"
            )
            discharge_activity_label.grid(row=5, columnspan=5)

            radionuclide = self.therapy_options_df.loc[
                self.odict["data"]["patient_details"]["type_therapy"], "radionuclide"
            ]
            activity_limit_for_discharge = self.radionuclide_options_df.loc[
                radionuclide, "activity_outpatient"
            ]
            if not np.isnan(activity_limit_for_discharge):
                if discharge_activity > activity_limit_for_discharge:
                    if activity_limit_for_discharge.is_integer():
                        activity_limit_for_discharge = int(activity_limit_for_discharge)
                        activity_limit_for_discharge_str = (
                            f"{activity_limit_for_discharge}"
                        )
                    else:
                        activity_limit_for_discharge_str = (
                            f"{activity_limit_for_discharge:.1f}"
                        )

                    discharge_activity_warning_str.set(
                        f"WARNING:\n>{activity_limit_for_discharge_str} MBq"
                    )
                    discharge_activity_warning_label["fg"] = "red"
                else:
                    discharge_activity_warning_str.set("OK")
                    discharge_activity_warning_label["fg"] = "black"

        def set_discharge_dose_rate_strings(self):
            discharge_dose_rate_1m = float(
                self.odict["data"]["patient_discharge"][
                    "calculated_discharge_dose_rate_1m"
                ]
            )
            discharge_dose_rate_1m_str.set(
                "Calculated discharge dose\nrate at 1 m: {:.1f} \N{GREEK SMALL LETTER MU}Sv/h".format(
                    discharge_dose_rate_1m
                )
            )
            discharge_dose_rate_1m_label.grid(row=6, columnspan=5)
            if discharge_dose_rate_1m > 25.0:
                discharge_dose_rate_1m_warning_str.set(
                    "WARNING:\n>25 \N{GREEK SMALL LETTER MU}Sv/h"
                )
                discharge_dose_rate_1m_warning_label["fg"] = "red"
            else:
                discharge_dose_rate_1m_warning_str.set("OK")
                discharge_dose_rate_1m_warning_label["fg"] = "black"

        def display_discharge_datetime(self):
            discharge_hour_entry.delete(0, tk.END)
            discharge_minute_entry.delete(0, tk.END)
            discharge_day_entry.delete(0, tk.END)
            discharge_month_entry.delete(0, tk.END)
            discharge_year_entry.delete(0, tk.END)
            if "actual_datetime" in self.odict["data"]["patient_discharge"]:
                discharge_datetime = str2datetime(
                    self.odict["data"]["patient_discharge"]["actual_datetime"]
                )
                discharge_hour_entry.insert(0, discharge_datetime.hour)
                discharge_minute_entry.insert(0, discharge_datetime.minute)
                discharge_day_entry.insert(0, discharge_datetime.day)
                discharge_month_entry.insert(0, discharge_datetime.month)
                discharge_year_entry.insert(0, str(discharge_datetime.year).zfill(4))

        window = tk.Toplevel()
        window.geometry("273x295")
        window.title("Patient Discharge")
        if WINDOWS_OS:
            window.resizable(width=False, height=False)
            window.iconbitmap(self.SOFTWARE_ICON)

        if not self.odict["data"]["patient_discharge"]["recommended_datetime"]:
            compute_recommended_discharge(self)
        recommended_datetime_str = (
            str2datetime(
                self.odict["data"]["patient_discharge"]["recommended_datetime"][
                    "datetime"
                ]
            )
            .strftime("%d %b %Y, %I:%M %p")
            .lstrip("0")
            .replace(", 0", ", ")
        )

        frame = tk.Frame(window)
        tk.Label(frame, text=f"Recommended discharge: {recommended_datetime_str}").grid(
            columnspan=7
        )

        tk.Label(frame, text="Actual Discharge", font="Arial 10 bold").grid(
            row=1, columnspan=7, pady=(10, 5)
        )
        tk.Label(frame, text="Time").grid(row=2, sticky="W")
        discharge_hour_entry = tk.Entry(frame, width=5)
        discharge_hour_entry.grid(row=2, column=1)
        tk.Label(frame, text="(hh)").grid(row=2, column=2, padx=(0, 5))
        discharge_minute_entry = tk.Entry(frame, width=3)
        discharge_minute_entry.grid(row=2, column=3)
        tk.Label(frame, text="(mm)").grid(row=2, column=4, padx=(0, 5))

        tk.Label(frame, text="Date").grid(row=3, sticky="W")
        discharge_day_entry = tk.Entry(frame, width=5)
        discharge_day_entry.grid(row=3, column=1)
        tk.Label(frame, text="(DD)").grid(row=3, column=2, padx=(0, 5))
        discharge_month_entry = tk.Entry(frame, width=3)
        discharge_month_entry.grid(row=3, column=3)
        tk.Label(frame, text="(MM)").grid(row=3, column=4, padx=(0, 5))
        discharge_year_entry = tk.Entry(frame, width=5)
        discharge_year_entry.grid(row=3, column=5)
        tk.Label(frame, text="(YYYY)").grid(row=3, column=6, padx=(0, 5))

        if "actual_datetime" in self.odict["data"]["patient_discharge"]:
            discharge_datetime = str2datetime(
                self.odict["data"]["patient_discharge"]["actual_datetime"]
            )
            discharge_hour_entry.insert(0, discharge_datetime.hour)
            discharge_minute_entry.insert(0, discharge_datetime.minute)
            discharge_day_entry.insert(0, discharge_datetime.day)
            discharge_month_entry.insert(0, discharge_datetime.month)
            discharge_year_entry.insert(0, str(discharge_datetime.year).zfill(4))
        else:
            admin_datetime = str2datetime(
                self.odict["data"]["administration_details"]["administration_datetime"]
            )
            discharge_month_entry.insert(0, admin_datetime.month)
            discharge_year_entry.insert(0, str(admin_datetime.year).zfill(4))

        tk.Button(
            frame, text="Submit", command=lambda: set_patient_discharge(self)
        ).grid(row=4, columnspan=7, pady=(10, 20))

        discharge_activity_str = tk.StringVar()
        discharge_activity_label = tk.Label(
            frame, textvariable=discharge_activity_str, justify=tk.RIGHT
        )
        discharge_activity_label.grid(sticky="E")

        discharge_activity_warning_str = tk.StringVar()
        discharge_activity_warning_str.set("")
        discharge_activity_warning_label = tk.Label(
            frame, textvariable=discharge_activity_warning_str
        )
        discharge_activity_warning_label.grid(row=5, column=5, columnspan=2)

        discharge_dose_rate_1m_str = tk.StringVar()
        discharge_dose_rate_1m_label = tk.Label(
            frame, textvariable=discharge_dose_rate_1m_str, justify=tk.RIGHT
        )
        discharge_dose_rate_1m_label.grid(sticky="E")

        discharge_dose_rate_1m_warning_str = tk.StringVar()
        discharge_dose_rate_1m_warning_str.set("")
        discharge_dose_rate_1m_warning_label = tk.Label(
            frame, textvariable=discharge_dose_rate_1m_warning_str
        )
        discharge_dose_rate_1m_warning_label.grid(row=6, column=5, columnspan=2)

        button_ok = tk.Button(frame, text="OK", command=window.withdraw)

        if "discharge_activity" in self.odict["data"]["patient_discharge"]:
            set_discharge_activity_strings(self)
            set_discharge_dose_rate_strings(self)
            button_ok.grid(row=7, columnspan=7, pady=(10, 0))

        frame.grid(pady=(10, 0), padx=(10, 0))

    # Restrictions
    def restrictions_window(self):
        window = tk.Toplevel()
        window.geometry("690x590")
        window.title("Restrictions")
        if WINDOWS_OS:
            window.resizable(width=False, height=False)
            window.iconbitmap(self.SOFTWARE_ICON)

        num_treatments_in_year = float(
            self.odict["data"]["patient_details"]["num_treatments_in_year"]
        )
        if num_treatments_in_year.is_integer():
            num_treatments_in_year = int(num_treatments_in_year)

        tk.Label(
            window,
            text=f"NB. You selected {num_treatments_in_year} treatment(s) in a year",
        ).grid(columnspan=4, sticky="e", pady=(0, 20))
        tk.Label(window, text="Type", font="Arial 10 bold").grid(row=1)
        tk.Label(window, text="End", font="Arial 10 bold").grid(row=1, column=1)
        tk.Label(window, text="Applicable?", font="Arial 10 bold").grid(row=1, column=2)

        restr = RestrictionsWindow(self)
        restr.make_display(self, window)

        tk.Button(
            window,
            text="Submit",
            command=lambda: restr.submit_restrictions(self, window),
        ).grid(row=restr.df.shape[0] + 2, columnspan=4, pady=(20, 0))

    # Comments
    def comments(self):
        def submit_comments(self):
            submitted_text = text.get("1.0", "end")
            submitted_text_lines = submitted_text[
                :-1
            ].splitlines()  # ignore the newline that gets added on the end of the text.get()
            submitted_text = "\n".join(submitted_text_lines)
            self.odict["data"]["additional_comments_to_patient"] = submitted_text
            window.withdraw()
            self.odict["data"]["reports_generated"] = "0"
            self.unsaved_data = True
            if int(self.odict["data"]["patient_finished"]):
                self.odict["data"]["patient_finished"] = "0"
                self.odict["data"]["patient_finished_by"] = "0"
                self.viewing_completed_patient_label.place_forget()

            filepath = self.filepath
            if filepath is not None:
                self.root.title(f"*{os.path.basename(filepath)} - {__program_name__}")

            self.update_buttons()

        def load_default_comments():
            text.delete("1.0", "end")
            text.insert("1.0", self.init_vals["patient_handout_comments"])

        window = tk.Toplevel()
        window.geometry("500x240")
        window.title("Comments")
        if WINDOWS_OS:
            window.resizable(width=False, height=False)
            window.iconbitmap(self.SOFTWARE_ICON)

        text = tk.Text(window, width=54, height=9, wrap=tk.WORD)
        tk.Button(window, text="Load default", command=load_default_comments).pack(
            padx=(360, 0), pady=(10, 5)
        )
        text.pack()

        tk.Button(window, text="Submit", command=lambda: submit_comments(self)).pack(
            pady=(10, 0)
        )

        if self.odict["data"]["additional_comments_to_patient"] != "0":
            if self.odict["data"]["additional_comments_to_patient"]:
                text.insert("1.0", self.odict["data"]["additional_comments_to_patient"])
        else:
            load_default_comments()

    @staticmethod
    def split_line_report(name_str, max_line_len):
        """Break up a long restriction name into multiple lines
        for writing in table in report.
        """
        if len(name_str) > max_line_len:
            out_str = ""
            leftover_str = name_str
            while len(leftover_str) > max_line_len:
                word_list = leftover_str.split()
                word_length = [len(x) for x in word_list]
                if any(x + 1 > max_line_len for x in word_length):
                    exit("A word is too long in a restriction name")
                cumulative_length = np.array(
                    [
                        np.sum(word_length[:j]) + j
                        for j in range(1, 1 + len(word_length))
                    ]
                )  # includes a space after each word

                jj = np.argmax(cumulative_length > max_line_len)

                out_str += " ".join(word_list[:jj]) + "\n"
                leftover_str = " " * 3 + " ".join(
                    word_list[jj:]
                )  # indent after first line
            out_str += leftover_str
        else:
            out_str = name_str
        return out_str

    # Generate Reports
    def generate_reports(self):
        from docx import Document
        from docx.shared import Inches, Pt

        def get_patient_handout(self):
            site = self.odict["data"]["patient_details"]["site"]
            if site in self.site_options_df.index.to_list():
                addressA = self.site_options_df.loc[site, "address_line1"]
                addressB = self.site_options_df.loc[site, "address_line2"]
                phone = self.site_options_df.loc[site, "phone"]
            else:
                addressA = ""
                addressB = ""
                phone = ""
            email = self.email
            url = self.url
            last_name, first_name = self.odict["data"]["patient_details"]["name"].split(
                "^"
            )
            last_name = last_name.upper()
            pid = self.odict["data"]["patient_details"]["id"]
            dob = str2date(self.odict["data"]["patient_details"]["dob"])
            dob = dob.strftime("%d %b, %Y")
            type_therapy = (
                self.odict["data"]["patient_details"]["type_therapy"]
                .split(",")[0]
                .split("(")[0]
            )

            admin_datetime = str2datetime(
                self.odict["data"]["administration_details"]["administration_datetime"]
            )
            admin_datetime_str = (
                admin_datetime.strftime("%I:%M %p on %d %b, %Y")
                .lstrip("0")
                .replace(" 0", " ")
            )
            a0 = self.odict["data"]["administration_details"]["administered_activity"]
            a0 = f"{float(a0):.0f}"

            inpatient = self.therapy_options_df.loc[
                self.odict["data"]["patient_details"]["type_therapy"], "inpatient"
            ]
            discharge_datetime_str = None
            if inpatient:
                discharge_datetime = str2datetime(
                    self.odict["data"]["patient_discharge"]["actual_datetime"]
                )
                discharge_datetime_str = (
                    discharge_datetime.strftime("%I:%M %p on %d %b, %Y")
                    .lstrip("0")
                    .replace(" 0", " ")
                )

            df = pd.DataFrame.from_dict(
                self.odict["data"]["restrictions"]["restriction"]
            )
            df["restriction_period"] = df["restriction_period"].astype(float)
            df.sort_values(by="restriction_period", inplace=True)

            generic = self.therapy_options_df.loc[
                self.odict["data"]["patient_details"]["type_therapy"],
                "generic_clearance",
            ]
            generic_str = (
                ""
                if generic
                else " and actual measurements of the radiation coming from your body"
            )
            inpatient_str = (
                "are discharged from hospital" if inpatient else "leave the department"
            )

            instructions = (
                "The table below contains details of any restrictions you have on your prolonged close contact with others.  "
                "You must follow these instructions to ensure that the radiation exposure to your family and friends "
                "will be well below the legal limit.  "
                "Detailed explanation of the types of restrictions can be found in the information sheet "
                "that you were provided by your Nuclear Medicine Doctor.  "
                f"Please contact the Nuclear Medicine Department at {site} if you have any questions."
            )

            additional_comments = self.odict["data"]["additional_comments_to_patient"]
            if additional_comments == "":
                additional_comments = None

            document = Document()
            font = document.styles["Normal"].font
            font.name = "Arial"
            section = document.sections[0]

            # header and footer
            header = section.header
            p = header.paragraphs[0]
            r = p.add_run()
            if os.path.isfile(self.report_logo):
                r.add_picture(self.report_logo, width=Inches(2))
            else:
                messagebox.showinfo(
                    "Information",
                    "Organisation logo image file not found and not included in the patient handout.",
                )

            footer = section.footer
            p = footer.paragraphs[0]
            r = p.add_run()
            r.text = "For more information"
            s = p.add_run()
            s.font.bold = True
            s.text = f"\n{site}\n{addressA}\n{addressB}\n{phone}\n{email}\n{url}\t\tMade using {__program_name__} version {__version__} ({__release_date__})"
            s.font.size = Pt(8)

            section.footer_distance = Inches(0.1)

            # title
            document.add_heading(
                "Radiation Protection Instructions following Radionuclide Therapy", 0
            )

            # patient and treatment details
            table = document.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Patient details"
            hdr_cells[1].text = "Treatment details"

            # set the font size for the heading row of the table
            for hdr_cell in hdr_cells:
                r = hdr_cell.paragraphs[0].runs
                r[0].font.size = Pt(10)
                r[0].font.bold = True

            # fill the table
            row_cells = table.add_row().cells
            row_cells[0].text = (
                f"\u2022 First name: {first_name}\n\u2022 Last name: {last_name}\n\u2022 ID number: {pid}\n\u2022 Date of birth: {dob}"
            )
            row_cells[1].text = (
                f"\u2022 Type of therapy: {type_therapy}\n\u2022 Site: {site}\n\u2022 Date/time of administration: {admin_datetime_str}\n\u2022 Activity administered: {a0} MBq"
            )
            if discharge_datetime_str is not None:
                row_cells[
                    1
                ].text += (
                    f"\n\u2022 Date/time of patient discharge: {discharge_datetime_str}"
                )

            # set the font size for the table
            for row_cell in row_cells:
                r = row_cell.paragraphs[0].runs
                r[0].font.size = Pt(10)

            # instructions
            document.add_heading("Close Contact Restrictions", level=2)

            p = document.add_paragraph()
            r = p.add_run()
            r.text = instructions

            r = p.runs
            for r in r:
                r.font.size = Pt(12)

            table = document.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Restriction"
            hdr_cells[1].text = "Applies until"
            hdr_cells[2].text = "Expired?"

            # set the font size for the heading row of the table
            for hdr_cell in hdr_cells:
                r = hdr_cell.paragraphs[0].runs
                r[0].font.size = Pt(12)
                r[0].font.bold = True

            # fill the table
            col1 = ""
            col2 = ""
            col3 = ""
            n_rows = len(df[df["applies"].isin([1, "1"])])
            if n_rows != 0:
                ii = 1
                for index, row in df.iterrows():
                    if int(row["applies"]):
                        name_str = "\N{BULLET} {}".format(row["name"])
                        MAX_NAME_LEN = 48  # max number of chars before we break a line... we must break the line before Word does... will depend on page margins, font size, etc
                        name_str_corrected = Gui.split_line_report(
                            name_str, MAX_NAME_LEN
                        )
                        n_new_lines = name_str_corrected.count("\n")
                        col1 += name_str_corrected

                        datetime_end = str2datetime(row["datetime_end"])
                        # display end of restriction as ceil of datetime_end to hour resolution
                        datetime_end_display = (
                            (
                                datetime_end.replace(minute=0, second=0, microsecond=0)
                                + timedelta(hours=1)
                            )
                            if datetime_end.minute != 0
                            else datetime_end
                        )
                        # don't use platform-specific strftime format codes
                        # (without zero-padded is dash on Unix and hash on Windows)
                        col2 += (
                            datetime_end_display.strftime("%I %p on %d %b, %Y")
                            .lstrip("0")
                            .replace(" 0", " ")
                        )
                        col2 += "\n" * n_new_lines

                        if discharge_datetime_str is not None:
                            if datetime_end_display <= discharge_datetime:
                                col3 += "Yes"
                        col3 += "\n" * n_new_lines

                        if ii < n_rows:
                            col1 += "\n"
                            col2 += "\n"
                            col3 += "\n"
                        ii += 1
            else:
                col1 += "No restrictions."

            row_cells = table.add_row().cells
            row_cells[0].text = col1
            row_cells[1].text = col2
            row_cells[2].text = col3

            # set the font size for the table
            for row_cell in row_cells:
                r = row_cell.paragraphs[0].runs
                r[0].font.size = Pt(12)

            table.style = document.styles["Table Grid"]

            for cell in table.columns[2].cells:
                cell.width = Inches(0.7)
            for cell in table.columns[1].cells:
                cell.width = Inches(2.1)
            for cell in table.columns[0].cells:
                cell.width = Inches(3.8)

            # additional comments
            if additional_comments is not None:
                document.add_heading("Additional Comments and Instructions", level=2)

                p = document.add_paragraph(additional_comments)

                # set the font size
                r = p.runs
                for r in r:
                    r.font.size = Pt(12)

            # sign off
            p = document.add_paragraph("\nMedical Physicist/Doctor: ")
            r = p.add_run()
            r.text = "\t\t\t\t\t\t"
            r.underline = True
            r = p.add_run()
            r.text = "\n\nSignature: "
            r = p.add_run()
            r.text = "\t\t\t\t"
            r.underline = True
            r = p.add_run()
            r.text = "    Date: "
            r = p.add_run()
            r.text = "\t\t\t"
            r.underline = True

            # set the font size
            r = p.runs
            for r in r:
                r.font.size = Pt(12)

            # set the page margins for the document
            sections = document.sections
            for section in sections:
                section.right_margin = Inches(1)

            return document

        def get_summary_report(self):
            site = self.odict["data"]["patient_details"]["site"]
            if site in self.site_options_df.index.to_list():
                addressA = self.site_options_df.loc[site, "address_line1"]
                addressB = self.site_options_df.loc[site, "address_line2"]
                phone = self.site_options_df.loc[site, "phone"]
            else:
                addressA = ""
                addressB = ""
                phone = ""
            email = self.email
            url = self.url
            last_name, first_name = self.odict["data"]["patient_details"]["name"].split(
                "^"
            )
            last_name = last_name.upper()
            pid = self.odict["data"]["patient_details"]["id"]
            dob = str2date(self.odict["data"]["patient_details"]["dob"])
            dob = dob.strftime("%d %b, %Y")
            pregnancy_excluded = int(
                self.odict["data"]["patient_details"]["pregnancy_excluded"]
            )
            breastfeeding_excluded = int(
                self.odict["data"]["patient_details"]["breastfeeding_excluded"]
            )
            hygiene_status = int(self.odict["data"]["patient_details"]["hygiene"])
            type_therapy = (
                self.odict["data"]["patient_details"]["type_therapy"]
                .split(",")[0]
                .split("(")[0]
            )

            admin_datetime = str2datetime(
                self.odict["data"]["administration_details"]["administration_datetime"]
            )
            admin_datetime_str = (
                admin_datetime.strftime("%I:%M %p on %d %b, %Y")
                .lstrip("0")
                .replace(" 0", " ")
            )
            a0_str = "{:.0f}".format(
                float(
                    self.odict["data"]["administration_details"][
                        "administered_activity"
                    ]
                )
            )

            discharge_datetime_str = None
            recommended_discharge_datetime_str = None
            recommended_discharge_based_on = None
            activity_limit_for_discharge = None
            discharge_activity_str = None
            calculated_discharge_dose_rate_1m_str = None
            inpatient = self.therapy_options_df.loc[
                self.odict["data"]["patient_details"]["type_therapy"], "inpatient"
            ]
            if inpatient:
                discharge_datetime = str2datetime(
                    self.odict["data"]["patient_discharge"]["actual_datetime"]
                )
                discharge_datetime_str = (
                    discharge_datetime.strftime("%I:%M %p on %d %b, %Y")
                    .lstrip("0")
                    .replace(" 0", " ")
                )
                recommended_discharge_datetime = str2datetime(
                    self.odict["data"]["patient_discharge"]["recommended_datetime"][
                        "datetime"
                    ]
                )
                recommended_discharge_datetime_str = (
                    recommended_discharge_datetime.strftime("%I:%M %p on %d %b, %Y")
                    .lstrip("0")
                    .replace(" 0", " ")
                )
                recommended_discharge_based_on = self.odict["data"][
                    "patient_discharge"
                ]["recommended_datetime"]["based_on"]
                if recommended_discharge_based_on == "25 uSv/h at 1 m":
                    recommended_discharge_based_on = (
                        "25 \N{GREEK SMALL LETTER MU}Sv/h at 1 m"
                    )
                rn = self.therapy_options_df.loc[
                    self.odict["data"]["patient_details"]["type_therapy"],
                    "radionuclide",
                ]
                activity_limit_for_discharge = self.radionuclide_options_df.loc[
                    rn, "activity_outpatient"
                ]
                if np.isnan(activity_limit_for_discharge):
                    activity_limit_for_discharge = None
                discharge_activity = float(
                    self.odict["data"]["patient_discharge"]["discharge_activity"]
                )
                discharge_activity_str = f"{discharge_activity:.0f}"
                calculated_discharge_dose_rate_1m = float(
                    self.odict["data"]["patient_discharge"][
                        "calculated_discharge_dose_rate_1m"
                    ]
                )
                calculated_discharge_dose_rate_1m_str = (
                    f"{calculated_discharge_dose_rate_1m:.1f}"
                )

            measurement_distance_str = None
            first_dose_rate_measurement_str = None
            last_dose_rate_measurement_str = None
            generic = self.therapy_options_df.loc[
                self.odict["data"]["patient_details"]["type_therapy"],
                "generic_clearance",
            ]
            if not generic:
                measurement_distance = float(
                    self.odict["data"]["clearance_data"]["measurement_distance"]
                )
                if measurement_distance.is_integer():
                    measurement_distance = int(measurement_distance)
                    measurement_distance_str = f"{measurement_distance}"
                else:
                    measurement_distance_str = f"{measurement_distance:.1f}"
                first_dose_rate_measurement = float(
                    self.odict["data"]["clearance_data"]["measurement"][0][
                        "doserate_corrected"
                    ]
                )
                first_dose_rate_measurement_str = f"{first_dose_rate_measurement:.1f}"
                last_dose_rate_measurement = float(
                    self.odict["data"]["clearance_data"]["measurement"][-1][
                        "doserate_corrected"
                    ]
                )
                last_dose_rate_measurement_str = f"{last_dose_rate_measurement:.1f}"

            df = pd.DataFrame.from_dict(
                self.odict["data"]["restrictions"]["restriction"]
            )
            df["restriction_period"] = df["restriction_period"].astype(float)
            df.sort_values(by="restriction_period", inplace=True)

            num_treatments_in_year = float(
                self.odict["data"]["patient_details"]["num_treatments_in_year"]
            )
            if num_treatments_in_year.is_integer():
                num_treatments_in_year = int(num_treatments_in_year)
                num_treatments_in_year_str = f"{num_treatments_in_year}"
            else:
                num_treatments_in_year_str = f"{num_treatments_in_year:.1f}"

            plural_str = "" if num_treatments_in_year == 1 else "s"
            inpatient_str = "upon discharge" if inpatient else "at home"
            if generic:
                generic_str = (
                    "The calculated close contact restrictions below are based on the administered activity "
                    "and a generic clearance rate."
                )
            else:
                generic_str = (
                    "The calculated close contact restrictions below are based on actual clearance and exposure rates, "
                    "derived from several measurements of dose rate from this patient over a period of time."
                )

            instructions_1 = f"{generic_str}  It was assumed the patient is receiving {num_treatments_in_year_str} treatment{plural_str} in a year."
            instructions_2 = (
                "The precautions and restrictions necessary to limit radiation exposure of other persons "
                "have been fully explained to this patient.  "
                "The patient has also been provided with written instructions "
                "which outline the procedures to be followed "
                "and the duration of time for which specific restrictions will apply."
            )

            document = Document()
            document.styles["Normal"].font.name = "Arial"
            section = document.sections[0]

            # header and footer
            header = section.header
            p = header.paragraphs[0]
            r = p.add_run()
            if os.path.isfile(self.report_logo):
                r.add_picture(self.report_logo, width=Inches(2))
            else:
                messagebox.showinfo(
                    "Information",
                    "Organisation logo image file not found and not included in the summary report.",
                )

            footer = section.footer
            p = footer.paragraphs[0]
            r = p.add_run()
            r.text = "For more information"
            s = p.add_run()
            s.font.bold = True
            s.text = f"\n{site}\n{addressA}\n{addressB}\n{phone}\n{email}\n{url}\t\tMade using {__program_name__} version {__version__} ({__release_date__})"
            s.font.size = Pt(8)

            section.footer_distance = Inches(0.1)

            # title
            document.add_heading(
                "Radiation Protection Measures following Radionuclide Therapy", 0
            )

            # patient and treatment details
            table = document.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Patient details"
            hdr_cells[1].text = "Treatment details"

            # set the font size for the heading row of the table
            for hdr_cell in hdr_cells:
                r = hdr_cell.paragraphs[0].runs
                r[0].font.size = Pt(8)
                r[0].font.bold = True

            # fill the table
            row_cells = table.add_row().cells
            row_cells[0].text = (
                f"\u2022 First name: {first_name}\n\u2022 Last name: {last_name}\n\u2022 ID number: {pid}\n\u2022 Date of birth: {dob}"
            )
            row_cells[1].text = (
                f"\u2022 Type of therapy: {type_therapy}\n\u2022 Site: {site}\n\u2022 Date/time of administration: {admin_datetime_str}\n\u2022 Activity administered: {a0_str} MBq"
            )
            if discharge_datetime_str is not None:
                row_cells[
                    1
                ].text += (
                    f"\n\u2022 Date/time of patient discharge: {discharge_datetime_str}"
                )

            if pregnancy_excluded:
                row_cells[0].text += "\n\u2022 Pregnancy excluded"
            if breastfeeding_excluded:
                row_cells[0].text += "\n\u2022 Breastfeeding excluded"
            if hygiene_status:
                row_cells[
                    0
                ].text += f"\n\u2022 The patient can be properly cared for (and good hygiene maintained) {inpatient_str}"

            if recommended_discharge_based_on is not None:
                row_cells[
                    1
                ].text += f"\n\u2022 Recommended date/time of discharge based on {recommended_discharge_based_on}: {recommended_discharge_datetime_str}"
                row_cells[
                    1
                ].text += f"\n\u2022 Calculated activity retained at time of discharge: {discharge_activity_str} MBq"
            if activity_limit_for_discharge is not None:
                if discharge_activity > activity_limit_for_discharge:
                    row_cells[
                        1
                    ].text += f"- exceeds the recommended limit of {activity_limit_for_discharge:.0f} MBq"
            if measurement_distance_str is not None:
                row_cells[
                    1
                ].text += "\n\u2022 First (initial) dose rate measurement at {} m: {} \N{GREEK SMALL LETTER MU}Sv/h".format(
                    measurement_distance_str, first_dose_rate_measurement_str
                )
                row_cells[
                    1
                ].text += "\n\u2022 Last (final) dose rate measurement at {} m: {} \N{GREEK SMALL LETTER MU}Sv/h".format(
                    measurement_distance_str, last_dose_rate_measurement_str
                )
            if calculated_discharge_dose_rate_1m_str is not None:
                row_cells[
                    1
                ].text += "\n\u2022 Calculated discharge dose rate at 1 m: {} \N{GREEK SMALL LETTER MU}Sv/h".format(
                    calculated_discharge_dose_rate_1m_str
                )
                if calculated_discharge_dose_rate_1m > 25.0:
                    row_cells[
                        1
                    ].text += " - exceeds the recommended limit of 25 \N{GREEK SMALL LETTER MU}Sv/h"

            # set the font size for the table
            for row_cell in row_cells:
                r = row_cell.paragraphs[0].runs
                r[0].font.size = Pt(8)

            # instructions
            document.add_heading("Restrictions", level=2)

            p = document.add_paragraph()
            r = p.add_run()
            r.text = instructions_1

            r = p.runs
            for r in r:
                r.font.size = Pt(10)

            # restrictions without formatting
            n_rows = len(df[df["applies"].isin([1, "1"])])
            if n_rows != 0:
                restrictions_plain = ""
                ii = 1
                for index, row in df.iterrows():
                    if int(row["applies"]):
                        restrictions_plain += row["name"] + ": "

                        datetime_end = str2datetime(row["datetime_end"])
                        # display end of restriction as ceil of datetime_end to hour resolution
                        datetime_end_display = (
                            (
                                datetime_end.replace(minute=0, second=0, microsecond=0)
                                + timedelta(hours=1)
                            )
                            if datetime_end.minute != 0
                            else datetime_end
                        )
                        # don't use platform-specific strftime format codes
                        # (without zero-padded is dash on Unix and hash on Windows)
                        restrictions_plain += (
                            datetime_end_display.strftime("%I %p on %d %b, %Y")
                            .lstrip("0")
                            .replace(" 0", " ")
                        )

                        if ii < n_rows:
                            restrictions_plain += "\n"
                        ii += 1
            else:
                restrictions_plain = "No restrictions."

            border_str = 80 * "_"
            restrictions_plain = f"{border_str}\n{restrictions_plain}\n{border_str}"

            p = document.add_paragraph()
            r = p.add_run()
            r.text = restrictions_plain

            r = p.runs
            for r in r:
                r.font.size = Pt(10)

            p = document.add_paragraph()
            r = p.add_run()
            r.text = instructions_2

            r = p.runs
            for r in r:
                r.font.size = Pt(10)

            # sign off
            p = document.add_paragraph("\n\nMedical Physicist/Doctor: ")
            r = p.add_run()
            r.text = "\t\t\t\t"
            r.underline = True
            r = p.add_run()
            r.text = "    Signature: "
            r = p.add_run()
            r.text = "\t\t"
            r.underline = True
            r = p.add_run()
            r.text = "    Date: "
            r = p.add_run()
            r.text = "\t"
            r.underline = True

            # set the font size in the sign-off
            r = p.runs
            for r in r:
                r.font.size = Pt(10)

            # set the page margins for the document
            sections = document.sections
            for section in sections:
                section.right_margin = Inches(1)

            return document

        last_name, first_name = self.odict["data"]["patient_details"]["name"].split("^")
        admin_datetime = str2datetime(
            self.odict["data"]["administration_details"]["administration_datetime"]
        )
        filename = (
            last_name.upper() + "_" + first_name + "_" + admin_datetime.strftime("%b%Y")
        )

        default_dir = self.init_vals["patient_handout_directory"]
        if not os.path.isdir(default_dir):
            default_dir = os.getcwd()
            messagebox.showinfo(
                "Information",
                "The default directory for patient handouts was not found.",
            )

        want_export_pdf = int(self.other_settings["export_pdf"])
        if want_export_pdf:
            word_avail = True
            try:
                word_available(WINDOWS_OS)
            except Exception as e:
                word_avail = False
                messagebox.showinfo(
                    "Information",
                    f"Unable to export docx files to PDF:\n    {e}\nTo stop seeing this message, uncheck the box in 'Settings \u2192 Other'.",
                )
        export_pdf = want_export_pdf and word_avail

        f = filedialog.asksaveasfilename(
            initialdir=(
                default_dir
                if self.previous_report_directory is None
                else self.previous_report_directory
            ),
            title="Save patient handout",
            initialfile="Patient_handout_" + filename + ".docx",
            filetypes=[("Microsoft Word Document (*.docx)", "*.docx")],
        )
        if f != "" and f != ():
            self.previous_report_directory = os.path.dirname(f)
            patient_handout = get_patient_handout(self)
            try:
                patient_handout.save(f)
            except Exception as e:
                messagebox.showerror(
                    "Error",
                    f"{e}",
                )
                # Better to raise than return so there is a more verbose
                # error message in CLI session.
                raise
            if export_pdf:
                base, _ = os.path.splitext(os.path.basename(f))
                initfile = base + ".pdf"
                fp = filedialog.asksaveasfilename(
                    initialdir=(
                        default_dir
                        if self.previous_report_directory is None
                        else self.previous_report_directory
                    ),
                    title="Save patient handout",
                    initialfile=initfile,
                    filetypes=[("PDF Files", "*.pdf")],
                )
                if fp != "" and fp != ():
                    try:
                        docx2pdf(f, fp)
                    except Exception as e:
                        messagebox.showwarning(
                            "Warning - Failed PDF Export",
                            f"Failed to export docx patient handout to PDF:\n    {e}",
                        )
                        # PDF export is not essential
                        # so don't raise or return

        default_dir = self.init_vals["summary_report_directory"]
        if not os.path.isdir(default_dir):
            default_dir = os.getcwd()
            messagebox.showinfo(
                "Information",
                "The default directory for summary reports was not found.",
            )

        f = filedialog.asksaveasfilename(
            initialdir=(
                default_dir
                if self.previous_report_directory is None
                else self.previous_report_directory
            ),
            title="Save summary report",
            initialfile="Summary_report_" + filename + ".docx",
            filetypes=[("Microsoft Word Document (*.docx)", "*.docx")],
        )
        if f != "" and f != ():
            self.previous_report_directory = os.path.dirname(f)
            summary_report = get_summary_report(self)
            try:
                summary_report.save(f)
            except Exception as e:
                messagebox.showerror(
                    "Error",
                    f"{e}",
                )
                # Better to raise than return so there is a more verbose
                # error message in CLI session.
                raise
            if export_pdf:
                base, _ = os.path.splitext(os.path.basename(f))
                initfile = base + ".pdf"
                fp = filedialog.asksaveasfilename(
                    initialdir=(
                        default_dir
                        if self.previous_report_directory is None
                        else self.previous_report_directory
                    ),
                    title="Save summary report",
                    initialfile=initfile,
                    filetypes=[("PDF Files", "*.pdf")],
                )
                if fp != "" and fp != ():
                    try:
                        docx2pdf(f, fp)
                    except Exception as e:
                        messagebox.showwarning(
                            "Warning - Failed PDF Export",
                            f"Failed to export docx summary report to PDF:\n    {e}",
                        )
                        # PDF export is not essential
                        # so don't raise or return

        if not int(self.odict["data"]["reports_generated"]):
            self.unsaved_data = True
        self.odict["data"]["reports_generated"] = "1"
        self.update_buttons()

    # Finish
    def finish_patient(self):
        filepath = self.filepath
        if filepath is None:
            last_name, first_name = self.odict["data"]["patient_details"]["name"].split(
                "^"
            )
            admin_datetime = str2datetime(
                self.odict["data"]["administration_details"]["administration_datetime"]
            )
            filename = (
                last_name.upper()
                + "_"
                + first_name
                + "_"
                + admin_datetime.strftime("%b%Y")
            )
        else:
            filename = os.path.basename(filepath)

        default_dir = self.init_vals["data_directory"]
        if not os.path.isdir(default_dir):
            default_dir = os.getcwd()
            messagebox.showinfo(
                "Information",
                "The default directory for XML file output was not found.",
            )

        filepath_new = filedialog.asksaveasfilename(
            initialdir=(
                default_dir
                if self.previous_data_directory is None
                else self.previous_data_directory
            ),
            title="Save XML file",
            initialfile=filename,
            filetypes=[("XML file (*.xml)", "*.xml")],
        )
        if filepath_new != "" and filepath_new != ():
            self.previous_data_directory = os.path.dirname(filepath_new)
            self.odict["data"]["patient_finished"] = "1"
            self.odict["data"]["patient_finished_by"] = getpass.getuser()
            self.save_odict_to_xml(filepath_new)
            self.new_patient()

    # Update the display and availability of buttons in the GUI Main Menu according to the data present and missing in self.odict
    def update_buttons(self):
        if self.action_required_patient():
            self.button_patient["bg"] = "light yellow"
        else:
            self.button_patient["bg"] = "white smoke"

        if self.action_forbidden_administration():
            self.button_administration["state"] = "disabled"
            self.button_administration["bg"] = "white smoke"
        else:
            self.button_administration["state"] = "normal"
            if self.action_required_administration():
                self.button_administration["bg"] = "light yellow"
            else:
                self.button_administration["bg"] = "white smoke"

        if self.action_forbidden_clearance():
            self.button_clearance["state"] = "disabled"
            self.button_clearance["bg"] = "white smoke"
        else:
            self.button_clearance["state"] = "normal"
            if self.action_required_clearance():
                self.button_clearance["bg"] = "light yellow"
            else:
                self.button_clearance["bg"] = "white smoke"

        if self.action_forbidden_discharge():
            self.button_discharge["state"] = "disabled"
            self.button_discharge["bg"] = "white smoke"
        else:
            self.button_discharge["state"] = "normal"
            if self.action_required_discharge():
                self.button_discharge["bg"] = "light yellow"
            else:
                self.button_discharge["bg"] = "white smoke"

        if self.action_forbidden_restrictions():
            self.button_restrictions["state"] = "disabled"
            self.button_restrictions["bg"] = "white smoke"
        else:
            self.button_restrictions["state"] = "normal"
            if self.action_required_restrictions():
                self.button_restrictions["bg"] = "light yellow"
            else:
                self.button_restrictions["bg"] = "white smoke"

        if self.action_forbidden_comments():
            self.button_comments["state"] = "disabled"
            self.button_comments["bg"] = "white smoke"
        else:
            self.button_comments["state"] = "normal"
            if self.action_required_comments():
                self.button_comments["bg"] = "light yellow"
            else:
                self.button_comments["bg"] = "white smoke"

        if self.action_forbidden_reports():
            self.button_reports["state"] = "disabled"
            self.button_reports["bg"] = "white smoke"
        else:
            self.button_reports["state"] = "normal"
            if not int(self.odict["data"]["reports_generated"]):
                self.button_reports["bg"] = "light yellow"
            else:
                self.button_reports["bg"] = "white smoke"

        if (not int(self.odict["data"]["reports_generated"])) or self.action_required():
            self.button_finish["state"] = "disabled"
            self.button_finish["bg"] = "white smoke"
        else:
            self.button_finish["state"] = "normal"
            if not int(self.odict["data"]["patient_finished"]):
                self.button_finish["bg"] = "light yellow"
            else:
                self.button_finish["bg"] = "white smoke"

    # "Action forbidden" meaning data is missing which must be present to open the window
    def action_forbidden_administration(self):
        return "type_therapy" not in self.odict["data"]["patient_details"]

    def action_forbidden_clearance(self):
        return (
            "administration_datetime"
            not in self.odict["data"]["administration_details"]
            or "administered_activity"
            not in self.odict["data"]["administration_details"]
            or "type_therapy" not in self.odict["data"]["patient_details"]
        )

    def action_forbidden_discharge(self):
        return (
            "type_therapy" not in self.odict["data"]["patient_details"]
            or "administration_datetime"
            not in self.odict["data"]["administration_details"]
            or "administered_activity"
            not in self.odict["data"]["administration_details"]
            or "curve_fit" not in self.odict["data"]["clearance_data"]
        )

    def action_forbidden_restrictions(self):
        return (
            "num_treatments_in_year" not in self.odict["data"]["patient_details"]
            or "administration_datetime"
            not in self.odict["data"]["administration_details"]
            or "curve_fit" not in self.odict["data"]["clearance_data"]
        )

    def action_forbidden_common_comments_reports(self):
        a = all(
            k in self.odict["data"]["patient_details"]
            for k in (
                "name",
                "id",
                "dob",
                "type_therapy",
                "site",
                "num_treatments_in_year",
                "pregnancy_excluded",
                "breastfeeding_excluded",
                "hygiene",
            )
        )
        b = all(
            k in self.odict["data"]["administration_details"]
            for k in (
                "calibrated_activity",
                "calibration_datetime",
                "administration_datetime",
                "administered_activity",
            )
        )
        c = all(
            k in self.odict["data"]["clearance_data"]
            for k in ("measurement", "curve_fit")
        )
        if "measurement" in self.odict["data"]["clearance_data"]:
            if isinstance(self.odict["data"]["clearance_data"]["measurement"], list):
                if (
                    "hours_elapsed"
                    not in self.odict["data"]["clearance_data"]["measurement"][0]
                ):
                    c = False
            else:
                if (
                    "hours_elapsed"
                    not in self.odict["data"]["clearance_data"]["measurement"]
                ):
                    c = False
        d = all(
            k in self.odict["data"]["patient_discharge"]
            for k in (
                "actual_datetime",
                "discharge_activity",
                "calculated_discharge_dose_rate_1m",
                "calculated_discharge_dose_rate_xm",
            )
        )
        if (
            "datetime"
            not in self.odict["data"]["patient_discharge"]["recommended_datetime"]
        ):
            d = False
        if "type_therapy" in self.odict["data"]["patient_details"]:
            if self.therapy_options_df.loc[
                self.odict["data"]["patient_details"]["type_therapy"],
                "generic_clearance",
            ]:
                c = True
            if not self.therapy_options_df.loc[
                self.odict["data"]["patient_details"]["type_therapy"], "inpatient"
            ]:
                d = True
        e = "restriction" in self.odict["data"]["restrictions"]
        if e:
            e = (
                "restriction_period"
                in self.odict["data"]["restrictions"]["restriction"][0]
            )
        forbidden = not (a and b and c and d and e)
        return forbidden

    def action_forbidden_comments(self):
        # inactivate the comments button until the business end
        forbidden = self.action_forbidden_common_comments_reports()
        if self.odict["data"]["additional_comments_to_patient"] != "0":
            forbidden = False
        return forbidden

    def action_forbidden_reports(self):
        forbidden = self.action_forbidden_common_comments_reports()
        f = self.odict["data"]["additional_comments_to_patient"] != "0"
        forbidden = forbidden or (not f)
        return forbidden

    # "Action required" meaning data is missing which can be filled via the window
    def action_required_patient(self):
        return not all(
            k in self.odict["data"]["patient_details"]
            for k in (
                "name",
                "id",
                "dob",
                "sex",
                "type_therapy",
                "site",
                "num_treatments_in_year",
                "pregnancy_excluded",
                "breastfeeding_excluded",
                "hygiene",
            )
        )

    def action_required_administration(self):
        return not all(
            k in self.odict["data"]["administration_details"]
            for k in (
                "calibrated_activity",
                "calibration_datetime",
                "administration_datetime",
                "administered_activity",
            )
        )

    def action_required_clearance(self):
        a = all(
            k in self.odict["data"]["clearance_data"]
            for k in ("detector_calibration_factor", "measurement", "curve_fit")
        )
        if "measurement" in self.odict["data"]["clearance_data"]:
            if isinstance(self.odict["data"]["clearance_data"]["measurement"], list):
                if (
                    "hours_elapsed"
                    not in self.odict["data"]["clearance_data"]["measurement"][0]
                ):
                    a = False
            else:
                if (
                    "hours_elapsed"
                    not in self.odict["data"]["clearance_data"]["measurement"]
                ):
                    a = False
        if "curve_fit" in self.odict["data"]["clearance_data"]:
            if (
                "sum_of_squared_residuals"
                not in self.odict["data"]["clearance_data"]["curve_fit"]
            ):
                a = False
        if "type_therapy" in self.odict["data"]["patient_details"]:
            if self.therapy_options_df.loc[
                self.odict["data"]["patient_details"]["type_therapy"],
                "generic_clearance",
            ]:
                a = True
        return not a

    def action_required_discharge(self):
        a = all(
            k in self.odict["data"]["patient_discharge"]
            for k in (
                "actual_datetime",
                "discharge_activity",
                "calculated_discharge_dose_rate_1m",
                "calculated_discharge_dose_rate_xm",
            )
        )
        b = (
            "datetime"
            in self.odict["data"]["patient_discharge"]["recommended_datetime"]
        )
        c = a and b
        if "type_therapy" in self.odict["data"]["patient_details"]:
            if not self.therapy_options_df.loc[
                self.odict["data"]["patient_details"]["type_therapy"], "inpatient"
            ]:
                c = True
        return not c

    def action_required_restrictions(self):
        e = "restriction" in self.odict["data"]["restrictions"]
        if e:
            e = (
                "restriction_period"
                in self.odict["data"]["restrictions"]["restriction"][0]
            )
        return not e

    def action_required_comments(self):
        return self.odict["data"]["additional_comments_to_patient"] == "0"

    def action_required(self):
        return any(
            [
                self.action_required_patient(),
                self.action_required_administration(),
                self.action_required_clearance(),
                self.action_required_discharge(),
                self.action_required_restrictions(),
                self.action_required_comments(),
            ]
        )

    @staticmethod
    def administered_activity(
        calib_activity: float,
        calib_datetime: datetime,
        admin_datetime: datetime,
        half_life: float,
        residual=None,
    ):
        hrs = (admin_datetime - calib_datetime).total_seconds() / 3600
        a0 = calib_activity * np.exp(-np.log(2) * hrs / half_life)
        # inf - inf = nan, make it return inf to trigger an error
        if residual is not None and not np.isinf(a0):
            residual_activity, residual_datetime = residual
            hrs = (residual_datetime - admin_datetime).total_seconds() / 3600
            a0 -= residual_activity * np.exp(np.log(2) * hrs / half_life)
        return a0

    @staticmethod
    def discharge_activity(
        admin_datetime: datetime,
        discharge_datetime: datetime,
        model: str,
        meaningful_parameters: list,
        a0: float,
    ):
        hrs = (discharge_datetime - admin_datetime).total_seconds() / 3600
        discharge_activity = 0
        if model == "biexponential":
            _, fraction_1, half_life_1, half_life_2 = meaningful_parameters
            discharge_activity = func_biexp(
                hrs, a0, fraction_1, np.log(2) / half_life_1, np.log(2) / half_life_2
            )
        elif model == "exponential":
            params = np.array([a0, np.log(2) / meaningful_parameters[1]])
            discharge_activity = func_exp(hrs, *params)
        return discharge_activity

    @staticmethod
    def discharge_dose_rate(
        admin_datetime,
        discharge_datetime,
        model,
        meaningful_parameters,
        measurement_distance,
    ):
        from glowgreen import Clearance_1m

        hrs = (discharge_datetime - admin_datetime).total_seconds() / 3600

        discharge_dose_rate_xm = 0
        if model == "biexponential":
            fit_parameters = np.array(
                [
                    meaningful_parameters[0],
                    meaningful_parameters[1],
                    np.log(2) / meaningful_parameters[2],
                    np.log(2) / meaningful_parameters[3],
                ]
            )
            discharge_dose_rate_xm = func_biexp(hrs, *fit_parameters)
        elif model == "exponential":
            fit_parameters = np.array(
                [meaningful_parameters[0], np.log(2) / meaningful_parameters[1]]
            )
            discharge_dose_rate_xm = func_exp(hrs, *fit_parameters)

        discharge_dose_rate_1m = 0
        cfit = Clearance_1m(model, meaningful_parameters, measurement_distance)
        if model == "biexponential":
            fit_parameters = np.array(
                [
                    cfit.model_params[0],
                    cfit.model_params[1],
                    np.log(2) / cfit.model_params[2],
                    np.log(2) / cfit.model_params[3],
                ]
            )
            discharge_dose_rate_1m = func_biexp(hrs, *fit_parameters)
        elif model == "exponential":
            fit_parameters = np.array(
                [cfit.model_params[0], np.log(2) / cfit.model_params[1]]
            )
            discharge_dose_rate_1m = func_exp(hrs, *fit_parameters)

        return discharge_dose_rate_1m, discharge_dose_rate_xm


def main():
    Gui()


if __name__ == "__main__":
    main()
