import fluidity_globals

import json
import os
import threading
import hashlib
import base64
import time as tm
import re
from datetime import datetime
from datetime import time
import random

import pandas as pd
pd.set_option('display.max_colwidth', None)

import PyPDF2
from docx import Document
from detoxify import Detoxify
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity

from llm_guard.input_scanners import BanCode, PromptInjection
from llm_guard.util import configure_logger
configure_logger("WARNING")

# Initialize the desired Detoxify model
# 'original' is a good general-purpose model. Other options include 'unbiased', 'multilingual'.
detoxifier = Detoxify('original')


class LLM:
    def __init__(self, client, model_name):
        self.client = client
        self.model_name = model_name

    def request(self, role, user_prompt, context=None):
        try:
            # Craft the messages to pass to chat.completions.create
            prompt = [{'role':'system', 'content': role}]

            # Append the prompt/response history if there is a context
            if context is not None:
                for p, r in context:
                    prompt.append({'role': 'user', 'content': p})
                    prompt.append({'role': 'system', 'content': r})

            prompt.append({'role': 'user', 'content': user_prompt})

            response = self.client.chat.completions.create(
                model=self.model_name,
                messages=prompt,
                temperature=0.1
            )

            return response.choices[0].message.content.strip()

        except Exception as e:
            return f'Sorry, I encountered the following error: \n {e}'

# Global cache for LLMs. Populate as required, in client code
class LLMCache:
    def __init__(self):
        self.llms = {}
        self.default = None

    def addLLM(self, name, llm):
        self.llms[name] = llm

    def setDefault(self, default):
        self.default = default

    def getLLM(self, name):
        return self.llms[name]
    
    def getDefaultLLM(self):
        return self.llms[self.default]


llmCache = LLMCache()


class WebUtils:
    @staticmethod
    def parse_headers(request):
        headers = {}
        lines = request.split('\r\n')
        for line in lines[1:]:
            if ": " in line:
                key, value = line.split(': ', 1)
                headers[key] = value

        return headers
    
    @staticmethod
    def generate_accept_key(sec_websocket_key):
        GUID = '258EAFA5-E914-47DA-95CA-C5AB0DC85B11'
        sha1 = hashlib.sha1((sec_websocket_key + GUID).encode()).digest()
        accept_key = base64.b64encode(sha1).decode()
        return accept_key

    @staticmethod
    def recv_frame(sock):
        # Read the first 2 bytes
        first_two = sock.recv(2)
        if len(first_two) < 2:
            return None
        b1, b2 = first_two
        fin = b1 & 0x80
        opcode = b1 & 0x0F
        mask = b2 & 0x80
        payload_length = b2 & 0x7F

        # Extended payload length
        if payload_length == 126:
            payload_length = int.from_bytes(sock.recv(2), 'big')
        elif payload_length == 127:
            payload_length = int.from_bytes(sock.recv(8), 'big')

        # Masking key (client to server frames are masked)
        mask_key = sock.recv(4)

        payload_data = sock.recv(payload_length)
        # Unmask payload
        decoded = bytes(b ^ mask_key[i % 4] for i, b in enumerate(payload_data))
        return decoded.decode().strip()

    @staticmethod
    def send_frame(sock, message):
        # Build a frame (text, no masking from server)
        payload = message.strip().encode()
        header = bytearray()

        # FIN + text frame
        header.append(0x81)

        # Payload length
        length = len(payload)
        if length <= 125:
            header.append(length)
        elif length <= 65535:
            header.append(126)
            header.extend(length.to_bytes(2, 'big'))
        else:
            header.append(127)
            header.extend(length.to_bytes(8, 'big'))

        # Send header + payload
        sock.sendall(header + payload)


class Utils:
    llm = None

    @staticmethod
    def chunk_text(text, chunk_size=200, overlap=50):
        tokens = text.split()
        chunks = []
        start = 0
        while start < len(tokens):
            end = min(start + chunk_size, len(tokens))
            chunk = " ".join(tokens[start:end])
            chunks.append(chunk)
            start += chunk_size - overlap

        return chunks
    
    @staticmethod
    def retrieve_top_pct(embed_model, query, vector_db, top_pct):
        query_emb = embed_model.encode([query])[0]
        vectors = [emb for _, emb in vector_db]

        # Compute cosine similarity
        similarities = cosine_similarity([query_emb], vectors)[0]

        # Get indices of top K
        k = round(len(vectors) * top_pct)
        top_k_indices = similarities.argsort()[-k:][::-1]

        return [(vector_db[i][0], similarities[i]) for i in top_k_indices]

    @staticmethod
    def print(text, remoteClient):
        if remoteClient is None:
            print(text)
        else:
            WebUtils.send_frame(remoteClient, text)

    @staticmethod
    def input(text, remoteClient):
        if remoteClient is None:
            return input(text)
        else:
            WebUtils.send_frame(remoteClient, text)
            return WebUtils.recv_frame(remoteClient)

    @staticmethod
    def llmFormat(text, type, format):
        if Utils.llm is None:
            Utils.llm = llmCache.getDefaultLLM()
        
        role = "You are an AI assistant skilled at converting names, dates and numbers from one format to another."
        instructions = f"Please extract the {type} from the string '{text}' and convert it to the format '{format}'. If there is no {type} present, return 'Error'."
        instructions += f" Provide only the correctly formatted {type} with no extra information, introduction or summary."
        
        return Utils.llm.request(role, instructions)

    @staticmethod
    def thread_task(task : "Task", delay : float):
        # Allow time for all threads to get started
        tm.sleep(delay)

        # Perform the task
        task.execute()

    @staticmethod
    def main_thread_task(tg : "TaskGroup"):
        for i, task in enumerate(tg.tasks):
            thread = threading.Thread(target=Utils.thread_task, args=(task, (len(tg.tasks) - i) * 0.25))
            tg.task_thread[task.name()] = thread
            thread.start()

        for task_name, thread in tg.task_thread.items():
            thread.join()

    @staticmethod
    def run_tg(tg):
        main_thread = threading.Thread(target=Utils.main_thread_task, args=(tg,))
        main_thread.start()
        main_thread.join()

    @staticmethod
    def is_email(single_input):
        email = single_input.replace(' ', '').replace('\n', '').replace('\t', '').replace('\r', '')
        return re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', email)

    @staticmethod
    def is_toxic(single_input):
        single_result = detoxifier.predict(single_input)
        for key, value in single_result.items():
            if value >= 0.5:
                return True
            
        return False
    
    @staticmethod
    def is_llm_unsafe(single_input):
        if Utils.is_email(single_input):
            return False
        
        for scanner in [PromptInjection(), BanCode()]:
            sanitized_prompt, is_valid, risk_score = scanner.scan(single_input)
            if risk_score >= 0.5:
                return True
            
        return False
    
    @staticmethod
    def get_current_date():
        return datetime.now()
    
    @staticmethod
    def get_current_time():
        return datetime.now().time()
    
    @staticmethod
    def get_current_session():
        session = "morning"

        time_now = Utils.get_current_time()
        if time_now > time(hour=18, minute=0, second=0):
            session = "evening"
        elif time_now > time(hour=12, minute=0, second=0):
            session = "afternoon"

        return session

    @staticmethod
    def is_number(s : str):
        try:
            float(s)
            return True
        except ValueError:
            return False
        
    @staticmethod
    def is_date(s : str, format : str):
        try:
            datetime.strptime(s, format)
            return True
        except ValueError:
            return False
        
    @staticmethod
    def get_random_numbers(start, end, n):
        r_nums = random.sample(range(start, end + 1), n)
        r_nums.sort()
        return r_nums
        
    @staticmethod
    def get_number(s : str):
        for w in s.split():
            if Utils.is_number(w):
                if '.' in w:
                    return float(w)
                else:
                    return int(w)
            
        return None
    
    @staticmethod
    def get_ordinal(n):
        r = n % 100
        if 11 <= r <= 19:
            return str(n) + "th"
        
        r = n % 10
        if r == 1:
            return str(n) + "st"
        if r == 2:
            return str(n) + "nd"
        if r == 3:
            return str(n) + "rd"
        
        return str(n) + "th"

    @staticmethod
    def get_list(s : str):
        c_first = s.find('{')
        c_last = s.find('}') + 1
        json_str = s[c_first:c_last]
        
        s_dict = json.loads(json_str)

        l = []
        for k, v in s_dict.items():
            l.append(v)

        return l

    @staticmethod
    def get_dict(s : str):
        c_first = s.find('{')
        c_last = s.rfind('}') + 1
        json_str = s[c_first:c_last]
        
        s_dict = json.loads(json_str)

        d = {}
        for k, v in s_dict.items():
            d[k] = v

        return d

    @staticmethod
    def get_json(s : str):
        c_first = s.find('{')
        c_last = s.rfind('}') + 1

        return s[c_first:c_last]

    @staticmethod
    def subst_placeholders(s : str, task_outputs):
        try:
            # Placeholders are of the form {$task_name} or {$task_name:$key}
            placeholders = re.findall(r"\{([^}]*)\}", s)
            placeholders = [p.strip() for p in placeholders if len(p.strip()) > 0]
            
            for p in placeholders:
                val = ""
                if p.lower() == "today":
                    # Get today's date in DayOfWeek, dd-Mmm-yyyy format
                    val = datetime.now().strftime("%A, %d-%b-%Y")
                elif ':' in p:
                    # Look up list or dictionary entry
                    k, v = p.split(':')
                    if Utils.is_number(v):
                        val = task_outputs[k][int(v)]
                    else:
                        val = task_outputs[k][v]
                else:
                    val = task_outputs[p]

                if isinstance(val, pd.DataFrame):
                    val = val.to_json()

                # Substitute placeholder for value
                s = s.replace('{' + p + '}', str(val))

        except:
            raise ValueError("subst_placeholders(): Unable to substitute for '" + p + "' in string: " + s)

        return s


class Task:
    def __init__(self, name, instructions, output_type="string", role=None, llm_or_code="llm", llm=None):
        self._name = name
       
        self.instructions = instructions.strip()
        if self.instructions[-1] != '.' and llm_or_code == "llm":
            self.instructions = self.instructions + '.'

        self.output_type = output_type.lower()
        if llm_or_code == "llm":
            self.role = "You are an AI assistant who just provides answers to questions without any introduction or summary" \
                if role is None else role.strip()
        else:
            self.role = None
        
        self.llm_or_code = llm_or_code.lower()
        self.llm = llm
        self.context = None
        self.remote_client = None

        # This is a reference to the workflow's task_outputs map
        self.workflow_task_outputs = None

        # The next task to execute
        self.next_task = None

    def to_dict(self):
        task_dict = {}

        task_dict["name"] = self._name
        task_dict["type"] = "Task"
        task_dict["instructions"] = self.instructions
        task_dict["role"] = self.role
        task_dict["outputType"] = self.output_type
        task_dict["llmOrCode"] = self.llm_or_code
        task_dict["maintainContext"] = True if self.context is not None else False
        task_dict["nextTask"] = self.next_task

        return task_dict
    
    @staticmethod
    def from_dict(task_dict):
        task = Task(task_dict["name"], task_dict["instructions"], task_dict["outputType"],
            task_dict["role"], task_dict["llmOrCode"])
        
        task.setNextTask(task_dict["nextTask"])
        if task_dict["maintainContext"]:
            task.maintainContext()

        return task

    def setRemoteClient(self, remote_client):
        self.remote_client = remote_client

    def maintainContext(self):
        self.context = []

    def name(self) -> str:
        return self._name
    
    def setNextTask(self, next_task):
        self.next_task = next_task

    def execute(self):
        if self.llm_or_code == "llm":
            # Substitute placeholders in role and instructions. This needs to be done each time the
            # task is executed, because the contents of task_outputs may change
            role = Utils.subst_placeholders(self.role, self.workflow_task_outputs)
            instructions = Utils.subst_placeholders(self.instructions, self.workflow_task_outputs)

            if self.output_type[:3] in ["str", "num"]:
                instructions = instructions + " Give the answer only, with no other explanation or information."
            elif self.output_type == "list":
                instructions = instructions + " Give your response in this JSON format: {$(item number) : $item}."
            elif self.output_type[:4] == "dict":
                instructions = instructions + " Give your response in this JSON format: {$key : $value}."
            elif self.output_type[:4] == "json":
                instructions = instructions + " Give your response in JSON format."

            # Run the LLM request
            if self.llm is None:
                self.llm = llmCache.getDefaultLLM()

            if self.context is None:
                # Simple call
                self.output = self.llm.request(role, instructions)
            else:
                # Send and update the context, updating the context with the response
                self.output = self.llm.request(role, instructions, self.context)
                self.context.append((instructions, self.output))

            if self.output_type[:3] in ["str", "obj"]:
                self.workflow_task_outputs[self._name] = self.output
            elif self.output_type[:3] == "num":
                self.workflow_task_outputs[self._name] = Utils.get_number(self.output)
            elif self.output_type[:4] == "dict":
                self.workflow_task_outputs[self._name] = Utils.get_dict(self.output)
            elif self.output_type == "list":
                self.workflow_task_outputs[self._name] = Utils.get_list(self.output)
            elif self.output_type == "json":
                self.workflow_task_outputs[self._name] = Utils.get_json(self.output)
            else:
                raise ValueError("Task:execute(): unknown LLM return type: " + self.output_type)
        else:
            # Instructions contain Python code to execute
            # Can also be a filename containing the code
            if self.instructions.startswith("file:"):
                with open(self.instructions[5:], 'r') as file:
                    self.instructions = file.read()
                
            namespace = {
                'next_task' : None,
                'task_outputs' : self.workflow_task_outputs,
                'Utils' : Utils,
                'remote_client' : self.remote_client
            }
            exec(self.instructions.strip().replace('\t', '    '), namespace)

            if namespace['result'] is not None:
                self.output = namespace['result']
                self.workflow_task_outputs[self._name] = self.output
            else:
                self.workflow_task_outputs[self._name] = None

            if namespace["next_task"] is not None:
                self.next_task = namespace["next_task"]

    def run(self):
        self.execute()


# A set of tasks to be run in parallel
#
class TaskGroup(Task):
    def __init__(self, name, instructions, output_type="string", role=None, llm_or_code="llm", llm=None):
        raise ValueError("TaskGroup(): constructor for class Task cannot be used")
    
    def to_dict(self):
        taskgroup_dict = {}

        taskgroup_dict["name"] = self._name
        taskgroup_dict["type"] = "TaskGroup"
        taskgroup_dict["tasks"] = [task.to_dict() for task in self.tasks]
        taskgroup_dict["nextTask"] = self.next_task

        return taskgroup_dict
    
    @staticmethod
    def from_dict(taskgroup_dict):
        task_group = TaskGroup(taskgroup_dict["name"])
        task_group.tasks = [Task.from_dict(task_dict) for task_dict in taskgroup_dict["tasks"]]
        task_group.setNextTask(taskgroup_dict["nextTask"])

        return task_group
    
    def __init__(self, name):
        self._name = name
        self.tasks = []
        self.workflow_task_outputs = None
        self.task_thread = {}
        self.next_task = None

    def execute(self):
        main_thread = threading.Thread(target=Utils.main_thread_task, args=(self,))
        main_thread.start()
        main_thread.join()

    def get_task_by_name(self, task_name):
        for t_name, t in zip([t.name() for t in self.tasks], self.tasks):
            t_name = t_name.replace(self._name + '-', '')
            if t_name == task_name:
                return t
            
        return None
    
    def updateTasks(self):
        for task in self.tasks:
            task.workflow_task_outputs = self.workflow_task_outputs

    def addTask(self, task : Task):
        # Cannot add tasks after task group has been added to a workflow
        if self.workflow_task_outputs is not None:
            raise ValueError("TaskGroup.addTask(): cannot add new tasks after addition to workflow")

        # Must have unique name
        if self.get_task_by_name(task.name()) is not None:
            raise ValueError("Tasks in a task group must have unique names")

        task._name = self._name + '-' + task.name()
        self.tasks.append(task)

    def removeTask(self, task_name):
        pos = None
        for i, t in enumerate(self.tasks):
            if t.name() == self.name + '-' + task_name:
                pos = i
                break

        if pos is not None:
            del self.tasks[pos]
        else:
            raise ValueError(f"No task named {task_name} was found in workflow {self.name}")


class Branch(Task):
    def __init__(self, name, instructions, output_type="string", role=None, llm_or_code="llm", llm=None):
        super().__init__(name, instructions, output_type, role, "code", None)

    def to_dict(self):
        task_dict = {}

        task_dict["name"] = self._name
        task_dict["type"] = "Branch"
        task_dict["instructions"] = self.instructions
        task_dict["nextTasks"] = self.next_tasks

        return task_dict
    
    @staticmethod
    def from_dict(task_dict):
        branch = Branch(task_dict["name"], task_dict["instructions"])
        branch.setNextTasks(task_dict["nextTasks"])

        return branch

    # All the possible options. Used when generating a graph
    def setNextTasks(self, next_tasks):
        self.next_tasks = next_tasks
        
    def execute(self):
        # Instructions contain code to execute and set next_task. Does not
        # write to task_outputs
        if self.instructions.startswith("file:"):
            with open(self.instructions[5:], 'r') as file:
                self.instructions = file.read()

        namespace = {
            'next_task' : None,
            'task_outputs' : self.workflow_task_outputs,
            'Utils' : Utils,
            'remote_client' : self.remote_client
        }
        exec(self.instructions.strip().replace('\t', '    '), namespace)

        if namespace['next_task'] is not None:
            self.next_task = namespace['next_task']


class Workflow:
    def __init__(self, name):
        self._name = name
        self.tasks = []
        self.task_thread = {}
        self.task_outputs = {}
        self.remote_client = None
        self.first_task_name = None

    def setFirstTask(self, first_task_name):
        self.first_task_name = first_task_name

    def dumpJSON(self, file_path):
        file_path = file_path.strip()
        if file_path[-1] != os.sep:
            file_path += os.sep

        file_path += self._name + ".json"
        wf_dict = self.to_dict()
        with open(file_path, 'w') as f:
            json.dump(wf_dict, f, indent=4)

    @staticmethod
    def retrieveJSON(file_path):
        file_path = file_path.strip()
        with open(file_path, 'r') as f:
            return f.read()

    def to_dict(self):
        wf_dict = {}

        wf_dict["name"] = self._name
        wf_dict["type"] = "Workflow"
        wf_dict["firstTask"] = self.first_task_name
        wf_dict["tasks"] = []

        task = self.get_task_by_name(self.first_task_name)
        while task is not None and task.name() not in [t["name"] for t in wf_dict["tasks"]]:
            wf_dict["tasks"].append(task.to_dict())
            task = self.get_task_by_name(task.next_task)

        return wf_dict
    
    @staticmethod
    def from_dict(wf_dict):
        wf = Workflow(wf_dict["name"])
        wf.setFirstTask(wf_dict["firstTask"])

        for task_dict in wf_dict["tasks"]:
            task_class_obj = globals()[task_dict["type"]]
            task = getattr(task_class_obj, "from_dict")(task_dict)
            wf.addTask(task)

        return wf

    def setRemoteClient(self, remote_client):
        self.remote_client = remote_client

        # Update all tasks
        for task in self.tasks:
            task.setRemoteClient(self.remote_client)

    def name(self):
        return self._name
    
    def outputs(self):
        return self.task_outputs
    
    def execute(self, first_task_name=None):
        # Find first task
        if first_task_name is None:
            first_task_name = self.first_task_name

        task = self.get_task_by_name(first_task_name)
        while task is not None:
            task.run()
            if task.next_task is not None and task.next_task.lower() == "stop":
                return
            
            task = self.get_task_by_name(task.next_task)
        
    def run(self, first_task_name=None):
        self.execute(first_task_name)

    def get_task_by_name(self, task_name):
        for t in self.tasks:
            if t.name() == task_name:
                return t
            
        return None
    
    def addTasks(self, *tasks):
        for task in tasks:
            self.addTask(task)

    def addTask(self, task : Task):
        # Must have unique name
        if self.get_task_by_name(task.name()) is not None:
            raise ValueError("Tasks in a workflow must have unique names")
            
        task.workflow_task_outputs = self.task_outputs
        task.remote_client = self.remote_client

        if isinstance(task, TaskGroup):
            task.updateTasks()
            
        self.tasks.append(task)

    def removeTask(self, task_name):
        pos = None
        for i, t in enumerate(self.tasks):
            if t.name() == task_name:
                pos = i
                break

        if pos is not None:
            del self.tasks[pos]
        else:
            raise ValueError(f"No task named {task_name} was found in workflow {self.name}")
        

# Makes repeated cmd line or web chat requests until user
# enters the information required in a recognisable format
class UserInputTask(Task):
    def __init__(self, name, instructions, output_type="string", role=None, llm_or_code="llm", llm=None):
        raise NameError("class UserInput: cannot use Task() constructor")
    
    def to_dict(self):
        task_dict = {}

        task_dict["name"] = self._name
        task_dict["type"] = "UserInputTask"
        task_dict["dataType"] = self.data_type
        task_dict["format"] = self.format
        task_dict["prompt1"] = self.prompt1
        task_dict["prompt2"] = self.prompt2
        task_dict["toxicMsg"] = self.toxic_msg
        task_dict["greeting"] = self.greeting
        task_dict["nextTask"] = self.next_task

        return task_dict
    
    @staticmethod
    def from_dict(task_dict):
        task = UserInputTask(task_dict["name"], None, task_dict["dataType"], task_dict["format"],
            task_dict["prompt1"], task_dict["prompt2"],
            task_dict["toxicMsg"], task_dict["greeting"])
        
        task.setNextTask(task_dict["nextTask"])
        return task
        
    def __init__(
            self, name, instructions = None,
            dataType = None,                    # E.g. name, date, zip code
            format = None,                      # 'dd-Mmm-yy', 'ddddd', 'first_name surname' etc.
            prompt1 = None,                     # Initial request
            prompt2 = None,                     # Requests after first failure,
            toxicMsg = None,                    # Response to toxic input
            greeting = False
    ):
        self._name = name
        self.instructions = instructions
        self.data_type = dataType
        self.format = format
        self.prompt1 = prompt1
        self.prompt2 = prompt2
        self.toxic_msg = toxicMsg
        self.greeting = greeting
        self.remote_client = None
        self.validation_function = None # Overrides default safety checking, e.g. email addresses can be incorrectly flagged as code injection 

        self.next_task = None

    # Implement in subclasses
    def validate(self):
        pass

    def setValidationFunction(self, fn):
        self.validation_function = fn

    def is_unsafe(self, text):
        if self.validation_function is not None:
            return not self.validation_function(text)
        
        postcode_or_email = False
        try:
            postcode_or_email = self.data_type is not None and self.data_type.lower() == "email" \
            or "postcode" in text.lower() or "post code" in ' '.join(text.lower().split())

        except Exception as e:
            print(f"Error in is_unsafe(): {e}")
            return False

        return Utils.is_toxic(text) or (not postcode_or_email and Utils.is_llm_unsafe(text))

    def execute(self):
        if self.data_type is not None:
            self.prompt1 = Utils.subst_placeholders(self.prompt1.strip(), self.workflow_task_outputs)
            self.prompt2 = Utils.subst_placeholders(self.prompt2.strip(), self.workflow_task_outputs)

            prompt1 = "Good " + Utils.get_current_session() + ", " + self.prompt1.lower() + ": " if self.greeting else self.prompt1 + ": "

            user_input = Utils.input(prompt1, self.remote_client)
            while self.is_unsafe(user_input):
                user_input = Utils.input(self.toxic_msg + ": ", self.remote_client)
            
            fmt_user_input = Utils.llmFormat(user_input, self.data_type, self.format)

            while fmt_user_input == 'Error':
                user_input = Utils.input(self.prompt2 + ": ", self.remote_client)
                
                while self.is_unsafe(user_input):
                    user_input = Utils.input(self.toxic_msg + ": ", self.remote_client)

                fmt_user_input = Utils.llmFormat(user_input, self.data_type, self.format)

            # Acceptable content
            self.output = fmt_user_input
            self.workflow_task_outputs[self._name] = self.output

        self.validate()

# Asks the user to input 3 randomly chosen characters from the password
# Instances provide code to implement obtaining the password, possibly from a database
class PasswordInputTask(UserInputTask):
    def to_dict(self):
        task_dict = {}

        task_dict["name"] = self._name
        task_dict["type"] = "PasswordInputTask"
        task_dict["instructions"] = self.instructions
        task_dict["nextTask"] = self.next_task

        return task_dict
    
    @staticmethod
    def from_dict(task_dict):
        task = PasswordInputTask(task_dict["name"], task_dict["instructions"])
        task.setNextTask(task_dict["nextTask"])
        return task
    
    def getPwd(self):
        if self.instructions.startswith("file:"):
            with open(self.instructions[5:], 'r') as file:
                self.instructions = file.read()

        namespace = {
            'password' : None,
            'task_outputs' : self.workflow_task_outputs,
            'Utils' : Utils,
            'remote_client' : self.remote_client
        }
        
        exec(self.instructions.strip().replace('\t', '    '), namespace)
        return namespace['password']

    def validate(self):
        password = self.getPwd()
        if password is None:
            Utils.print("Failed to identify user. Goodbye.", self.remote_client)
            exit()

        char_positions = Utils.get_random_numbers(1, len(password), 3)
        char_ordinals = list(map(Utils.get_ordinal, char_positions))

        pwd_string = Utils.input(
            f"Please enter the {char_ordinals[0]}, {char_ordinals[1]} and {char_ordinals[2]} characters of your password: ",
            self.remote_client)
        for pwd_char, char_pos in zip(pwd_string.split(), char_positions):
             if password[char_pos - 1] != pwd_char:
                    Utils.print("Password not validated", self.remote_client)
                    self.workflow_task_outputs[self._name] = 'Error'
                    return

        Utils.print("Password validated", self.remote_client)
        self.workflow_task_outputs[self._name] = 'OK'


# Caches CSV files as DataFrames, any other file types are read into a string, e.g. PDF files. Status of
# each file is stored in task_outputs[self.name]['<filename> Status']
class FileLoaderTask(UserInputTask):
    def __init__(self, name, files):
        super().__init__(name, None, None, None, None, None, False)

        self.files = files      # Dictionary mapping file name => file path

    def to_dict(self):
        task_dict = {}

        task_dict["name"] = self._name
        task_dict["type"] = "FileLoaderTask"
        task_dict["files"] = self.files
        task_dict["nextTask"] = self.next_task

        return task_dict
    
    @staticmethod
    def from_dict(task_dict):
        task = FileLoaderTask(task_dict["name"], task_dict["files"])
        task.setNextTask(task_dict["nextTask"])
        return task

    def execute(self):
        self.workflow_task_outputs[self._name] = {}
        for file_name, file_path in self.files.items():

            file_name = file_name.strip()
            file_path = file_path.strip()

            if file_path.lower().endswith(".csv"):
                self.workflow_task_outputs[self._name][file_name] = pd.read_csv(file_path)
            elif file_path.lower().endswith(".docx"):
                # Open the Word document
                doc = Document(file_path)

                # Extract text from each paragraph
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)

                # Combine paragraphs into a single string
                self.workflow_task_outputs[self._name][file_name] = '\n'.join(full_text)
            elif file_path.lower().endswith(".pdf"):
                text = ""
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text += page.extract_text()

                self.workflow_task_outputs[self._name][file_name] = text
            else:
                with open(file_path, 'r') as f:
                    self.workflow_task_outputs[self._name][file_name] = f.read()

            self.workflow_task_outputs[self._name][file_name + ' Status'] = None


# Chunks all pre loaded files, creates vector embeddings and stores them in a vector database, then creates
# a context from the top_pct (10% by default) of chunks by relevance to the (pre loaded) query
class RAGRetrievalTask(UserInputTask):
    def __init__(
            self, name,
            loader,         # the name of the file loader task
            query,          # query string including {} placeholders
            top_pct = 0.1
        ):
        super().__init__(name, None, None, None, None, None, False)

        self.loader = loader.strip()
        self.query = query.strip()
        self.top_pct = top_pct

        self.embed_model = None
        self.embeddings = None
        self.vector_db = None

    def setCache(self, ragCache):
        self.embed_model = ragCache.embed_model
        self.embeddings = ragCache.embeddings
        self.vector_db = ragCache.vector_db

    def to_dict(self):
        task_dict = {}

        task_dict["name"] = self._name
        task_dict["type"] = "RAGRetrievalTask"
        task_dict["loader"] = self.loader
        task_dict["query"] = self.query
        task_dict["topPct"] = self.top_pct
        task_dict["nextTask"] = self.next_task

        return task_dict
    
    @staticmethod
    def from_dict(task_dict):
        task = RAGRetrievalTask(task_dict["name"], task_dict["loader"], task_dict["query"], task_dict["topPct"])
        task.setNextTask(task_dict["nextTask"])

        return task

    def execute(self):
        if self.vector_db is None:
            # Combine all files into a single string
            text = ""
            for file_name in self.workflow_task_outputs[self.loader]:
                file = self.workflow_task_outputs[self.loader][file_name]
                if isinstance(file, pd.DataFrame):
                    text += file.to_string(index=False)
                else:
                    text += str(file)

            # Chunk the document
            chunks = Utils.chunk_text(text)

            # Generate embeddings
            self.embed_model = SentenceTransformer('all-MiniLM-L6-v2')
            self.embeddings = self.embed_model.encode(chunks)
            self.vector_db = list(zip(chunks, self.embeddings))       

        query_str = Utils.subst_placeholders(self.query, self.workflow_task_outputs)
        top_chunks = Utils.retrieve_top_pct(self.embed_model, query_str, self.vector_db, self.top_pct)
        context = "\n\n".join([chunk for chunk, score in top_chunks])

        self.workflow_task_outputs[self._name] = context


# Useful global cache class for e.g. chat apps in which all threads query the same file(s)
class RAGCache:
    def __init__(self, file_paths):
        self.file_paths = file_paths
        self.files = []

        self.embed_model = None
        self.embeddings = None
        self.vector_db = None

    def load(self):
        # Get the text for each file
        for file_path in self.file_paths:
            file_path = file_path.strip()

            if file_path.lower().endswith(".csv"):
                self.files.append(pd.read_csv(file_path))
            elif file_path.lower().endswith(".docx"):
                # Open the Word document
                doc = Document(file_path)

                # Extract text from each paragraph
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)

                # Combine paragraphs into a single string
                self.files.append('\n'.join(full_text))
            elif file_path.lower().endswith(".pdf"):
                text = ""
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text += page.extract_text()

                self.files.append(text)
            else:
                with open(file_path, 'r') as f:
                    self.files.append(f.read())

        # Combine all files into a single string
        text = ""
        for file in self.files:
            if isinstance(file, pd.DataFrame):
                text += file.to_string(index=False)
            else:
                text += str(file)

        # Chunk the document
        chunks = Utils.chunk_text(text)

        # Generate embeddings
        self.embed_model = SentenceTransformer('all-MiniLM-L6-v2')
        self.embeddings = self.embed_model.encode(chunks)
        self.vector_db = list(zip(chunks, self.embeddings))       


# Makes repeated cmd line or web chat requests for submission to the LLM
class ChatbotTask(UserInputTask):
    def to_dict(self):
        task_dict = {}

        task_dict["name"] = self._name
        task_dict["response"] = self.response
        task_dict["type"] = "ChatbotTask"
        task_dict["prompt"] = self.prompt
        task_dict["toxicMsg"] = self.toxic_msg
        task_dict["greeting"] = self.greeting
        task_dict["initialPromptOnly"] = self.initialPromptOnly
        task_dict["nextTask"] = self.next_task

        return task_dict
    
    @staticmethod
    def from_dict(task_dict):
        task = ChatbotTask(task_dict["name"], None, task_dict["response"], task_dict["prompt1"],
            task_dict["toxicMsg"], task_dict["greeting"], task_dict["initialPromptOnly"])
        
        task.setNextTask(task_dict["nextTask"])
        return task
        
    def __init__(
            self, name,
            response = None,                   # The LLM task answering the question
            prompt = None,                     # Initial request
            toxicMsg = None,                   # Response to toxic input
            greeting = False,
            initialPromptOnly = False
    ):
        self._name = name
        self.response = response
        self.prompt = prompt
        self.toxic_msg = toxicMsg
        self.greeting = greeting
        self.initialPromptOnly = initialPromptOnly
        self.greeted = False
        self.prompted = False

        self.validation_function = None
        self.data_type = None

        self.remote_client = None
        self.next_task = None

    def execute(self):
        response_msg = ""
        if self.response in self.workflow_task_outputs:
            response_msg = self.workflow_task_outputs[self.response]
            if not self.initialPromptOnly:
                Utils.print(response_msg, self.remote_client)

        self.prompt = Utils.subst_placeholders(self.prompt.strip(), self.workflow_task_outputs)

        prompt = ""
        if not (self.initialPromptOnly and self.prompted):
            if self.greeting and not self.greeted:
                prompt = "Good " + Utils.get_current_session() + ", " + self.prompt.lower() + ": "
                self.greeted = True
            else:
                prompt = self.prompt + ": "

            self.prompted = True

        user_input = Utils.input(prompt if prompt != "" else response_msg, self.remote_client)
        while user_input is None:
            user_input = Utils.input("", self.remote_client)

        while self.is_unsafe(user_input):
            user_input = Utils.input(self.toxic_msg + ": ", self.remote_client)
        
        # Acceptable content
        print(f"ChatbotTask.execute(): Received user_input: {user_input}")
        self.output = user_input
        self.workflow_task_outputs[self._name] = self.output

# Requests the LLM to extract specific information fields from text 
class FormTask(Task):
    def __init__(self, name, instructions, output_type="string", role=None, llm_or_code="llm", llm=None):
        raise NameError("class FormTask: cannot use Task() constructor")
    
    def __init__(
        self,
        name,
        fields,         # Dictionary mapping form field name => name in LLM instructions
        formats=None    # Maps form field name => to format (e.g. "HH:mm:ss")
    ):
        self._name = name
        self.fields = fields
        self.formats = formats
        
        self.workflow_task_outputs = None
        self.next_task = None

    def to_dict(self):
        task_dict = {}

        task_dict["name"] = self._name
        task_dict["type"] = "FormTask"
        task_dict["fields"] = self.fields
        task_dict["formats"] = self.formats
        task_dict["nextTask"] = self.next_task

        return task_dict
    
    @staticmethod
    def from_dict(task_dict):
        task = FormTask(task_dict["name"], task_dict["fields"], task_dict["formats"])
        task.setNextTask(task_dict["nextTask"])

        return task

    def execute(self):
        if self._name not in self.workflow_task_outputs:
            self.workflow_task_outputs[self._name] = {k : 'Unknown' for k, v in self.fields.items()}

        fill_or_update = self.workflow_task_outputs[self._name].get("FillOrUpdate", "fill").lower()

        n = 0
        if fill_or_update == "fill":
            n = sum(1 for k, v in self.workflow_task_outputs[self._name].items() if k in self.fields.keys() and v == 'Unknown')
        else:
            n = sum(1 for k, v in self.workflow_task_outputs[self._name].items() if k in self.fields.keys())

        fieldsStr = ""

        # Build up a string of required fields
        i = 0
        for k, v in self.fields.items():
            if self.workflow_task_outputs[self._name][k] != 'Unknown' and fill_or_update == 'fill':
                continue

            sep = ", "
            if i == n - 2:
                sep = " and "
            elif i == n - 1:
                sep = ""

            formatStr = ""
            if self.formats is not None and k in self.formats:
                formatStr = " in the format " + self.formats[k]

            fieldsStr += v + formatStr + sep
            i += 1

        self.workflow_task_outputs[self._name]["Fields"] = fieldsStr


class FormFillTask(FormTask):
    def __init__(
        self,
        name,
        formName,
        formResponseName,
        fields,
        validators=None     # Maps form field name => function to validate what the LLM returns
    ):
        super().__init__(name, fields=None, formats=None)

        self.form_name = formName
        self.form_response_name = formResponseName
        # The LLM uses underscores instead of spaces in returned field names
        self.fields = {k : '_'.join(v.split()) for k, v in fields.items()}
        self.validators = validators

    def to_dict(self):
        task_dict = {}

        task_dict["name"] = self._name
        task_dict["type"] = "FormFillTask"
        task_dict["formName"] = self.form_name
        task_dict["formResponseName"] = self.form_response_name
        task_dict["fields"] = self.fields
        task_dict["validators"] = self.validators
        task_dict["nextTask"] = self.next_task

        return task_dict
    
    @staticmethod
    def from_dict(task_dict):
        task = FormFillTask(task_dict["name"], task_dict["formName"], task_dict["formResponseName"],
            task_dict["fields"], task_dict["validators"])
    
        task.setNextTask(task_dict["nextTask"])
        return task

    def execute(self):
        fill_or_update = self.workflow_task_outputs[self.form_name].get("FillOrUpdate", "fill").lower()
        for k, v in self.fields.items():
            if fill_or_update == "update":
                val = self.workflow_task_outputs[self.form_response_name][v]

                if self.validators is not None and k in self.validators:
                    if not getattr(fluidity_globals, self.validators[k])(val):
                        # Validation of field returned by LLM failed
                        val = 'Unknown'

                if val != 'Unknown':
                    self.workflow_task_outputs[self.form_name][k] = val
            else:
                if self.workflow_task_outputs[self.form_name][k] == 'Unknown':
                    val = self.workflow_task_outputs[self.form_response_name][v]

                    if self.validators is not None and k in self.validators:
                        if not getattr(fluidity_globals, self.validators[k])(val):
                            # Validation of field returned by LLM failed
                            val = 'Unknown'

                    self.workflow_task_outputs[self.form_name][k] = val

