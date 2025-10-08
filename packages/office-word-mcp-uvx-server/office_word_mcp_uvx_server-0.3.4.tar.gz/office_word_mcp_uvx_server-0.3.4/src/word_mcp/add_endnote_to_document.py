"""向Word文档添加尾注的功能模块。"""

import os
from typing import Dict, Any
from docx import Document
from .exceptions import FileError, DocumentError, ValidationError


def add_endnote_to_document(
    filepath: str,
    paragraph_index: int,
    endnote_text: str
) -> Dict[str, Any]:
    """
    向Word文档的指定段落添加尾注。

    Args:
        filepath: Word 文档文件路径（.docx 格式，必须是绝对路径）
        paragraph_index: 段落索引（从0开始）
        endnote_text: 尾注文本内容

    Returns:
        包含操作结果的字典，包括：
        - message: 操作结果消息
        - file_path: 文件路径
        - paragraph_index: 段落索引
        - endnote_text: 尾注文本
        - endnote_number: 尾注编号
        - endnote_symbol: 使用的尾注符号

    Raises:
        FileError: 当文件不存在或格式不支持时
        DocumentError: 当文档操作失败时
        ValidationError: 当参数验证失败时
    """
    try:
        # 验证参数
        if not filepath:
            raise ValidationError("文件路径不能为空")

        if not endnote_text:
            raise ValidationError("尾注文本不能为空")

        # 验证段落索引
        try:
            paragraph_index = int(paragraph_index)
        except (ValueError, TypeError):
            raise ValidationError("段落索引必须是整数")

        if paragraph_index < 0:
            raise ValidationError("段落索引不能为负数")

        # 检查文件是否存在
        if not os.path.exists(filepath):
            raise FileError(f"文件不存在: {filepath}")

        # 检查文件扩展名
        if not filepath.lower().endswith('.docx'):
            raise FileError("文件格式不支持，只支持.docx格式")

        # 检查文件是否可写
        if not os.access(filepath, os.W_OK):
            raise FileError(f"文件不可写: {filepath}")

        # 打开文档
        doc = Document(filepath)

        # 验证段落索引
        if paragraph_index >= len(doc.paragraphs):
            raise ValidationError(f"段落索引无效。文档共有{len(doc.paragraphs)}个段落（索引0-{len(doc.paragraphs)-1}）")

        paragraph = doc.paragraphs[paragraph_index]

        # 查找现有尾注数量以确定编号和符号
        endnote_count = _count_existing_endnotes(doc)
        endnote_number = endnote_count + 1
        endnote_symbol = _get_endnote_symbol(endnote_number)

        try:
            # 添加尾注引用到段落
            run = paragraph.add_run()
            run.text = endnote_symbol
            run.font.superscript = True

            # 确保尾注部分存在
            _ensure_endnote_section(doc)

            # 添加尾注文本
            endnote_para = doc.add_paragraph(f"{endnote_symbol} {endnote_text}")

            # 设置尾注样式（如果存在）
            try:
                endnote_para.style = "Endnote Text"
            except KeyError:
                # 如果没有尾注样式，使用普通样式
                pass

        except Exception as e:
            raise DocumentError(f"添加尾注时发生错误: {str(e)}")

        # 保存文档
        doc.save(filepath)

        return {
            "message": f"成功向段落{paragraph_index}添加尾注",
            "file_path": filepath,
            "paragraph_index": paragraph_index,
            "endnote_text": endnote_text,
            "endnote_number": endnote_number,
            "endnote_symbol": endnote_symbol
        }

    except (FileError, ValidationError) as e:
        raise
    except Exception as e:
        raise DocumentError(f"添加尾注时发生错误: {str(e)}")


def _count_existing_endnotes(doc) -> int:
    """计算文档中现有尾注的数量"""
    count = 0

    # 查找尾注部分
    endnote_section_found = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip().lower()
        if text in ['endnotes:', 'endnotes', '尾注:', '尾注']:
            endnote_section_found = True
            continue
        elif endnote_section_found:
            # 计算尾注部分的条目
            if paragraph.text.strip():
                # 检查是否以尾注符号开头
                first_char = paragraph.text[0]
                if first_char in ['†', '‡', '§', '¶', '*', '**', '***'] or first_char.isdigit():
                    count += 1
                elif not paragraph.text.strip():
                    # 空行可能表示尾注部分结束
                    break

    return count


def _get_endnote_symbol(number: int) -> str:
    """获取尾注符号"""
    # 传统尾注符号序列
    symbols = ['†', '‡', '§', '¶']

    if number <= len(symbols):
        return symbols[number - 1]
    elif number <= len(symbols) * 2:
        # 双符号
        idx = (number - len(symbols) - 1) % len(symbols)
        return symbols[idx] * 2
    else:
        # 使用数字或星号
        return '*' * min(number - len(symbols) * 2, 5)


def _ensure_endnote_section(doc):
    """确保文档中存在尾注部分"""
    # 检查是否已存在尾注部分
    endnote_section_found = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip().lower()
        if text in ['endnotes:', 'endnotes', '尾注:', '尾注']:
            endnote_section_found = True
            break

    if not endnote_section_found:
        # 添加分页符
        doc.add_page_break()

        # 添加尾注标题
        endnote_heading = doc.add_heading("尾注:", level=1)

        # 可选：添加空行
        doc.add_paragraph("")