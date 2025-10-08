"""在指定文本附近插入标题的功能模块。"""

import os
from typing import Dict, Any
from docx import Document
from .exceptions import FileError, DocumentError, ValidationError


def insert_header_near_text(
    filepath: str,
    target_text: str,
    header_title: str,
    position: str = 'after',
    header_style: str = 'Heading 1'
) -> Dict[str, Any]:
    """
    在包含目标文本的第一个段落之前或之后插入标题。

    Args:
        filepath: Word 文档文件路径（.docx 格式，必须是绝对路径）
        target_text: 要查找的目标文本
        header_title: 要插入的标题文本
        position: 插入位置，'before' 或 'after'（默认 'after'）
        header_style: 标题样式名称（默认 'Heading 1'）

    Returns:
        包含操作结果的字典，包括：
        - message: 操作结果消息
        - file_path: 文件路径
        - target_text: 目标文本
        - header_title: 插入的标题
        - position: 插入位置
        - header_style: 使用的标题样式
        - found: 是否找到目标文本

    Raises:
        FileError: 当文件不存在或格式不支持时
        DocumentError: 当文档操作失败时
        ValidationError: 当参数验证失败时
    """
    try:
        # 验证参数
        if not filepath:
            raise ValidationError("文件路径不能为空")

        if not target_text:
            raise ValidationError("目标文本不能为空")

        if not header_title:
            raise ValidationError("标题文本不能为空")

        if position not in ['before', 'after']:
            raise ValidationError("位置参数必须是 'before' 或 'after'")

        # 检查文件是否存在
        if not os.path.exists(filepath):
            raise FileError(f"文件不存在: {filepath}")

        # 检查文件扩展名
        if not filepath.lower().endswith('.docx'):
            raise FileError("文件格式不支持，只支持.docx格式")

        # 检查文件是否可写
        if not os.access(filepath, os.W_OK):
            raise FileError(f"文件不可写: {filepath}")

        # 打开文档
        doc = Document(filepath)

        # 查找包含目标文本的段落
        found = False
        target_paragraph = None

        for i, paragraph in enumerate(doc.paragraphs):
            if target_text in paragraph.text:
                found = True
                target_paragraph = paragraph
                break

        if not found:
            return {
                "message": f"未找到包含目标文本的段落: {target_text}",
                "file_path": filepath,
                "target_text": target_text,
                "header_title": header_title,
                "position": position,
                "header_style": header_style,
                "found": False
            }

        # 创建新的标题段落
        try:
            # 尝试使用指定样式
            new_header = doc.add_paragraph(header_title)
            try:
                new_header.style = header_style
            except KeyError:
                # 如果样式不存在，使用默认标题样式
                try:
                    new_header.style = 'Heading 1'
                    header_style = 'Heading 1'
                except KeyError:
                    # 如果 Heading 1 也不存在，保持Normal样式但加粗
                    new_header.style = 'Normal'
                    for run in new_header.runs:
                        run.bold = True
                    header_style = 'Normal (加粗)'

            # 移动段落到正确位置
            if position == 'before':
                target_paragraph._element.addprevious(new_header._element)
            else:
                target_paragraph._element.addnext(new_header._element)

        except Exception as e:
            raise DocumentError(f"插入标题时发生错误: {str(e)}")

        # 保存文档
        doc.save(filepath)

        return {
            "message": f"成功在包含'{target_text}'的段落{position}插入标题'{header_title}'",
            "file_path": filepath,
            "target_text": target_text,
            "header_title": header_title,
            "position": position,
            "header_style": header_style,
            "found": True
        }

    except (FileError, ValidationError) as e:
        raise
    except Exception as e:
        raise DocumentError(f"插入标题时发生错误: {str(e)}")