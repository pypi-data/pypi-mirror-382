"""向Word文档添加标题的功能模块。"""

import os
from typing import Dict, Any
from docx import Document
from docx.shared import Pt
from .exceptions import FileError, DocumentError, ValidationError


def add_heading(
    filepath: str,
    text: str,
    level: int = 1
) -> Dict[str, Any]:
    """
    向Word文档添加一个新标题。

    Args:
        filepath: Word 文档文件路径（.docx 格式，必须是绝对路径）
        text: 标题文本内容
        level: 标题级别（1-9，其中1是最高级别）

    Returns:
        包含操作结果的字典，包括：
        - message: 操作结果消息
        - file_path: 文件路径
        - text: 添加的标题文本
        - level: 标题级别
        - style_used: 实际使用的样式
        - text_length: 文本长度
        - word_count: 词数

    Raises:
        FileError: 当文件不存在或格式不支持时
        DocumentError: 当文档操作失败时
        ValidationError: 当参数验证失败时
    """
    try:
        # 验证参数
        if not filepath:
            raise ValidationError("文件路径不能为空")

        if not text:
            raise ValidationError("标题文本不能为空")

        # 验证并转换level参数
        try:
            level = int(level)
        except (ValueError, TypeError):
            raise ValidationError("标题级别必须是整数")

        if level < 1 or level > 9:
            raise ValidationError("标题级别必须在1-9之间")

        # 检查文件是否存在
        if not os.path.exists(filepath):
            raise FileError(f"文件不存在: {filepath}")

        # 检查文件扩展名
        if not filepath.lower().endswith('.docx'):
            raise FileError("文件格式不支持，只支持.docx格式")

        # 检查文件是否可写
        if not os.access(filepath, os.W_OK):
            raise FileError(f"文件不可写: {filepath}")

        # 打开文档
        doc = Document(filepath)

        # 添加标题
        style_used = ""
        try:
            # 尝试使用内置标题样式
            heading = doc.add_heading(text, level=level)
            style_used = f"Heading {level}"

        except Exception as style_error:
            # 如果标题样式失败，使用直接格式化
            try:
                paragraph = doc.add_paragraph(text)
                paragraph.style = doc.styles['Normal']

                # 设置加粗
                for run in paragraph.runs:
                    run.bold = True

                # 根据级别设置字体大小
                if level == 1:
                    font_size = Pt(16)
                elif level == 2:
                    font_size = Pt(14)
                elif level == 3:
                    font_size = Pt(13)
                else:
                    font_size = Pt(12)

                for run in paragraph.runs:
                    run.font.size = font_size

                style_used = f"直接格式化 (级别 {level}, {font_size.pt}pt, 加粗)"

            except Exception as e:
                raise DocumentError(f"添加标题时发生错误: {str(e)}")

        # 保存文档
        doc.save(filepath)

        # 计算统计信息
        text_length = len(text)
        word_count = len(text.split()) if text else 0

        return {
            "message": f"成功向文档添加{level}级标题: {text}",
            "file_path": filepath,
            "text": text,
            "level": level,
            "style_used": style_used,
            "text_length": text_length,
            "word_count": word_count
        }

    except (FileError, ValidationError) as e:
        raise
    except Exception as e:
        raise DocumentError(f"添加标题时发生错误: {str(e)}")