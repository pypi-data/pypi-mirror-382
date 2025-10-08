"""获取Word文档结构大纲的功能模块。"""

import os
from typing import Dict, Any, List
from docx import Document
from .exceptions import FileError, DocumentError, ValidationError


def get_document_outline(filepath: str) -> Dict[str, Any]:
    """
    获取Word文档的结构大纲，包括段落和表格的结构信息。

    Args:
        filepath: Word 文档文件路径（.docx 格式，必须是绝对路径）

    Returns:
        包含文档结构的字典，包括：
        - message: 操作结果消息
        - file_path: 文件路径
        - outline: 文档结构大纲
        - summary: 结构统计摘要

    Raises:
        FileError: 当文件不存在或格式不支持时
        DocumentError: 当文档操作失败时
        ValidationError: 当参数验证失败时
    """
    try:
        # 验证文件路径
        if not filepath:
            raise ValidationError("文件路径不能为空")

        # 检查文件是否存在
        if not os.path.exists(filepath):
            raise FileError(f"文件不存在: {filepath}")

        # 检查文件扩展名
        if not filepath.lower().endswith('.docx'):
            raise FileError("文件格式不支持，只支持.docx格式")

        # 打开文档
        doc = Document(filepath)

        # 构建文档大纲
        outline = {
            "paragraphs": [],
            "tables": [],
            "headings": []
        }

        # 统计信息
        heading_count = 0
        normal_paragraph_count = 0
        table_count = 0

        # 分析段落
        for i, paragraph in enumerate(doc.paragraphs):
            text_preview = paragraph.text[:100] + ("..." if len(paragraph.text) > 100 else "")
            style_name = paragraph.style.name if paragraph.style else "Normal"

            paragraph_info = {
                "index": i,
                "text_preview": text_preview,
                "style": style_name,
                "is_heading": style_name.startswith("Heading") or style_name.startswith("标题"),
                "character_count": len(paragraph.text),
                "word_count": len(paragraph.text.split()) if paragraph.text else 0
            }

            outline["paragraphs"].append(paragraph_info)

            # 如果是标题，也添加到标题列表
            if paragraph_info["is_heading"]:
                heading_level = 1  # 默认级别
                if "Heading" in style_name:
                    try:
                        # 尝试从样式名提取级别，如 "Heading 1"
                        level_str = style_name.split()[-1]
                        heading_level = int(level_str)
                    except (ValueError, IndexError):
                        heading_level = 1

                outline["headings"].append({
                    "paragraph_index": i,
                    "level": heading_level,
                    "text": paragraph.text,
                    "style": style_name
                })
                heading_count += 1
            else:
                normal_paragraph_count += 1

        # 分析表格
        for i, table in enumerate(doc.tables):
            table_info = {
                "index": i,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "preview_data": []
            }

            # 获取表格预览数据（前3行，前3列）
            max_rows_preview = min(3, len(table.rows))
            max_cols_preview = min(3, len(table.columns))

            for row_idx in range(max_rows_preview):
                row_data = []
                for col_idx in range(max_cols_preview):
                    try:
                        cell = table.cell(row_idx, col_idx)
                        cell_text = cell.text[:30] + ("..." if len(cell.text) > 30 else "")
                        row_data.append(cell_text)
                    except IndexError:
                        row_data.append("N/A")

                table_info["preview_data"].append(row_data)

            outline["tables"].append(table_info)
            table_count += 1

        # 创建结构摘要
        summary = {
            "total_paragraphs": len(doc.paragraphs),
            "heading_count": heading_count,
            "normal_paragraph_count": normal_paragraph_count,
            "table_count": table_count,
            "has_structured_headings": heading_count > 0,
            "has_tables": table_count > 0,
            "document_sections": len(doc.sections)
        }

        return {
            "message": f"成功获取文档结构大纲: {filepath}",
            "file_path": filepath,
            "outline": outline,
            "summary": summary
        }

    except (FileError, ValidationError) as e:
        raise
    except Exception as e:
        raise DocumentError(f"获取文档结构大纲时发生错误: {str(e)}")