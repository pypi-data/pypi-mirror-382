"""在指定文本附近插入行或段落的功能模块。"""

import os
from typing import Dict, Any, Optional
from docx import Document
from .exceptions import FileError, DocumentError, ValidationError


def insert_line_or_paragraph_near_text(
    filepath: str,
    target_text: str,
    line_text: str,
    position: str = 'after',
    line_style: Optional[str] = None
) -> Dict[str, Any]:
    """
    在包含目标文本的第一个段落之前或之后插入新的行或段落。

    Args:
        filepath: Word 文档文件路径（.docx 格式，必须是绝对路径）
        target_text: 要查找的目标文本
        line_text: 要插入的行或段落文本
        position: 插入位置，'before' 或 'after'（默认 'after'）
        line_style: 段落样式名称，如果为None则使用目标段落的样式

    Returns:
        包含操作结果的字典，包括：
        - message: 操作结果消息
        - file_path: 文件路径
        - target_text: 目标文本
        - line_text: 插入的文本
        - position: 插入位置
        - line_style: 使用的样式
        - found: 是否找到目标文本

    Raises:
        FileError: 当文件不存在或格式不支持时
        DocumentError: 当文档操作失败时
        ValidationError: 当参数验证失败时
    """
    try:
        # 验证参数
        if not filepath:
            raise ValidationError("文件路径不能为空")

        if not target_text:
            raise ValidationError("目标文本不能为空")

        if not line_text:
            raise ValidationError("插入文本不能为空")

        if position not in ['before', 'after']:
            raise ValidationError("位置参数必须是 'before' 或 'after'")

        # 检查文件是否存在
        if not os.path.exists(filepath):
            raise FileError(f"文件不存在: {filepath}")

        # 检查文件扩展名
        if not filepath.lower().endswith('.docx'):
            raise FileError("文件格式不支持，只支持.docx格式")

        # 检查文件是否可写
        if not os.access(filepath, os.W_OK):
            raise FileError(f"文件不可写: {filepath}")

        # 打开文档
        doc = Document(filepath)

        # 查找包含目标文本的段落
        found = False
        target_paragraph = None

        for i, paragraph in enumerate(doc.paragraphs):
            if target_text in paragraph.text:
                found = True
                target_paragraph = paragraph
                break

        if not found:
            return {
                "message": f"未找到包含目标文本的段落: {target_text}",
                "file_path": filepath,
                "target_text": target_text,
                "line_text": line_text,
                "position": position,
                "line_style": line_style,
                "found": False
            }

        # 确定要使用的样式
        if line_style is None:
            # 使用目标段落的样式
            style_to_use = target_paragraph.style
            actual_style = target_paragraph.style.name if target_paragraph.style else 'Normal'
        else:
            style_to_use = line_style
            actual_style = line_style

        # 创建新的段落
        try:
            new_paragraph = doc.add_paragraph(line_text)

            # 设置样式
            try:
                if isinstance(style_to_use, str):
                    new_paragraph.style = style_to_use
                else:
                    new_paragraph.style = style_to_use
                    actual_style = style_to_use.name if hasattr(style_to_use, 'name') else str(style_to_use)
            except KeyError:
                # 如果样式不存在，使用Normal样式
                new_paragraph.style = 'Normal'
                actual_style = 'Normal (样式不存在，使用默认)'

            # 移动段落到正确位置
            if position == 'before':
                target_paragraph._element.addprevious(new_paragraph._element)
            else:
                target_paragraph._element.addnext(new_paragraph._element)

        except Exception as e:
            raise DocumentError(f"插入段落时发生错误: {str(e)}")

        # 保存文档
        doc.save(filepath)

        return {
            "message": f"成功在包含'{target_text}'的段落{position}插入段落'{line_text}'",
            "file_path": filepath,
            "target_text": target_text,
            "line_text": line_text,
            "position": position,
            "line_style": actual_style,
            "found": True
        }

    except (FileError, ValidationError) as e:
        raise
    except Exception as e:
        raise DocumentError(f"插入段落时发生错误: {str(e)}")