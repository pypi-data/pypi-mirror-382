"""向Word文档添加图片的功能模块。"""

import os
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Inches
from .exceptions import FileError, DocumentError, ValidationError


def add_picture(
    filepath: str,
    image_path: str,
    width: Optional[float] = None
) -> Dict[str, Any]:
    """
    向Word文档添加一张图片。

    Args:
        filepath: Word 文档文件路径（.docx 格式，必须是绝对路径）
        image_path: 图片文件路径（必须是绝对路径）
        width: 可选的图片宽度（英寸），如果不指定则使用原始大小

    Returns:
        包含操作结果的字典，包括：
        - message: 操作结果消息
        - file_path: 文件路径
        - image_path: 图片路径
        - width: 图片宽度（英寸）
        - image_size: 图片文件大小（KB）
        - image_format: 图片格式

    Raises:
        FileError: 当文件不存在或格式不支持时
        DocumentError: 当文档操作失败时
        ValidationError: 当参数验证失败时
    """
    try:
        # 验证参数
        if not filepath:
            raise ValidationError("文件路径不能为空")

        if not image_path:
            raise ValidationError("图片路径不能为空")

        # 检查文档文件是否存在
        if not os.path.exists(filepath):
            raise FileError(f"文档文件不存在: {filepath}")

        # 检查文档文件扩展名
        if not filepath.lower().endswith('.docx'):
            raise FileError("文档格式不支持，只支持.docx格式")

        # 获取绝对路径以便更好地诊断
        abs_filepath = os.path.abspath(filepath)
        abs_image_path = os.path.abspath(image_path)

        # 检查图片文件是否存在
        if not os.path.exists(abs_image_path):
            raise FileError(f"图片文件不存在: {abs_image_path}")

        # 检查图片文件大小
        try:
            image_size_bytes = os.path.getsize(abs_image_path)
            image_size_kb = image_size_bytes / 1024

            if image_size_bytes <= 0:
                raise FileError(f"图片文件为空: {abs_image_path}")

        except Exception as e:
            raise FileError(f"检查图片文件时发生错误: {str(e)}")

        # 检查图片格式（通过扩展名）
        image_ext = os.path.splitext(abs_image_path)[1].lower()
        supported_formats = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.tif']

        if image_ext not in supported_formats:
            raise FileError(f"不支持的图片格式: {image_ext}，支持的格式: {', '.join(supported_formats)}")

        # 检查文档文件是否可写
        if not os.access(abs_filepath, os.W_OK):
            raise FileError(f"文档文件不可写: {abs_filepath}")

        # 验证width参数
        if width is not None:
            try:
                width = float(width)
                if width <= 0:
                    raise ValidationError("图片宽度必须大于0")
            except (ValueError, TypeError):
                raise ValidationError("图片宽度必须是数字")

        # 打开文档
        doc = Document(abs_filepath)

        # 添加图片
        try:
            if width is not None:
                # 使用指定宽度
                doc.add_picture(abs_image_path, width=Inches(width))
                width_info = f"{width} 英寸"
            else:
                # 使用原始大小
                doc.add_picture(abs_image_path)
                width_info = "原始大小"

        except Exception as e:
            error_type = type(e).__name__
            error_msg = str(e)
            raise DocumentError(f"添加图片失败 ({error_type}): {error_msg}")

        # 保存文档
        doc.save(abs_filepath)

        return {
            "message": f"成功向文档添加图片: {os.path.basename(image_path)}",
            "file_path": filepath,
            "image_path": image_path,
            "width": width_info,
            "image_size": f"{image_size_kb:.2f} KB",
            "image_format": image_ext.upper().replace('.', '')
        }

    except (FileError, ValidationError) as e:
        raise
    except Exception as e:
        raise DocumentError(f"添加图片时发生错误: {str(e)}")