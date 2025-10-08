"""Search and replace functionality for Word documents."""

import logging
import os
from typing import Dict, Any

from docx import Document

from .exceptions import SearchReplaceError, DocumentError, FileError

logger = logging.getLogger(__name__)


def replace_placeholder_in_cell(cell, placeholder: str, replacement: str) -> None:
    """
    替换表格单元格中的占位符。

    Args:
        cell: Word 表格单元格对象
        placeholder: 要查找的占位符文本
        replacement: 替换文本
    """
    for p in cell.paragraphs:
        # 将一个段落的runs拼接起来，处理分散的占位符
        full_text = ''.join(run.text for run in p.runs)
        if placeholder in full_text:
            # 清除旧内容并添加新内容
            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = full_text.replace(placeholder, replacement)


def search_and_replace_in_document(
    filepath: str,
    replacements: Dict[str, str]
) -> Dict[str, Any]:
    """
    在 Word 文档中搜索并替换文本。

    Args:
        filepath: Word 文档文件路径
        replacements: 替换映射字典，键为要查找的文本，值为替换文本

    Returns:
        包含操作结果的字典

    Raises:
        FileError: 当文件不存在或无法访问时
        DocumentError: 当文档操作失败时
        SearchReplaceError: 当搜索替换操作失败时
    """
    try:
        # 验证是否为绝对路径
        if not os.path.isabs(filepath):
            raise FileError(f"文件路径必须是绝对路径，当前路径: {filepath}。任何时候都必须使用绝对路径。")

        # 检查文件是否存在
        if not os.path.exists(filepath):
            raise FileError(f"文件不存在: {filepath}。请确认文件路径正确，任何时候都必须使用绝对路径。")

        if not filepath.lower().endswith('.docx'):
            raise FileError(f"不支持的文件格式，仅支持 .docx 文件: {filepath}。任何时候都必须使用绝对路径。")

        # 打开文档
        try:
            doc = Document(filepath)
        except Exception as e:
            raise DocumentError(f"无法打开文档 {filepath}: {str(e)}")

        replacement_count = 0
        total_replacements = 0

        # 在普通段落中进行替换
        for paragraph in doc.paragraphs:
            for placeholder, replacement in replacements.items():
                if placeholder in paragraph.text:
                    # 计算替换次数
                    count = paragraph.text.count(placeholder)
                    replacement_count += count
                    total_replacements += count

                    # 执行替换
                    paragraph.text = paragraph.text.replace(placeholder, replacement)

        # 在表格中进行替换
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, replacement in replacements.items():
                        # 检查单元格中所有段落是否包含占位符
                        cell_has_placeholder = False
                        for paragraph in cell.paragraphs:
                            cell_text = ''.join(run.text for run in paragraph.runs)
                            if placeholder in cell_text:
                                cell_has_placeholder = True
                                # 计算替换次数
                                count = cell_text.count(placeholder)
                                replacement_count += count
                                total_replacements += count
                                break

                        # 如果找到占位符，执行替换
                        if cell_has_placeholder:
                            replace_placeholder_in_cell(cell, placeholder, replacement)

        # 保存修改后的文档
        try:
            doc.save(filepath)
        except Exception as e:
            raise DocumentError(f"无法保存文档 {filepath}: {str(e)}")

        logger.info(f"成功在文档 {filepath} 中完成 {total_replacements} 处替换")

        return {
            "message": f"成功完成搜索和替换操作",
            "file_path": filepath,
            "total_replacements": total_replacements,
            "replacements_made": dict(replacements)
        }

    except (FileError, DocumentError, SearchReplaceError) as e:
        logger.error(str(e))
        raise
    except Exception as e:
        logger.error(f"搜索替换操作失败: {str(e)}")
        raise SearchReplaceError(f"搜索替换操作失败: {str(e)}")