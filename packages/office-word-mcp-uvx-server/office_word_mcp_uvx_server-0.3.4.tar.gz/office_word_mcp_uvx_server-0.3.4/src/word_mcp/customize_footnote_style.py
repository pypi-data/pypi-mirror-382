"""自定义Word文档脚注样式的功能模块。"""

import os
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from .exceptions import FileError, DocumentError, ValidationError


def customize_footnote_style(
    filepath: str,
    numbering_format: str = "1, 2, 3",
    start_number: int = 1,
    font_name: Optional[str] = None,
    font_size: Optional[int] = None
) -> Dict[str, Any]:
    """
    自定义Word文档中脚注的编号和格式。

    Args:
        filepath: Word 文档文件路径（.docx 格式，必须是绝对路径）
        numbering_format: 编号格式（"1, 2, 3", "i, ii, iii", "a, b, c", "*, **, ***"）
        start_number: 起始编号
        font_name: 可选的字体名称
        font_size: 可选的字体大小（点数）

    Returns:
        包含操作结果的字典，包括：
        - message: 操作结果消息
        - file_path: 文件路径
        - numbering_format: 编号格式
        - start_number: 起始编号
        - footnotes_updated: 更新的脚注数量
        - style_applied: 是否应用了字体样式

    Raises:
        FileError: 当文件不存在或格式不支持时
        DocumentError: 当文档操作失败时
        ValidationError: 当参数验证失败时
    """
    try:
        # 验证参数
        if not filepath:
            raise ValidationError("文件路径不能为空")

        # 验证起始编号
        try:
            start_number = int(start_number)
            if start_number < 1:
                raise ValidationError("起始编号必须大于0")
        except (ValueError, TypeError):
            raise ValidationError("起始编号必须是整数")

        # 验证字体大小
        if font_size is not None:
            try:
                font_size = int(font_size)
                if font_size <= 0 or font_size > 200:
                    raise ValidationError("字体大小必须在1-200之间")
            except (ValueError, TypeError):
                raise ValidationError("字体大小必须是整数")

        # 验证编号格式
        valid_formats = ["1, 2, 3", "i, ii, iii", "a, b, c", "*, **, ***"]
        if numbering_format not in valid_formats:
            raise ValidationError(f"编号格式必须是以下之一: {', '.join(valid_formats)}")

        # 检查文件是否存在
        if not os.path.exists(filepath):
            raise FileError(f"文件不存在: {filepath}")

        # 检查文件扩展名
        if not filepath.lower().endswith('.docx'):
            raise FileError("文件格式不支持，只支持.docx格式")

        # 检查文件是否可写
        if not os.access(filepath, os.W_OK):
            raise FileError(f"文件不可写: {filepath}")

        # 打开文档
        doc = Document(filepath)

        try:
            # 创建或获取脚注样式
            style_applied = _create_or_update_footnote_style(doc, font_name, font_size)

            # 查找现有脚注引用
            footnote_refs = _find_footnote_references(doc)

            # 生成新的格式符号
            format_symbols = _generate_format_symbols(numbering_format, len(footnote_refs), start_number)

            # 更新脚注引用和文本
            footnotes_updated = _update_footnote_formatting(doc, footnote_refs, format_symbols, start_number)

        except Exception as e:
            raise DocumentError(f"自定义脚注样式时发生错误: {str(e)}")

        # 保存文档
        doc.save(filepath)

        return {
            "message": f"成功自定义脚注样式，更新了{footnotes_updated}个脚注",
            "file_path": filepath,
            "numbering_format": numbering_format,
            "start_number": start_number,
            "footnotes_updated": footnotes_updated,
            "style_applied": style_applied
        }

    except (FileError, ValidationError) as e:
        raise
    except Exception as e:
        raise DocumentError(f"自定义脚注样式时发生错误: {str(e)}")


def _create_or_update_footnote_style(doc, font_name: Optional[str], font_size: Optional[int]) -> bool:
    """创建或更新脚注样式"""
    try:
        footnote_style_name = "Footnote Text"

        # 尝试获取现有样式
        try:
            footnote_style = doc.styles[footnote_style_name]
        except KeyError:
            # 创建新样式
            footnote_style = doc.styles.add_style(footnote_style_name, WD_STYLE_TYPE.PARAGRAPH)

        # 应用字体设置
        if font_name:
            footnote_style.font.name = font_name

        if font_size:
            footnote_style.font.size = Pt(font_size)

        return font_name is not None or font_size is not None

    except Exception:
        return False


def _find_footnote_references(doc) -> list:
    """查找文档中的脚注引用"""
    footnote_refs = []

    for para_idx, paragraph in enumerate(doc.paragraphs):
        for run_idx, run in enumerate(paragraph.runs):
            text = run.text

            # 查找上标的脚注引用
            if run.font.superscript:
                for char in text:
                    if char.isdigit() or char in ['¹', '²', '³', '⁴', '⁵', '⁶', '⁷', '⁸', '⁹']:
                        footnote_refs.append((para_idx, run_idx, char))
                        break

    return footnote_refs


def _generate_format_symbols(numbering_format: str, count: int, start_number: int) -> list:
    """生成格式符号列表"""
    symbols = []

    for i in range(count):
        number = i + start_number

        if numbering_format == "1, 2, 3":
            symbols.append(str(number))
        elif numbering_format == "i, ii, iii":
            symbols.append(_int_to_roman(number).lower())
        elif numbering_format == "a, b, c":
            symbols.append(_int_to_letter(number).lower())
        elif numbering_format == "*, **, ***":
            symbols.append('*' * number)

    return symbols


def _int_to_roman(num: int) -> str:
    """将整数转换为罗马数字"""
    values = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    numerals = ["M", "CM", "D", "CD", "C", "XC", "L", "XL", "X", "IX", "V", "IV", "I"]

    result = ""
    for i, value in enumerate(values):
        while num >= value:
            result += numerals[i]
            num -= value
    return result


def _int_to_letter(num: int) -> str:
    """将整数转换为字母"""
    if num <= 26:
        return chr(ord('A') + num - 1)
    else:
        # 对于大于26的数字，使用双字母
        first = (num - 1) // 26
        second = (num - 1) % 26 + 1
        return chr(ord('A') + first - 1) + chr(ord('A') + second - 1)


def _update_footnote_formatting(doc, footnote_refs: list, format_symbols: list, start_number: int) -> int:
    """更新脚注格式"""
    updated_count = 0

    # 更新脚注引用
    for i, (para_idx, run_idx, old_char) in enumerate(footnote_refs):
        if i < len(format_symbols):
            try:
                paragraph = doc.paragraphs[para_idx]
                run = paragraph.runs[run_idx]

                # 替换引用符号
                symbol = format_symbols[i]
                if len(symbol) <= 3:  # 对于较短的符号使用上标
                    run.text = run.text.replace(old_char, symbol)
                    run.font.superscript = True
                else:
                    run.text = run.text.replace(old_char, symbol)
                    run.font.superscript = False

                updated_count += 1

            except (IndexError, AttributeError):
                continue

    # 更新脚注文本部分
    footnote_section_found = False
    footnote_index = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip().lower()
        if text in ['footnotes:', 'footnotes', '脚注:', '脚注']:
            footnote_section_found = True
            continue
        elif footnote_section_found and paragraph.text.strip():
            # 更新脚注文本的编号
            if footnote_index < len(format_symbols):
                symbol = format_symbols[footnote_index]

                # 提取脚注文本（去掉旧的编号）
                old_text = paragraph.text
                if ' ' in old_text:
                    footnote_content = old_text.split(' ', 1)[1]
                    paragraph.text = f"{symbol} {footnote_content}"

                footnote_index += 1

    return updated_count