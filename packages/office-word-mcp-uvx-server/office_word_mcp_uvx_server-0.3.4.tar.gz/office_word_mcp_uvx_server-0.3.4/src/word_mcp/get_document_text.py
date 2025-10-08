"""获取Word文档所有文本的功能模块。"""

import os
from typing import Dict, Any
from docx import Document
from .exceptions import FileError, DocumentError, ValidationError


def get_document_text(filepath: str) -> Dict[str, Any]:
    """
    从Word文档中提取所有文本内容，包括段落和表格中的文本。

    Args:
        filepath: Word 文档文件路径（.docx 格式，必须是绝对路径）

    Returns:
        包含文档文本的字典，包括：
        - message: 操作结果消息
        - file_path: 文件路径
        - text: 提取的所有文本内容
        - paragraph_count: 段落数量
        - table_count: 表格数量
        - total_characters: 总字符数
        - total_words: 总词数

    Raises:
        FileError: 当文件不存在或格式不支持时
        DocumentError: 当文档操作失败时
        ValidationError: 当参数验证失败时
    """
    try:
        # 验证文件路径
        if not filepath:
            raise ValidationError("文件路径不能为空")

        # 检查文件是否存在
        if not os.path.exists(filepath):
            raise FileError(f"文件不存在: {filepath}")

        # 检查文件扩展名
        if not filepath.lower().endswith('.docx'):
            raise FileError("文件格式不支持，只支持.docx格式")

        # 打开文档
        doc = Document(filepath)

        # 提取段落文本
        paragraph_texts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # 忽略空段落
                paragraph_texts.append(paragraph.text)

        # 提取表格文本
        table_texts = []
        table_count = 0
        for table in doc.tables:
            table_count += 1
            table_content = []
            for row in table.rows:
                row_content = []
                for cell in row.cells:
                    cell_text = []
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            cell_text.append(paragraph.text)
                    row_content.append(' '.join(cell_text) if cell_text else '')
                table_content.append(' | '.join(row_content))
            if table_content:
                table_texts.append('\n'.join(table_content))

        # 合并所有文本
        all_texts = []

        # 添加段落文本
        if paragraph_texts:
            all_texts.extend(paragraph_texts)

        # 添加表格文本
        if table_texts:
            all_texts.append("\n--- 表格内容 ---")
            all_texts.extend(table_texts)

        # 生成最终文本
        final_text = '\n'.join(all_texts)

        # 计算统计信息
        total_characters = len(final_text)
        total_words = len(final_text.split()) if final_text else 0

        return {
            "message": f"成功提取文档文本: {filepath}",
            "file_path": filepath,
            "text": final_text,
            "paragraph_count": len(paragraph_texts),
            "table_count": table_count,
            "total_characters": total_characters,
            "total_words": total_words
        }

    except (FileError, ValidationError) as e:
        raise
    except Exception as e:
        raise DocumentError(f"提取文档文本时发生错误: {str(e)}")