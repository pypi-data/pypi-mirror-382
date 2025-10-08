"""向Word文档添加脚注的功能模块。"""

import os
from typing import Dict, Any
from docx import Document
from .exceptions import FileError, DocumentError, ValidationError


def add_footnote_to_document(
    filepath: str,
    paragraph_index: int,
    footnote_text: str
) -> Dict[str, Any]:
    """
    向Word文档的指定段落添加脚注。

    Args:
        filepath: Word 文档文件路径（.docx 格式，必须是绝对路径）
        paragraph_index: 段落索引（从0开始）
        footnote_text: 脚注文本内容

    Returns:
        包含操作结果的字典，包括：
        - message: 操作结果消息
        - file_path: 文件路径
        - paragraph_index: 段落索引
        - footnote_text: 脚注文本
        - footnote_number: 脚注编号
        - method_used: 使用的实现方法

    Raises:
        FileError: 当文件不存在或格式不支持时
        DocumentError: 当文档操作失败时
        ValidationError: 当参数验证失败时
    """
    try:
        # 验证参数
        if not filepath:
            raise ValidationError("文件路径不能为空")

        if not footnote_text:
            raise ValidationError("脚注文本不能为空")

        # 验证段落索引
        try:
            paragraph_index = int(paragraph_index)
        except (ValueError, TypeError):
            raise ValidationError("段落索引必须是整数")

        if paragraph_index < 0:
            raise ValidationError("段落索引不能为负数")

        # 检查文件是否存在
        if not os.path.exists(filepath):
            raise FileError(f"文件不存在: {filepath}")

        # 检查文件扩展名
        if not filepath.lower().endswith('.docx'):
            raise FileError("文件格式不支持，只支持.docx格式")

        # 检查文件是否可写
        if not os.access(filepath, os.W_OK):
            raise FileError(f"文件不可写: {filepath}")

        # 打开文档
        doc = Document(filepath)

        # 验证段落索引
        if paragraph_index >= len(doc.paragraphs):
            raise ValidationError(f"段落索引无效。文档共有{len(doc.paragraphs)}个段落（索引0-{len(doc.paragraphs)-1}）")

        paragraph = doc.paragraphs[paragraph_index]

        # 查找现有脚注数量以确定编号
        footnote_count = _count_existing_footnotes(doc)
        footnote_number = footnote_count + 1

        method_used = ""

        try:
            # 尝试使用内置脚注功能（如果支持）
            try:
                # 添加脚注引用到段落末尾
                run = paragraph.add_run()
                # 尝试添加真正的脚注（某些python-docx版本可能支持）
                footnote = run.add_footnote(footnote_text)
                method_used = "内置脚注功能"

            except (AttributeError, NotImplementedError):
                # 回退到简化实现
                _add_footnote_simple(doc, paragraph, footnote_text, footnote_number)
                method_used = "简化脚注实现"

        except Exception as e:
            raise DocumentError(f"添加脚注时发生错误: {str(e)}")

        # 保存文档
        doc.save(filepath)

        return {
            "message": f"成功向段落{paragraph_index}添加脚注",
            "file_path": filepath,
            "paragraph_index": paragraph_index,
            "footnote_text": footnote_text,
            "footnote_number": footnote_number,
            "method_used": method_used
        }

    except (FileError, ValidationError) as e:
        raise
    except Exception as e:
        raise DocumentError(f"添加脚注时发生错误: {str(e)}")


def _count_existing_footnotes(doc) -> int:
    """计算文档中现有脚注的数量"""
    count = 0

    # 查找脚注标记
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text
            # 查找上标数字或脚注符号
            if run.font.superscript:
                # 简单计数上标字符
                for char in text:
                    if char.isdigit() or char in ['¹', '²', '³', '⁴', '⁵', '⁶', '⁷', '⁸', '⁹']:
                        count += 1
                        break

    # 也可以查找脚注部分中的条目
    footnote_section_found = False
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().lower() in ['footnotes:', 'footnotes', '脚注:', '脚注']:
            footnote_section_found = True
            continue
        elif footnote_section_found:
            # 计算脚注部分的条目
            if paragraph.text.strip() and (paragraph.text[0].isdigit() or paragraph.text[0] in ['¹', '²', '³', '⁴', '⁵']):
                count = max(count, _extract_footnote_number(paragraph.text))

    return count


def _extract_footnote_number(text: str) -> int:
    """从脚注文本中提取编号"""
    if not text:
        return 0

    # 查找数字编号
    if text[0].isdigit():
        try:
            # 提取开头的数字
            i = 0
            while i < len(text) and text[i].isdigit():
                i += 1
            return int(text[:i])
        except:
            return 0

    # 查找Unicode上标数字
    superscript_map = {
        '¹': 1, '²': 2, '³': 3, '⁴': 4, '⁵': 5,
        '⁶': 6, '⁷': 7, '⁸': 8, '⁹': 9, '⁰': 0
    }

    if text[0] in superscript_map:
        return superscript_map[text[0]]

    return 0


def _add_footnote_simple(doc, paragraph, footnote_text: str, footnote_number: int):
    """使用简化方法添加脚注"""
    # 添加脚注引用到段落
    run = paragraph.add_run()

    # 使用Unicode上标数字（1-9）或普通数字
    if footnote_number <= 9:
        superscript_digits = ['⁰', '¹', '²', '³', '⁴', '⁵', '⁶', '⁷', '⁸', '⁹']
        run.text = superscript_digits[footnote_number]
    else:
        run.text = str(footnote_number)
        run.font.superscript = True

    # 查找或创建脚注部分
    footnote_section_found = False
    footnote_section_paragraph = None

    for para in doc.paragraphs:
        if para.text.strip().lower() in ['footnotes:', 'footnotes', '脚注:', '脚注']:
            footnote_section_found = True
            footnote_section_paragraph = para
            break

    if not footnote_section_found:
        # 添加脚注部分
        doc.add_paragraph("")  # 空行
        footnote_section_paragraph = doc.add_paragraph("脚注:")
        footnote_section_paragraph.runs[0].bold = True

    # 添加脚注文本
    if footnote_number <= 9:
        superscript_digits = ['⁰', '¹', '²', '³', '⁴', '⁵', '⁶', '⁷', '⁸', '⁹']
        footnote_prefix = superscript_digits[footnote_number]
    else:
        footnote_prefix = str(footnote_number)

    footnote_para = doc.add_paragraph(f"{footnote_prefix} {footnote_text}")

    # 设置脚注样式（如果存在）
    try:
        footnote_para.style = "Footnote Text"
    except KeyError:
        # 如果没有脚注样式，使用普通样式但缩小字体
        try:
            for run in footnote_para.runs:
                run.font.size = run.font.size * 0.9 if run.font.size else None
        except:
            pass