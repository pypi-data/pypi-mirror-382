"""解除Word文档保护的功能模块（简化版本）。"""

import os
import json
import hashlib
from typing import Dict, Any
from docx import Document
from .exceptions import FileError, DocumentError, ValidationError


def unprotect_document(
    filepath: str,
    password: str
) -> Dict[str, Any]:
    """
    解除Word文档的密码保护（简化版本）。

    Args:
        filepath: Word 文档文件路径（.docx 格式，必须是绝对路径）
        password: 保护密码

    Returns:
        包含操作结果的字典，包括：
        - message: 操作结果消息
        - file_path: 文件路径
        - was_protected: 是否之前被保护
        - metadata_removed: 是否移除了元数据文件

    Raises:
        FileError: 当文件不存在或格式不支持时
        DocumentError: 当文档操作失败时
        ValidationError: 当参数验证失败时
    """
    try:
        # 验证参数
        if not filepath:
            raise ValidationError("文件路径不能为空")

        if not password:
            raise ValidationError("密码不能为空")

        # 检查文件是否存在
        if not os.path.exists(filepath):
            raise FileError(f"文件不存在: {filepath}")

        # 检查文件扩展名
        if not filepath.lower().endswith('.docx'):
            raise FileError("文件格式不支持，只支持.docx格式")

        # 检查文件是否可写
        if not os.access(filepath, os.W_OK):
            raise FileError(f"文件不可写: {filepath}")

        # 检查保护元数据文件
        base_path, _ = os.path.splitext(filepath)
        metadata_file = f"{base_path}.protection"

        if not os.path.exists(metadata_file):
            return {
                "message": "文档未被保护",
                "file_path": filepath,
                "was_protected": False,
                "metadata_removed": False
            }

        try:
            # 读取保护元数据
            with open(metadata_file, 'r', encoding='utf-8') as f:
                protection_data = json.load(f)

            # 验证密码
            provided_hash = hashlib.sha256(password.encode('utf-8')).hexdigest()
            stored_hash = protection_data.get('password_hash', '')

            if provided_hash != stored_hash:
                raise ValidationError("密码错误")

            # 打开文档
            doc = Document(filepath)

            # 移除保护标记（如果存在）
            paragraphs_to_remove = []
            for i, paragraph in enumerate(doc.paragraphs):
                if "此文档受密码保护" in paragraph.text:
                    paragraphs_to_remove.append(i)

            # 从后往前删除段落以避免索引变化
            for i in reversed(paragraphs_to_remove):
                try:
                    p = doc.paragraphs[i]._element
                    p.getparent().remove(p)
                except:
                    # 如果删除失败，继续
                    pass

            # 保存文档
            doc.save(filepath)

            # 删除保护元数据文件
            os.remove(metadata_file)
            metadata_removed = True

        except ValidationError:
            raise
        except Exception as e:
            raise DocumentError(f"解除文档保护时发生错误: {str(e)}")

        return {
            "message": f"成功解除文档密码保护",
            "file_path": filepath,
            "was_protected": True,
            "metadata_removed": metadata_removed
        }

    except (FileError, ValidationError) as e:
        raise
    except Exception as e:
        raise DocumentError(f"解除文档保护时发生错误: {str(e)}")