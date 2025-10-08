"""保护Word文档的功能模块（简化版本）。"""

import os
import json
import hashlib
from typing import Dict, Any
from docx import Document
from .exceptions import FileError, DocumentError, ValidationError


def protect_document(
    filepath: str,
    password: str
) -> Dict[str, Any]:
    """
    为Word文档添加密码保护（简化版本）。

    注意：这是一个简化的实现，主要用于演示。真正的加密需要专门的库。

    Args:
        filepath: Word 文档文件路径（.docx 格式，必须是绝对路径）
        password: 保护密码

    Returns:
        包含操作结果的字典，包括：
        - message: 操作结果消息
        - file_path: 文件路径
        - protection_type: 保护类型
        - password_hash: 密码哈希（用于验证）
        - metadata_file: 元数据文件路径

    Raises:
        FileError: 当文件不存在或格式不支持时
        DocumentError: 当文档操作失败时
        ValidationError: 当参数验证失败时
    """
    try:
        # 验证参数
        if not filepath:
            raise ValidationError("文件路径不能为空")

        if not password:
            raise ValidationError("密码不能为空")

        if len(password) < 4:
            raise ValidationError("密码长度至少为4个字符")

        # 检查文件是否存在
        if not os.path.exists(filepath):
            raise FileError(f"文件不存在: {filepath}")

        # 检查文件扩展名
        if not filepath.lower().endswith('.docx'):
            raise FileError("文件格式不支持，只支持.docx格式")

        # 检查文件是否可写
        if not os.access(filepath, os.W_OK):
            raise FileError(f"文件不可写: {filepath}")

        # 检查文档是否已经被保护
        base_path, _ = os.path.splitext(filepath)
        metadata_file = f"{base_path}.protection"

        if os.path.exists(metadata_file):
            raise ValidationError("文档已经被保护，请先解除保护")

        try:
            # 打开文档以验证其有效性
            doc = Document(filepath)

            # 生成密码哈希
            password_hash = hashlib.sha256(password.encode('utf-8')).hexdigest()

            # 创建保护元数据
            protection_data = {
                "type": "password",
                "password_hash": password_hash,
                "protected_at": str(os.path.getmtime(filepath)),
                "file_size": os.path.getsize(filepath),
                "method": "simplified"
            }

            # 保存保护元数据
            with open(metadata_file, 'w', encoding='utf-8') as f:
                json.dump(protection_data, f, indent=2)

            # 在文档中添加保护标记（可选）
            try:
                protection_paragraph = doc.add_paragraph()
                protection_paragraph.add_run("此文档受密码保护。").font.color.rgb = None
                protection_paragraph.runs[0].font.size = None  # 使用默认大小
                doc.save(filepath)
            except:
                # 如果添加标记失败，继续
                pass

        except Exception as e:
            # 清理可能创建的元数据文件
            if os.path.exists(metadata_file):
                try:
                    os.remove(metadata_file)
                except:
                    pass
            raise DocumentError(f"保护文档时发生错误: {str(e)}")

        return {
            "message": f"成功为文档添加密码保护",
            "file_path": filepath,
            "protection_type": "password",
            "password_hash": password_hash[:16] + "...",  # 只显示部分哈希
            "metadata_file": metadata_file
        }

    except (FileError, ValidationError) as e:
        raise
    except Exception as e:
        raise DocumentError(f"保护文档时发生错误: {str(e)}")