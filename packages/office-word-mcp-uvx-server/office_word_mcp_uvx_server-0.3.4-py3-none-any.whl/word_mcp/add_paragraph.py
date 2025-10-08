"""向Word文档添加段落的功能模块。"""

import os
from typing import Dict, Any, Optional
from docx import Document
from .exceptions import FileError, DocumentError, ValidationError


def add_paragraph(
    filepath: str,
    text: str,
    style: Optional[str] = None
) -> Dict[str, Any]:
    """
    向Word文档添加一个新段落。

    Args:
        filepath: Word 文档文件路径（.docx 格式，必须是绝对路径）
        text: 段落文本内容
        style: 可选的段落样式名称

    Returns:
        包含操作结果的字典，包括：
        - message: 操作结果消息
        - file_path: 文件路径
        - text: 添加的文本
        - style: 使用的样式
        - text_length: 文本长度
        - word_count: 词数

    Raises:
        FileError: 当文件不存在或格式不支持时
        DocumentError: 当文档操作失败时
        ValidationError: 当参数验证失败时
    """
    try:
        # 验证参数
        if not filepath:
            raise ValidationError("文件路径不能为空")

        if not text:
            raise ValidationError("段落文本不能为空")

        # 检查文件是否存在
        if not os.path.exists(filepath):
            raise FileError(f"文件不存在: {filepath}")

        # 检查文件扩展名
        if not filepath.lower().endswith('.docx'):
            raise FileError("文件格式不支持，只支持.docx格式")

        # 检查文件是否可写
        if not os.access(filepath, os.W_OK):
            raise FileError(f"文件不可写: {filepath}")

        # 打开文档
        doc = Document(filepath)

        # 添加段落
        try:
            paragraph = doc.add_paragraph(text)
            actual_style = 'Normal'  # 默认样式

            # 设置样式（如果提供）
            if style:
                try:
                    paragraph.style = style
                    actual_style = style
                except KeyError:
                    # 样式不存在，保持默认样式
                    actual_style = f"Normal ('{style}' 样式不存在)"

        except Exception as e:
            raise DocumentError(f"添加段落时发生错误: {str(e)}")

        # 保存文档
        doc.save(filepath)

        # 计算统计信息
        text_length = len(text)
        word_count = len(text.split()) if text else 0

        return {
            "message": f"成功向文档添加段落: {text[:50]}{'...' if len(text) > 50 else ''}",
            "file_path": filepath,
            "text": text,
            "style": actual_style,
            "text_length": text_length,
            "word_count": word_count
        }

    except (FileError, ValidationError) as e:
        raise
    except Exception as e:
        raise DocumentError(f"添加段落时发生错误: {str(e)}")