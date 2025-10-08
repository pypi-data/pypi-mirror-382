"""在 Word 文档中查找文本功能模块。"""

import logging
import os
import re
from typing import Dict, Any

from docx import Document

from .exceptions import DocumentError, FileError

logger = logging.getLogger(__name__)


def find_text_in_document(filepath: str, text_to_find: str, match_case: bool = True, whole_word: bool = False) -> Dict[str, Any]:
    """
    在 Word 文档中查找特定文本的出现位置。

    Args:
        filepath: Word 文档文件路径（.docx 格式，必须是绝对路径）
        text_to_find: 要查找的文本
        match_case: 是否区分大小写（默认 True）
        whole_word: 是否仅匹配完整单词（默认 False）

    Returns:
        包含查找结果的字典，包括：
        - message: 操作结果消息
        - file_path: 文件路径
        - query: 查找的文本
        - match_case: 是否区分大小写
        - whole_word: 是否仅匹配完整单词
        - occurrences: 查找结果列表
        - total_count: 总计找到的次数

    Raises:
        FileError: 当文件不存在或格式不支持时
        DocumentError: 当文档操作失败时

    Example:
        find_text_in_document(
            filepath="/path/to/document.docx",
            text_to_find="标准值",
            match_case=False,
            whole_word=True
        )
    """
    try:
        # 验证是否为绝对路径
        if not os.path.isabs(filepath):
            raise FileError(f"文件路径必须是绝对路径，当前路径: {filepath}。任何时候都必须使用绝对路径。")

        # 检查文件是否存在
        if not os.path.exists(filepath):
            raise FileError(f"文件不存在: {filepath}。请确认文件路径正确，任何时候都必须使用绝对路径。")

        if not filepath.lower().endswith('.docx'):
            raise FileError(f"不支持的文件格式，仅支持 .docx 文件: {filepath}。任何时候都必须使用绝对路径。")

        if not text_to_find:
            raise DocumentError("查找文本不能为空")

        # 打开文档
        try:
            doc = Document(filepath)
        except Exception as e:
            raise DocumentError(f"无法打开文档 {filepath}: {str(e)}")

        # 初始化结果
        results = {
            "message": "成功完成文本查找操作",
            "file_path": filepath,
            "query": text_to_find,
            "match_case": match_case,
            "whole_word": whole_word,
            "occurrences": [],
            "total_count": 0
        }

        # 搜索段落中的文本
        for i, para in enumerate(doc.paragraphs):
            if not para.text:  # 跳过空段落
                continue

            para_text = para.text
            search_text = text_to_find

            if not match_case:
                para_text = para_text.lower()
                search_text = search_text.lower()

            if whole_word:
                # 完整单词匹配：使用正则表达式
                pattern = r'\b' + re.escape(search_text) + r'\b'
                matches = list(re.finditer(pattern, para_text, re.IGNORECASE if not match_case else 0))

                for match in matches:
                    results["occurrences"].append({
                        "location_type": "paragraph",
                        "paragraph_index": i,
                        "position": match.start(),
                        "matched_text": para.text[match.start():match.end()],
                        "context": _get_context(para.text, match.start(), match.end())
                    })
                    results["total_count"] += 1
            else:
                # 子字符串匹配
                start_pos = 0
                while True:
                    pos = para_text.find(search_text, start_pos)
                    if pos == -1:
                        break

                    # 获取原始文本中的匹配部分
                    original_match = para.text[pos:pos + len(text_to_find)]

                    results["occurrences"].append({
                        "location_type": "paragraph",
                        "paragraph_index": i,
                        "position": pos,
                        "matched_text": original_match,
                        "context": _get_context(para.text, pos, pos + len(text_to_find))
                    })
                    results["total_count"] += 1
                    start_pos = pos + len(search_text)

        # 搜索表格中的文本
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        if not para.text:  # 跳过空段落
                            continue

                        para_text = para.text
                        search_text = text_to_find

                        if not match_case:
                            para_text = para_text.lower()
                            search_text = search_text.lower()

                        if whole_word:
                            # 完整单词匹配
                            pattern = r'\b' + re.escape(search_text) + r'\b'
                            matches = list(re.finditer(pattern, para_text, re.IGNORECASE if not match_case else 0))

                            for match in matches:
                                results["occurrences"].append({
                                    "location_type": "table",
                                    "location": f"表格 {table_idx + 1}, 行 {row_idx + 1}, 列 {col_idx + 1}",
                                    "table_index": table_idx,
                                    "row_index": row_idx,
                                    "column_index": col_idx,
                                    "paragraph_index": para_idx,
                                    "position": match.start(),
                                    "matched_text": para.text[match.start():match.end()],
                                    "context": _get_context(para.text, match.start(), match.end())
                                })
                                results["total_count"] += 1
                        else:
                            # 子字符串匹配
                            start_pos = 0
                            while True:
                                pos = para_text.find(search_text, start_pos)
                                if pos == -1:
                                    break

                                # 获取原始文本中的匹配部分
                                original_match = para.text[pos:pos + len(text_to_find)]

                                results["occurrences"].append({
                                    "location_type": "table",
                                    "location": f"表格 {table_idx + 1}, 行 {row_idx + 1}, 列 {col_idx + 1}",
                                    "table_index": table_idx,
                                    "row_index": row_idx,
                                    "column_index": col_idx,
                                    "paragraph_index": para_idx,
                                    "position": pos,
                                    "matched_text": original_match,
                                    "context": _get_context(para.text, pos, pos + len(text_to_find))
                                })
                                results["total_count"] += 1
                                start_pos = pos + len(search_text)

        logger.info(f"成功在文档 {filepath} 中查找文本 '{text_to_find}'，找到 {results['total_count']} 处匹配")
        return results

    except (FileError, DocumentError) as e:
        logger.error(str(e))
        raise
    except Exception as e:
        logger.error(f"查找文本失败: {str(e)}")
        raise DocumentError(f"查找文本失败: {str(e)}")


def _get_context(text: str, start_pos: int, end_pos: int, context_length: int = 50) -> str:
    """
    获取匹配文本的上下文。

    Args:
        text: 完整文本
        start_pos: 匹配开始位置
        end_pos: 匹配结束位置
        context_length: 上下文长度

    Returns:
        包含上下文的字符串
    """
    # 计算上下文范围
    context_start = max(0, start_pos - context_length)
    context_end = min(len(text), end_pos + context_length)

    # 提取上下文
    before = text[context_start:start_pos]
    match = text[start_pos:end_pos]
    after = text[end_pos:context_end]

    # 添加省略号标识
    prefix = "..." if context_start > 0 else ""
    suffix = "..." if context_end < len(text) else ""

    return f"{prefix}{before}【{match}】{after}{suffix}"