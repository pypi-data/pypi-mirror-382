"""获取Word文档信息的功能模块。"""

import os
from typing import Dict, Any
from docx import Document
from .exceptions import FileError, DocumentError, ValidationError


def get_document_info(filepath: str) -> Dict[str, Any]:
    """
    获取Word文档的详细信息，包括属性、统计信息等。

    Args:
        filepath: Word 文档文件路径（.docx 格式，必须是绝对路径）

    Returns:
        包含文档信息的字典，包括：
        - message: 操作结果消息
        - file_path: 文件路径
        - title: 文档标题
        - author: 文档作者
        - subject: 文档主题
        - keywords: 关键词
        - created: 创建时间
        - modified: 修改时间
        - last_modified_by: 最后修改者
        - revision: 修订版本
        - page_count: 页数（近似值：章节数）
        - word_count: 总字数
        - paragraph_count: 段落数
        - table_count: 表格数
        - file_size: 文件大小（字节）

    Raises:
        FileError: 当文件不存在或格式不支持时
        DocumentError: 当文档操作失败时
        ValidationError: 当参数验证失败时
    """
    try:
        # 验证文件路径
        if not filepath:
            raise ValidationError("文件路径不能为空")

        # 检查文件是否存在
        if not os.path.exists(filepath):
            raise FileError(f"文件不存在: {filepath}")

        # 检查文件扩展名
        if not filepath.lower().endswith('.docx'):
            raise FileError("文件格式不支持，只支持.docx格式")

        # 获取文件大小
        file_size = os.path.getsize(filepath)

        # 打开文档
        doc = Document(filepath)
        core_props = doc.core_properties

        # 计算字数统计
        word_count = 0
        for paragraph in doc.paragraphs:
            word_count += len(paragraph.text.split())

        # 计算表格中的字数
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        word_count += len(paragraph.text.split())

        return {
            "message": f"成功获取文档信息: {filepath}",
            "file_path": filepath,
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or "",
            "revision": core_props.revision or 0,
            "page_count": len(doc.sections),  # 近似值：章节数
            "word_count": word_count,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "file_size": file_size
        }

    except (FileError, ValidationError) as e:
        raise
    except Exception as e:
        raise DocumentError(f"获取文档信息时发生错误: {str(e)}")