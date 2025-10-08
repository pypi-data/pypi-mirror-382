"""格式化Word文档中表格的功能模块。"""

import os
from typing import Dict, Any, Optional, List
from docx import Document
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from .exceptions import FileError, DocumentError, ValidationError


def format_table(
    filepath: str,
    table_index: int,
    has_header_row: Optional[bool] = None,
    border_style: Optional[str] = None,
    shading: Optional[List[List[str]]] = None
) -> Dict[str, Any]:
    """
    格式化Word文档中的表格，包括边框、底纹和结构。

    Args:
        filepath: Word 文档文件路径（.docx 格式，必须是绝对路径）
        table_index: 表格索引（从0开始）
        has_header_row: 是否将第一行格式化为标题行
        border_style: 边框样式（'none', 'single', 'double', 'thick'）
        shading: 二维列表，指定每个单元格的背景色

    Returns:
        包含操作结果的字典，包括：
        - message: 操作结果消息
        - file_path: 文件路径
        - table_index: 表格索引
        - formatting_applied: 应用的格式选项
        - table_size: 表格大小（行x列）

    Raises:
        FileError: 当文件不存在或格式不支持时
        DocumentError: 当文档操作失败时
        ValidationError: 当参数验证失败时
    """
    try:
        # 验证参数
        if not filepath:
            raise ValidationError("文件路径不能为空")

        # 验证表格索引
        try:
            table_index = int(table_index)
        except (ValueError, TypeError):
            raise ValidationError("表格索引必须是整数")

        if table_index < 0:
            raise ValidationError("表格索引不能为负数")

        # 验证边框样式
        valid_border_styles = ['none', 'single', 'double', 'thick']
        if border_style and border_style not in valid_border_styles:
            raise ValidationError(f"边框样式必须是以下之一: {', '.join(valid_border_styles)}")

        # 检查文件是否存在
        if not os.path.exists(filepath):
            raise FileError(f"文件不存在: {filepath}")

        # 检查文件扩展名
        if not filepath.lower().endswith('.docx'):
            raise FileError("文件格式不支持，只支持.docx格式")

        # 检查文件是否可写
        if not os.access(filepath, os.W_OK):
            raise FileError(f"文件不可写: {filepath}")

        # 打开文档
        doc = Document(filepath)

        # 验证表格索引
        if table_index >= len(doc.tables):
            raise ValidationError(f"表格索引无效。文档共有{len(doc.tables)}个表格（索引0-{len(doc.tables)-1}）")

        table = doc.tables[table_index]

        # 记录格式化选项
        formatting_applied = []

        # 获取表格尺寸
        rows = len(table.rows)
        cols = len(table.columns) if rows > 0 else 0

        try:
            # 应用标题行格式
            if has_header_row and rows > 0:
                header_row = table.rows[0]
                for cell in header_row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                formatting_applied.append("标题行加粗")

            # 应用边框样式
            if border_style:
                _apply_border_style(table, border_style)
                formatting_applied.append(f"边框样式: {border_style}")

            # 应用底纹
            if shading and isinstance(shading, list):
                cells_shaded = _apply_table_shading(table, shading)
                formatting_applied.append(f"底纹应用到{cells_shaded}个单元格")

            # 设置表格样式（如果没有其他格式化）
            if not formatting_applied:
                try:
                    table.style = 'Table Grid'
                    formatting_applied.append("默认表格网格样式")
                except:
                    pass

        except Exception as e:
            raise DocumentError(f"格式化表格时发生错误: {str(e)}")

        # 保存文档
        doc.save(filepath)

        return {
            "message": f"成功格式化表格{table_index}",
            "file_path": filepath,
            "table_index": table_index,
            "formatting_applied": formatting_applied,
            "table_size": f"{rows}x{cols}"
        }

    except (FileError, ValidationError) as e:
        raise
    except Exception as e:
        raise DocumentError(f"格式化表格时发生错误: {str(e)}")


def _apply_border_style(table, border_style: str):
    """应用边框样式到表格"""
    try:
        tbl = table._tbl
        tblPr = tbl.tblPr

        # 创建表格边框元素
        tblBorders = OxmlElement('w:tblBorders')

        # 定义边框属性
        border_attrs = {}
        if border_style == 'none':
            border_attrs = {'val': 'none', 'sz': '0', 'space': '0', 'color': '000000'}
        elif border_style == 'single':
            border_attrs = {'val': 'single', 'sz': '4', 'space': '0', 'color': '000000'}
        elif border_style == 'double':
            border_attrs = {'val': 'double', 'sz': '4', 'space': '0', 'color': '000000'}
        elif border_style == 'thick':
            border_attrs = {'val': 'single', 'sz': '12', 'space': '0', 'color': '000000'}

        # 应用到所有边框
        border_types = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
        for border_type in border_types:
            border_element = OxmlElement(f'w:{border_type}')
            for attr, value in border_attrs.items():
                border_element.set(qn(f'w:{attr}'), value)
            tblBorders.append(border_element)

        # 移除现有边框并添加新边框
        existing_borders = tblPr.find(qn('w:tblBorders'))
        if existing_borders is not None:
            tblPr.remove(existing_borders)
        tblPr.append(tblBorders)

    except Exception:
        # 如果XML操作失败，回退到简单样式
        pass


def _apply_table_shading(table, shading: List[List[str]]) -> int:
    """应用底纹到表格单元格"""
    cells_shaded = 0

    # 定义颜色映射
    color_map = {
        'lightgray': 'D9D9D9',
        'lightgrey': 'D9D9D9',
        'gray': '808080',
        'grey': '808080',
        'lightblue': 'ADD8E6',
        'blue': '0000FF',
        'lightgreen': '90EE90',
        'green': '008000',
        'lightyellow': 'FFFFE0',
        'yellow': 'FFFF00',
        'lightred': 'FFB6C1',
        'red': 'FF0000',
        'white': 'FFFFFF',
        'black': '000000'
    }

    try:
        for row_idx, row_shading in enumerate(shading):
            if row_idx >= len(table.rows):
                break

            for col_idx, color in enumerate(row_shading):
                if col_idx >= len(table.rows[row_idx].cells):
                    break

                if color:
                    cell = table.rows[row_idx].cells[col_idx]

                    # 转换颜色名称为十六进制
                    if color.lower() in color_map:
                        hex_color = color_map[color.lower()]
                    elif color.startswith('#'):
                        hex_color = color[1:]  # 移除#号
                    else:
                        hex_color = color  # 假设已经是十六进制

                    # 应用底纹
                    try:
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), hex_color.upper())
                        tcPr.append(shd)
                        cells_shaded += 1
                    except:
                        # 如果XML操作失败，跳过这个单元格
                        pass

    except Exception:
        pass

    return cells_shaded