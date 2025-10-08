"""获取 Word 文档特定段落文本功能模块。"""

import logging
import os
from typing import Dict, Any

from docx import Document

from .exceptions import DocumentError, FileError

logger = logging.getLogger(__name__)


def get_paragraph_text_from_document(filepath: str, paragraph_index: int) -> Dict[str, Any]:
    """
    从 Word 文档中获取特定段落的文本。

    Args:
        filepath: Word 文档文件路径（.docx 格式，必须是绝对路径）
        paragraph_index: 段落索引（从0开始）

    Returns:
        包含段落信息的字典，包括：
        - message: 操作结果消息
        - file_path: 文件路径
        - paragraph_index: 段落索引
        - paragraph_text: 段落文本内容
        - paragraph_style: 段落样式名称
        - is_heading: 是否为标题段落
        - total_paragraphs: 文档总段落数
        - character_count: 段落字符数
        - word_count: 段落词数（按空格分割）

    Raises:
        FileError: 当文件不存在或格式不支持时
        DocumentError: 当文档操作失败时

    Example:
        get_paragraph_text_from_document(
            filepath="/path/to/document.docx",
            paragraph_index=5
        )
    """
    try:
        # 验证是否为绝对路径
        if not os.path.isabs(filepath):
            raise FileError(f"文件路径必须是绝对路径，当前路径: {filepath}。任何时候都必须使用绝对路径。")

        # 检查文件是否存在
        if not os.path.exists(filepath):
            raise FileError(f"文件不存在: {filepath}。请确认文件路径正确，任何时候都必须使用绝对路径。")

        if not filepath.lower().endswith('.docx'):
            raise FileError(f"不支持的文件格式，仅支持 .docx 文件: {filepath}。任何时候都必须使用绝对路径。")

        # 验证段落索引
        if not isinstance(paragraph_index, int):
            raise DocumentError("段落索引必须是整数类型")

        if paragraph_index < 0:
            raise DocumentError("段落索引必须是非负整数（从0开始）")

        # 打开文档
        try:
            doc = Document(filepath)
        except Exception as e:
            raise DocumentError(f"无法打开文档 {filepath}: {str(e)}")

        # 检查段落索引是否有效
        total_paragraphs = len(doc.paragraphs)
        if paragraph_index >= total_paragraphs:
            raise DocumentError(
                f"段落索引 {paragraph_index} 超出范围。文档总共有 {total_paragraphs} 个段落（索引范围: 0-{total_paragraphs-1}）"
            )

        # 获取指定段落
        paragraph = doc.paragraphs[paragraph_index]
        paragraph_text = paragraph.text

        # 获取段落样式信息
        style_name = paragraph.style.name if paragraph.style else "Normal"
        is_heading = style_name.startswith("Heading") if style_name else False

        # 计算统计信息
        character_count = len(paragraph_text)
        word_count = len(paragraph_text.split()) if paragraph_text.strip() else 0

        # 构建结果
        result = {
            "message": "成功获取段落文本",
            "file_path": filepath,
            "paragraph_index": paragraph_index,
            "paragraph_text": paragraph_text,
            "paragraph_style": style_name,
            "is_heading": is_heading,
            "total_paragraphs": total_paragraphs,
            "character_count": character_count,
            "word_count": word_count
        }

        logger.info(f"成功获取文档 {filepath} 第 {paragraph_index} 个段落的文本，长度: {character_count} 字符")
        return result

    except (FileError, DocumentError) as e:
        logger.error(str(e))
        raise
    except Exception as e:
        logger.error(f"获取段落文本失败: {str(e)}")
        raise DocumentError(f"获取段落文本失败: {str(e)}")


def get_document_paragraphs_summary(filepath: str) -> Dict[str, Any]:
    """
    获取 Word 文档所有段落的摘要信息。

    Args:
        filepath: Word 文档文件路径（.docx 格式，必须是绝对路径）

    Returns:
        包含文档段落摘要的字典

    Raises:
        FileError: 当文件不存在或格式不支持时
        DocumentError: 当文档操作失败时
    """
    try:
        # 验证是否为绝对路径
        if not os.path.isabs(filepath):
            raise FileError(f"文件路径必须是绝对路径，当前路径: {filepath}。任何时候都必须使用绝对路径。")

        # 检查文件是否存在
        if not os.path.exists(filepath):
            raise FileError(f"文件不存在: {filepath}。请确认文件路径正确，任何时候都必须使用绝对路径。")

        if not filepath.lower().endswith('.docx'):
            raise FileError(f"不支持的文件格式，仅支持 .docx 文件: {filepath}。任何时候都必须使用绝对路径。")

        # 打开文档
        try:
            doc = Document(filepath)
        except Exception as e:
            raise DocumentError(f"无法打开文档 {filepath}: {str(e)}")

        # 获取所有段落的摘要
        paragraphs_info = []
        total_characters = 0
        total_words = 0
        heading_count = 0

        for i, paragraph in enumerate(doc.paragraphs):
            paragraph_text = paragraph.text
            style_name = paragraph.style.name if paragraph.style else "Normal"
            is_heading = style_name.startswith("Heading") if style_name else False

            char_count = len(paragraph_text)
            word_count = len(paragraph_text.split()) if paragraph_text.strip() else 0

            # 只保留前100个字符作为预览
            text_preview = paragraph_text[:100] + "..." if len(paragraph_text) > 100 else paragraph_text

            paragraphs_info.append({
                "index": i,
                "text_preview": text_preview,
                "style": style_name,
                "is_heading": is_heading,
                "character_count": char_count,
                "word_count": word_count,
                "is_empty": not paragraph_text.strip()
            })

            total_characters += char_count
            total_words += word_count
            if is_heading:
                heading_count += 1

        # 构建摘要结果
        result = {
            "message": "成功获取文档段落摘要",
            "file_path": filepath,
            "total_paragraphs": len(doc.paragraphs),
            "total_characters": total_characters,
            "total_words": total_words,
            "heading_count": heading_count,
            "paragraphs": paragraphs_info
        }

        logger.info(f"成功获取文档 {filepath} 的段落摘要，共 {len(doc.paragraphs)} 个段落")
        return result

    except (FileError, DocumentError) as e:
        logger.error(str(e))
        raise
    except Exception as e:
        logger.error(f"获取段落摘要失败: {str(e)}")
        raise DocumentError(f"获取段落摘要失败: {str(e)}")