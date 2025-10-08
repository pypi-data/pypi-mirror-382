"""删除Word文档中指定段落的功能模块。"""

import os
from typing import Dict, Any
from docx import Document
from .exceptions import FileError, DocumentError, ValidationError


def delete_paragraph(
    filepath: str,
    paragraph_index: int
) -> Dict[str, Any]:
    """
    从Word文档中删除指定的段落。

    Args:
        filepath: Word 文档文件路径（.docx 格式，必须是绝对路径）
        paragraph_index: 要删除的段落索引（从0开始）

    Returns:
        包含操作结果的字典，包括：
        - message: 操作结果消息
        - file_path: 文件路径
        - paragraph_index: 删除的段落索引
        - paragraph_text: 被删除段落的文本（前50个字符）
        - total_paragraphs_before: 删除前的段落总数
        - total_paragraphs_after: 删除后的段落总数

    Raises:
        FileError: 当文件不存在或格式不支持时
        DocumentError: 当文档操作失败时
        ValidationError: 当参数验证失败时
    """
    try:
        # 验证参数
        if not filepath:
            raise ValidationError("文件路径不能为空")

        # 验证段落索引
        try:
            paragraph_index = int(paragraph_index)
        except (ValueError, TypeError):
            raise ValidationError("段落索引必须是整数")

        if paragraph_index < 0:
            raise ValidationError("段落索引不能为负数")

        # 检查文件是否存在
        if not os.path.exists(filepath):
            raise FileError(f"文件不存在: {filepath}")

        # 检查文件扩展名
        if not filepath.lower().endswith('.docx'):
            raise FileError("文件格式不支持，只支持.docx格式")

        # 检查文件是否可写
        if not os.access(filepath, os.W_OK):
            raise FileError(f"文件不可写: {filepath}")

        # 打开文档
        doc = Document(filepath)

        # 验证段落索引是否有效
        total_paragraphs_before = len(doc.paragraphs)
        if paragraph_index >= total_paragraphs_before:
            raise ValidationError(f"段落索引无效。文档共有{total_paragraphs_before}个段落（索引0-{total_paragraphs_before-1}）")

        # 获取要删除的段落信息
        target_paragraph = doc.paragraphs[paragraph_index]
        paragraph_text = target_paragraph.text[:50] + ("..." if len(target_paragraph.text) > 50 else "")

        # 删除段落
        try:
            # 获取段落的XML元素
            p = target_paragraph._element
            # 从其父元素中删除
            p.getparent().remove(p)

        except Exception as e:
            raise DocumentError(f"删除段落时发生错误: {str(e)}")

        # 保存文档
        doc.save(filepath)

        # 重新加载文档以获取准确的段落数
        doc_after = Document(filepath)
        total_paragraphs_after = len(doc_after.paragraphs)

        return {
            "message": f"成功删除段落{paragraph_index}: {paragraph_text}",
            "file_path": filepath,
            "paragraph_index": paragraph_index,
            "paragraph_text": paragraph_text,
            "total_paragraphs_before": total_paragraphs_before,
            "total_paragraphs_after": total_paragraphs_after
        }

    except (FileError, ValidationError) as e:
        raise
    except Exception as e:
        raise DocumentError(f"删除段落时发生错误: {str(e)}")