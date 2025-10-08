"""向Word文档添加表格的功能模块。"""

import os
from typing import Dict, Any, Optional, List
from docx import Document
from .exceptions import FileError, DocumentError, ValidationError


def add_table(
    filepath: str,
    rows: int,
    cols: int,
    data: Optional[List[List[str]]] = None
) -> Dict[str, Any]:
    """
    向Word文档添加一个表格。

    Args:
        filepath: Word 文档文件路径（.docx 格式，必须是绝对路径）
        rows: 表格行数
        cols: 表格列数
        data: 可选的二维数据数组来填充表格

    Returns:
        包含操作结果的字典，包括：
        - message: 操作结果消息
        - file_path: 文件路径
        - rows: 表格行数
        - cols: 表格列数
        - data_provided: 是否提供了数据
        - cells_filled: 填充的单元格数量

    Raises:
        FileError: 当文件不存在或格式不支持时
        DocumentError: 当文档操作失败时
        ValidationError: 当参数验证失败时
    """
    try:
        # 验证参数
        if not filepath:
            raise ValidationError("文件路径不能为空")

        # 验证行列数
        try:
            rows = int(rows)
            cols = int(cols)
        except (ValueError, TypeError):
            raise ValidationError("行数和列数必须是整数")

        if rows <= 0 or cols <= 0:
            raise ValidationError("行数和列数必须大于0")

        if rows > 100 or cols > 50:
            raise ValidationError("表格过大，行数不能超过100，列数不能超过50")

        # 检查文件是否存在
        if not os.path.exists(filepath):
            raise FileError(f"文件不存在: {filepath}")

        # 检查文件扩展名
        if not filepath.lower().endswith('.docx'):
            raise FileError("文件格式不支持，只支持.docx格式")

        # 检查文件是否可写
        if not os.access(filepath, os.W_OK):
            raise FileError(f"文件不可写: {filepath}")

        # 打开文档
        doc = Document(filepath)

        # 创建表格
        try:
            table = doc.add_table(rows=rows, cols=cols)

            # 尝试设置表格样式
            try:
                table.style = 'Table Grid'
            except KeyError:
                # 如果样式不存在，继续而不设置样式
                pass

        except Exception as e:
            raise DocumentError(f"创建表格时发生错误: {str(e)}")

        # 填充数据（如果提供）
        cells_filled = 0
        data_provided = False

        if data:
            data_provided = True
            try:
                for i, row_data in enumerate(data):
                    if i >= rows:
                        break  # 超出表格行数，停止填充
                    for j, cell_text in enumerate(row_data):
                        if j >= cols:
                            break  # 超出表格列数，停止填充
                        table.cell(i, j).text = str(cell_text)
                        cells_filled += 1
            except Exception as e:
                # 即使填充数据失败，表格已经创建，只记录错误
                pass

        # 保存文档
        doc.save(filepath)

        return {
            "message": f"成功向文档添加{rows}x{cols}表格",
            "file_path": filepath,
            "rows": rows,
            "cols": cols,
            "data_provided": data_provided,
            "cells_filled": cells_filled
        }

    except (FileError, ValidationError) as e:
        raise
    except Exception as e:
        raise DocumentError(f"添加表格时发生错误: {str(e)}")