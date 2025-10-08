"""格式化Word文档中指定文本的功能模块。"""

import os
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Pt, RGBColor
from .exceptions import FileError, DocumentError, ValidationError


def format_text(
    filepath: str,
    paragraph_index: int,
    start_pos: int,
    end_pos: int,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    underline: Optional[bool] = None,
    color: Optional[str] = None,
    font_size: Optional[int] = None,
    font_name: Optional[str] = None
) -> Dict[str, Any]:
    """
    格式化段落中指定范围的文本。

    Args:
        filepath: Word 文档文件路径（.docx 格式，必须是绝对路径）
        paragraph_index: 段落索引（从0开始）
        start_pos: 文本开始位置
        end_pos: 文本结束位置
        bold: 是否加粗（True/False）
        italic: 是否斜体（True/False）
        underline: 是否下划线（True/False）
        color: 文字颜色（如'red', 'blue'等）
        font_size: 字体大小（点数）
        font_name: 字体名称

    Returns:
        包含操作结果的字典，包括：
        - message: 操作结果消息
        - file_path: 文件路径
        - paragraph_index: 段落索引
        - target_text: 格式化的文本
        - format_applied: 应用的格式
        - start_pos: 开始位置
        - end_pos: 结束位置

    Raises:
        FileError: 当文件不存在或格式不支持时
        DocumentError: 当文档操作失败时
        ValidationError: 当参数验证失败时
    """
    try:
        # 验证参数
        if not filepath:
            raise ValidationError("文件路径不能为空")

        # 验证数值参数
        try:
            paragraph_index = int(paragraph_index)
            start_pos = int(start_pos)
            end_pos = int(end_pos)
            if font_size is not None:
                font_size = int(font_size)
        except (ValueError, TypeError):
            raise ValidationError("段落索引、开始位置、结束位置和字体大小必须是整数")

        if paragraph_index < 0:
            raise ValidationError("段落索引不能为负数")

        if start_pos < 0:
            raise ValidationError("开始位置不能为负数")

        if start_pos >= end_pos:
            raise ValidationError("开始位置必须小于结束位置")

        if font_size is not None and (font_size <= 0 or font_size > 200):
            raise ValidationError("字体大小必须在1-200之间")

        # 检查文件是否存在
        if not os.path.exists(filepath):
            raise FileError(f"文件不存在: {filepath}")

        # 检查文件扩展名
        if not filepath.lower().endswith('.docx'):
            raise FileError("文件格式不支持，只支持.docx格式")

        # 检查文件是否可写
        if not os.access(filepath, os.W_OK):
            raise FileError(f"文件不可写: {filepath}")

        # 打开文档
        doc = Document(filepath)

        # 验证段落索引
        if paragraph_index >= len(doc.paragraphs):
            raise ValidationError(f"段落索引无效。文档共有{len(doc.paragraphs)}个段落（索引0-{len(doc.paragraphs)-1}）")

        paragraph = doc.paragraphs[paragraph_index]
        text = paragraph.text

        # 验证文本位置
        if end_pos > len(text):
            raise ValidationError(f"结束位置无效。段落共有{len(text)}个字符")

        # 获取要格式化的文本
        target_text = text[start_pos:end_pos]

        if not target_text:
            raise ValidationError("指定位置没有文本可以格式化")

        # 格式化文本
        try:
            # 清除现有的运行并重新创建
            # 保存原始文本的三个部分：之前、目标、之后
            text_before = text[:start_pos]
            text_after = text[end_pos:]

            # 清除段落中的所有运行
            for run in paragraph.runs:
                run.clear()

            # 添加之前的文本（如果有）
            if text_before:
                paragraph.add_run(text_before)

            # 添加目标文本并应用格式
            target_run = paragraph.add_run(target_text)

            if bold is not None:
                target_run.bold = bold

            if italic is not None:
                target_run.italic = italic

            if underline is not None:
                target_run.underline = underline

            if color:
                # 定义常见颜色映射
                color_map = {
                    'red': RGBColor(255, 0, 0),
                    'blue': RGBColor(0, 0, 255),
                    'green': RGBColor(0, 128, 0),
                    'yellow': RGBColor(255, 255, 0),
                    'black': RGBColor(0, 0, 0),
                    'gray': RGBColor(128, 128, 128),
                    'grey': RGBColor(128, 128, 128),
                    'white': RGBColor(255, 255, 255),
                    'purple': RGBColor(128, 0, 128),
                    'orange': RGBColor(255, 165, 0),
                    'brown': RGBColor(165, 42, 42),
                    'pink': RGBColor(255, 192, 203)
                }

                try:
                    if color.lower() in color_map:
                        target_run.font.color.rgb = color_map[color.lower()]
                    else:
                        # 尝试从字符串解析RGB颜色
                        target_run.font.color.rgb = RGBColor.from_string(color)
                except Exception:
                    # 如果颜色设置失败，保持默认颜色
                    pass

            if font_size is not None:
                target_run.font.size = Pt(font_size)

            if font_name:
                target_run.font.name = font_name

            # 添加之后的文本（如果有）
            if text_after:
                paragraph.add_run(text_after)

        except Exception as e:
            raise DocumentError(f"格式化文本时发生错误: {str(e)}")

        # 保存文档
        doc.save(filepath)

        # 构建格式信息
        format_applied = {}
        if bold is not None:
            format_applied['bold'] = bold
        if italic is not None:
            format_applied['italic'] = italic
        if underline is not None:
            format_applied['underline'] = underline
        if color:
            format_applied['color'] = color
        if font_size is not None:
            format_applied['font_size'] = font_size
        if font_name:
            format_applied['font_name'] = font_name

        return {
            "message": f"成功格式化文本: '{target_text}'",
            "file_path": filepath,
            "paragraph_index": paragraph_index,
            "target_text": target_text,
            "format_applied": format_applied,
            "start_pos": start_pos,
            "end_pos": end_pos
        }

    except (FileError, ValidationError) as e:
        raise
    except Exception as e:
        raise DocumentError(f"格式化文本时发生错误: {str(e)}")